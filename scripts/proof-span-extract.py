#!/usr/bin/env python3
"""
Regression/Proof: PDF/DOCX → Spans (ohne Django).

  python3 scripts/proof-span-extract.py
"""
from __future__ import annotations

import importlib.util
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SERVICES = ROOT / 'Repo_abpe' / 'cv_extractor' / 'incoming' / 'services'
SAMPLES = ROOT / 'Repo_abpe' / 'cv_extractor' / 'samples'
OUT = Path('/tmp/cv_span_proof')
OUT.mkdir(parents=True, exist_ok=True)


def load(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def main() -> int:
    sys.path.insert(0, str(SERVICES.parent))
    pdf_mod = load('main_pdf_extractor', SERVICES / 'main_pdf_extractor.py')
    word_mod = load('main_word_extractor', SERVICES / 'main_word_extractor.py')
    ctrl_mod = load('main_pipeline_controller', SERVICES / 'main_pipeline_controller.py')

    results = {'ok': [], 'fail': []}

    # 1) PDF samples + controller normalize
    for p in sorted(SAMPLES.glob('*.pdf')):
        r = pdf_mod.PDFExtractor().extract(str(p))
        try:
            spans = ctrl_mod._normalize_spans(r.spans)
            assert spans, 'empty after normalize'
            assert isinstance(spans[0], dict)
            results['ok'].append(f'PDF {p.name}: {len(spans)} spans')
        except Exception as e:
            results['fail'].append(f'PDF {p.name}: {e}')

    # 2) DOCX dict spans must normalize (P0 regression)
    from docx import Document
    docx_path = OUT / 'proof_minimal.docx'
    doc = Document()
    doc.add_heading('Qualifikationsprofil Test', level=1)
    doc.add_paragraph('Max Mustermann')
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Programmiersprachen'
    table.cell(0, 1).text = 'Java, Python'
    table.cell(1, 0).text = 'Datenbanken'
    table.cell(1, 1).text = 'PostgreSQL'
    doc.save(docx_path)

    wr = word_mod.WordExtractor().extract(str(docx_path))
    if not wr.ok or not wr.spans:
        results['fail'].append(f'DOCX extract failed: ok={wr.ok} err={wr.error}')
    else:
        try:
            spans = ctrl_mod._normalize_spans(wr.spans)
            assert spans, 'DOCX normalize empty'
            results['ok'].append(
                f'DOCX normalize: {len(spans)} spans type0={type(wr.spans[0]).__name__} via={wr.meta.get("via", "python-docx")}'
            )
        except Exception as e:
            results['fail'].append(f'DOCX normalize: {e}')

        # Fallback path: table column_id when python-docx
        if wr.meta.get('via') != 'libreoffice_pdf':
            col_ids = {s.get('column_id') for s in wr.spans}
            if 0 in col_ids and 1 in col_ids:
                results['ok'].append(f'DOCX table column_id set: {sorted(col_ids)}')
            else:
                results['fail'].append(f'DOCX table column_id missing: {col_ids}')

    # 3) OCR source invariants (no tesseract binary required)
    src = (SERVICES / 'main_pdf_extractor.py').read_text(encoding='utf-8')
    if "x1=float(x + round(w))" in src or 'x1=float(x + round(w))' in src:
        results['ok'].append('OCR x1 uses line width w')
    else:
        results['fail'].append('OCR x1 still not using width w')
    if "x1=float(x + round(ln['h'] / scale))" in src:
        results['fail'].append('OCR x1 still uses height h (regression)')
    ocr_branch = src.split('if not all_spans:')[1].split('# ── Normaler Pfad')[0]
    if '_detect_columns_and_merge(ocr_spans)' in ocr_branch:
        results['ok'].append('OCR path calls column merge')
    else:
        results['fail'].append('OCR path missing column merge')

    report_path = Path('/opt/cursor/artifacts/span-proof-after-fix.json')
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding='utf-8')
    print(json.dumps(results, indent=2, ensure_ascii=False))
    print(f'\nWrote {report_path}')
    return 1 if results['fail'] else 0


if __name__ == '__main__':
    raise SystemExit(main())
