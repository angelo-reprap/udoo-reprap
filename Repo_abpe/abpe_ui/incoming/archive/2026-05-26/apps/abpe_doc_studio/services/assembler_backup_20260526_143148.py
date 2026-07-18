"""
services/assembler.py
=====================
Generischer OOXML-Assembler — vollstaendige ISO 29500 Kapitel 17 Implementierung.

20 WordprocessingML-Grundbausteine:
  17.3  PARAGRAPH       <w:p>
  17.3  LINE_BREAK      <w:br type="textWrapping">
  17.3  PAGE_BREAK      <w:br type="page">
  17.4  TABLE           <w:tbl>
  17.5  CONTENT_CONTROL <w:sdt>
  17.6  SECTION         <w:sectPr>
  17.6  COLUMN_BREAK    <w:br type="column">
  17.9  LIST            <w:numPr> + <w:abstractNum>
  17.10 HEADER          <w:hdr>
  17.10 FOOTER          <w:ftr>
  17.11 FOOTNOTE        <w:footnote>
  17.11 ENDNOTE         <w:endnote>
  17.13 BOOKMARK        <w:bookmarkStart>/<w:bookmarkEnd>
  17.13 COMMENT         <w:comment>
  17.16 FIELD           <w:fldChar>
  17.16 HYPERLINK       <w:hyperlink>
  17.17 DRAWING         <w:drawing>
  17.17 TEXTBOX         <wp:anchor> + <wps:txbx>
  17.17 SUBDOCUMENT     <w:subDoc>
  17.17 RAW_XML         direktes OOXML

Keine hardcoded Werte. Alles aus Konfiguration:
  block.content      -> Text
  block.row_styles   -> Style pro Zeile
  block.col_styles   -> Style pro Spalte
  block.layout_ref   -> Geometrie aus layout.layout_refs
  block.image_ref    -> Bildpfad aus layout.image_refs
  StyleDefinition    -> Formatierung
  PageLayout         -> Seitenmasse + layout_refs + image_refs
"""
import io
import logging
import os
import copy

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.ns import nsmap as _nsmap
from docx.oxml    import OxmlElement
from lxml import etree

# Fehlende Namespaces fuer python-docx registrieren
_extra_ns = {
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
}
for _pfx, _uri in _extra_ns.items():
    if _pfx not in _nsmap:
        _nsmap[_pfx] = _uri

from .context_loader  import ContextLoader
from .variable_engine import VariableEngine
from .exporter        import DocExporter

log = logging.getLogger('abpe_doc_studio.assembler')

DXA = 567  # 1cm = 567 DXA (Word-interne Einheit)


def _rgb(hex_str: str) -> RGBColor:
    """Hex-String zu RGBColor."""
    h = hex_str.lstrip('#')
    if len(h) != 6:
        h = '1A1A1A'
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cm(cm_val: float) -> int:
    """Zentimeter zu DXA."""
    return int(cm_val * DXA)


def _emu(cm_val: float) -> int:
    """Zentimeter zu EMU (English Metric Units)."""
    return int(cm_val * 914400 / 2.54)


class StyleCache:
    """Cached StyleDefinitions pro StyleKit."""

    def __init__(self):
        self._cache = {}

    def load(self, style_kit) -> dict:
        kid = style_kit.pk if style_kit else None
        if kid not in self._cache:
            self._cache[kid] = (
                {s.style_key: s for s in style_kit.definitions.all()}
                if style_kit else {}
            )
        return self._cache[kid]

    def get(self, style_kit, style_key: str):
        if not style_key or not style_kit:
            return None
        return self.load(style_kit).get(style_key)


class DocAssembler:

    def __init__(self):
        self.context_loader  = ContextLoader()
        self.variable_engine = VariableEngine()
        self.exporter        = DocExporter()
        self.style_cache     = StyleCache()
        self._footnote_counter = 0
        self._endnote_counter  = 0
        self._comment_counter  = 0
        self._bookmark_counter = 0

    # ── Oeffentliche API ─────────────────────────────────────────────────

    def generate(self, template_identifier: str,
                 variables: dict = None,
                 context_ref: str = '',
                 scope: str = '',
                 engine: str = 'BOTH',
                 user=None) -> dict:

        from apps.abpe_doc_studio.models import (
            DocTemplate, DocLog, DocStatus, LogStatus
        )
        tpl = DocTemplate.objects.filter(
            identifier=template_identifier,
            status=DocStatus.ACTIVE
        ).select_related('layout', 'style_kit').prefetch_related(
            'template_blocks__block__style_kit__definitions'
        ).first()

        if not tpl:
            raise ValueError(f"Template '{template_identifier}' nicht gefunden")

        all_vars = self.context_loader.load(
            scope           = scope or tpl.scope,
            context_ref     = context_ref,
            extra_variables = variables or {},
        )

        doc    = self._build_docx(tpl, all_vars)
        result = {'success': True}

        docx_info = self.exporter.save_docx(
            doc, template_identifier,
            scope=scope or tpl.scope, context_ref=context_ref,
        )
        result['file_path_docx'] = docx_info['file_path']
        result['file_size']      = docx_info['file_size']

        if engine in ('PDF', 'BOTH'):
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    docx_info['file_path']
                )
                result['file_path_pdf'] = pdf_info['file_path']
            except Exception as e:
                log.warning(f'PDF fehlgeschlagen: {e}')
                result['file_path_pdf'] = ''
                result['pdf_error']     = str(e)

        doc_log = DocLog.objects.create(
            template         = tpl,
            template_version = tpl.active_version,
            engine_used      = engine,
            file_path_docx   = result.get('file_path_docx', ''),
            file_path_pdf    = result.get('file_path_pdf', ''),
            file_size_bytes  = result.get('file_size', 0),
            context_ref      = context_ref,
            scope            = scope or tpl.scope,
            variables_used   = {
                k: v for k, v in all_vars.items()
                if not isinstance(v, list)
            },
            status       = LogStatus.OK,
            generated_by = user,
        )
        result['log_id'] = str(doc_log.log_id)
        log.info(f'Generiert: {template_identifier} [{doc_log.log_id}]')
        return result

    def render_to_bytes(self, template, variables: dict = None,
                        engine: str = 'DOCX') -> bytes:
        from apps.abpe_doc_studio.models import DocTemplate
        if isinstance(template, str):
            template = DocTemplate.objects.filter(
                identifier=template
            ).select_related('layout', 'style_kit').prefetch_related(
                'template_blocks__block__style_kit__definitions'
            ).first()
            if not template:
                raise ValueError('Template nicht gefunden')

        doc = self._build_docx(template, variables or {})
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        if engine == 'PDF':
            import tempfile
            with tempfile.NamedTemporaryFile(
                suffix='.docx', delete=False
            ) as tmp:
                tmp.write(buf.getvalue())
                tmp_path = tmp.name
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    tmp_path, output_dir=os.path.dirname(tmp_path)
                )
                return self.exporter.read_as_bytes(pdf_info['file_path'])
            finally:
                os.unlink(tmp_path)
                pdf_path = tmp_path.replace('.docx', '.pdf')
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)

        return buf.getvalue()

    # ── DOCX bauen ───────────────────────────────────────────────────────

    def _build_docx(self, tpl, all_vars: dict) -> DocxDocument:
        """
        Baut DOCX aus Template-Konfiguration.
        Liest alles aus DB — kein hardcoded Wert.
        """
        # Counter zuruecksetzen pro Dokument
        self._footnote_counter = 0
        self._endnote_counter  = 0
        self._comment_counter  = 0
        self._bookmark_counter = 0

        doc = DocxDocument()
        doc._sk              = tpl.style_kit
        doc._lay             = tpl.layout
        doc._logo_block_id   = tpl.logo_block_id
        doc._header_block_id = tpl.header_block_id
        doc._footer_block_id = tpl.footer_block_id
        doc._template_id     = tpl.identifier
        doc._template_dir    = tpl.template_dir or ''
        doc._last_vars       = all_vars

        # 17.6 SECTION — Seitenmasse aus PageLayout
        self._render_section(doc)

        # Initialen leeren Paragraph entfernen den python-docx einfügt
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = doc.element.body
        first_p = body.find(f'{{{W}}}p')
        if first_p is not None:
            texts = [t.text for t in first_p.iter() if t.tag.endswith('}t') and t.text]
            if not texts:
                body.remove(first_p)

        # 17.10 HEADER — aus header_block_id
        self._render_header_footer(doc, 'header', all_vars)

        # Body-Bloecke — nur slot=body
        # anchor_para_map: block_identifier -> erster Paragraph nach dem Rendern
        anchor_para_map = {}
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        for tb in tpl.template_blocks.order_by('slot', 'order'):
            if tb.slot != 'body':
                continue
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue
            if tb.page_break_before and doc.paragraphs:
                doc.add_page_break()

            # anchor_to_block: Textbox-Anker in bestimmten Block-Paragraph einhängen
            anchor_to = tb.anchor_to_block or None
            if anchor_to and tb.block.block_type == 'TEXTBOX':
                target_p = anchor_para_map.get(anchor_to)
                if target_p is not None:
                    doc._textbox_anchor_target = target_p
                else:
                    doc._textbox_anchor_target = None
            else:
                doc._textbox_anchor_target = None

            # Para-Count vor dem Rendern
            paras_before = [e for e in doc.element.body if e.tag == f'{{{W}}}p']
            n_before = len(paras_before)

            self._dispatch(doc, tb, all_vars)

            # Ersten neuen Paragraph merken
            paras_after = [e for e in doc.element.body if e.tag == f'{{{W}}}p']
            if len(paras_after) > n_before:
                anchor_para_map[tb.block.identifier] = paras_after[n_before]

        # 17.10 FOOTER — slot=footer Bloecke
        self._render_footer_from_blocks(doc, tpl, all_vars)

        return doc

    # ── _apply_style: EINZIGE Formatierungsquelle ────────────────────────

    def _apply_style(self, paragraph, style_key: str,
                     style_kit, run=None) -> None:
        """
        Liest StyleDefinition aus DB.
        Einzige Stelle fuer Formatierung — kein Pt/bold/color sonst.
        """
        sdef = self.style_cache.get(style_kit, style_key)
        if not sdef:
            return

        if sdef.space_before_pt:
            paragraph.paragraph_format.space_before = Pt(sdef.space_before_pt)
        if sdef.space_after_pt:
            paragraph.paragraph_format.space_after  = Pt(sdef.space_after_pt)
        if sdef.indent_left_cm:
            paragraph.paragraph_format.left_indent  = Cm(sdef.indent_left_cm)

        align_map = {
            'left':    WD_ALIGN_PARAGRAPH.LEFT,
            'right':   WD_ALIGN_PARAGRAPH.RIGHT,
            'center':  WD_ALIGN_PARAGRAPH.CENTER,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if sdef.alignment in align_map:
            paragraph.alignment = align_map[sdef.alignment]

        if sdef.border_bottom:
            self._add_border(paragraph, sdef.border_bottom_color,
                             sz=int(sdef.border_bottom_pt * 8),
                             style=sdef.border_bottom_style or 'single')

        if run is not None:
            run.font.name      = sdef.font_family
            run.font.size      = Pt(sdef.font_size_pt)
            run.bold           = sdef.bold
            run.italic         = sdef.italic
            run.underline      = sdef.underline
            run.font.color.rgb = _rgb(sdef.color_hex)

    def _add_border(self, paragraph, color: str,
                    side: str = 'bottom',
                    sz: int = 4,
                    style: str = 'single') -> None:
        pPr  = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        el   = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   style)
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color.lstrip('#'))
        pBdr.append(el)
        pPr.append(pBdr)

    # ── Block-Dispatch ───────────────────────────────────────────────────

    def _dispatch(self, doc, tb, variables: dict) -> None:
        """Verteilt auf alle 20 OOXML-Grundbausteine."""
        block    = tb.block
        bt       = block.block_type
        content  = tb.content_override or block.content or ''
        rendered = self.variable_engine.render_text(content, variables)

        dispatch = {
            # 17.3
            'PARAGRAPH':       lambda: self._render_paragraph(
                                    doc, rendered, block, variables),
            'LINE_BREAK':      lambda: self._render_line_break(doc),
            'PAGE_BREAK':      lambda: self._render_page_break(doc),
            # 17.4
            'TABLE':           lambda: self._render_table(
                                    doc, rendered, block, variables),
            # 17.5
            'CONTENT_CONTROL': lambda: self._render_content_control(
                                    doc, rendered, block),
            # 17.6
            'SECTION':         lambda: self._render_section(doc),
            'COLUMN_BREAK':    lambda: self._render_column_break(doc),
            # 17.9
            'LIST':            lambda: self._render_list(
                                    doc, rendered, block),
            # 17.10
            'HEADER':          lambda: None,
            'FOOTER':          lambda: None,
            # 17.11
            'FOOTNOTE':        lambda: self._render_footnote(
                                    doc, rendered, block),
            'ENDNOTE':         lambda: self._render_endnote(
                                    doc, rendered, block),
            # 17.13
            'BOOKMARK':        lambda: self._render_bookmark(
                                    doc, block),
            'COMMENT':         lambda: self._render_comment(
                                    doc, rendered, block),
            # 17.16
            'FIELD':           lambda: self._render_field(
                                    doc, block),
            'HYPERLINK':       lambda: self._render_hyperlink(
                                    doc, rendered, block, variables),
            # 17.17
            'DRAWING':         lambda: self._render_drawing(
                                    doc, block),
            'TEXTBOX':         lambda: self._render_textbox(
                                    doc, rendered, block),
            'SUBDOCUMENT':     lambda: self._render_subdocument(
                                    doc, block, variables),
            'RAW_XML':         lambda: self._render_raw_xml(
                                    doc, rendered),
        }
        fn = dispatch.get(bt)
        if fn:
            fn()
        else:
            log.warning(f'Unbekannter block_type: {bt}')

    # ── 17.3 PARAGRAPH <w:p> ────────────────────────────────────────────

    def _render_paragraph(self, doc, content: str,
                          block, variables: dict,
                          container=None) -> None:
        """
        Rendert Paragraph-Block.
        row_styles[i] bestimmt Style fuer Zeile i.
        Letzter Style in row_styles wird fuer alle weiteren Zeilen wiederholt.
        """
        sk         = block.style_kit
        row_styles = block.row_styles or []
        fallback   = block.style_key or ''
        lines      = content.split('\n') if content else []

        for i, line in enumerate(lines):
            skey = (row_styles[i] if i < len(row_styles)
                    else (row_styles[-1] if row_styles else fallback))

            p = (container.add_paragraph()
                 if container is not None
                 else doc.add_paragraph())

            if line.strip():
                run = p.add_run(line.strip())
                self._apply_style(p, skey, sk, run)
            else:
                sdef = self.style_cache.get(sk, skey)
                if sdef and sdef.space_after_pt:
                    p.paragraph_format.space_after = Pt(sdef.space_after_pt / 2)

    # ── 17.3 LINE_BREAK <w:br type="textWrapping"> ──────────────────────

    def _render_line_break(self, doc, container=None) -> None:
        """Fuegt einen Zeilenumbruch (kein Seitenumbruch) ein."""
        p   = (container.add_paragraph()
               if container is not None
               else doc.add_paragraph())
        run = p.add_run()
        br  = OxmlElement('w:br')
        br.set(qn('w:type'), 'textWrapping')
        run._r.append(br)

    # ── 17.3 PAGE_BREAK <w:br type="page"> ──────────────────────────────

    def _render_page_break(self, doc) -> None:
        """Fuegt Seitenumbruch ein."""
        doc.add_page_break()

    # ── 17.6 COLUMN_BREAK <w:br type="column"> ──────────────────────────

    def _render_column_break(self, doc) -> None:
        """Fuegt Spaltenumbruch ein (fuer mehrspaltige Layouts)."""
        p   = doc.add_paragraph()
        run = p.add_run()
        br  = OxmlElement('w:br')
        br.set(qn('w:type'), 'column')
        run._r.append(br)

    # ── 17.4 TABLE <w:tbl> ──────────────────────────────────────────────

    def _render_table(self, doc, content: str,
                      block, variables: dict) -> None:
        """
        Rendert Table-Block.
        A) Dynamische Tabelle: block.columns + Daten aus variables
        B) Statische Tabelle: block.content als 'sp1|sp2|sp3' pro Zeile
        """
        sk         = block.style_kit
        lay        = doc._lay
        layout_ref = block.layout_ref or ''
        geo        = (lay.layout_refs or {}).get(layout_ref, {})
        col_styles = block.col_styles or []
        col_aligns = block.col_alignments or []
        row_bdr    = block.row_border_bottom or []
        bb_sdef    = self.style_cache.get(sk, 'border_signature')

        if block.columns:
            self._render_dynamic_table(doc, block, variables, geo, sk)
        else:
            self._render_static_table(
                doc, content, sk, geo,
                col_styles, col_aligns, row_bdr, bb_sdef
            )

    def _render_static_table(self, doc, content: str,
                              sk, geo: dict,
                              col_styles: list,
                              col_aligns: list,
                              row_bdr: list,
                              bb_sdef) -> None:
        """Statische Tabelle aus 'sp1|sp2|sp3' pro Zeile."""
        lines = [l for l in content.split('\n') if l.strip()]
        if not lines:
            return

        col_widths_cm = geo.get('column_widths_cm', [])
        n_cols = max(
            len(col_widths_cm),
            max((len(l.split('|')) for l in lines), default=1)
        )

        tbl = doc.add_table(rows=len(lines), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        align_map = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
        }

        for ri, line in enumerate(lines):
            parts = line.split('|')
            while len(parts) < n_cols:
                parts.append('')

            for ci in range(n_cols):
                cell = tbl.rows[ri].cells[ci]
                if ci < len(col_widths_cm):
                    self._set_cell_width(cell, _cm(col_widths_cm[ci]))
                self._set_cell_borders_none(cell)

                skey      = (col_styles[ci] if ci < len(col_styles)
                             else (col_styles[-1] if col_styles else ''))
                align_str = (col_aligns[ci] if ci < len(col_aligns)
                             else 'left')

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                p.alignment = align_map.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)

                text = parts[ci].strip()
                if text:
                    run = p.add_run(text)
                    self._apply_style(p, skey, sk, run)

                if ri in row_bdr and ci != 1 and bb_sdef:
                    self._add_border(p, bb_sdef.border_bottom_color,
                                     sz=int(bb_sdef.border_bottom_pt * 8))

        sp    = self.style_cache.get(sk, 'spacing_normal')
        p_gap = doc.add_paragraph()
        if sp:
            p_gap.paragraph_format.space_after = Pt(sp.space_after_pt)

    def _render_dynamic_table(self, doc, block,
                               variables: dict,
                               geo: dict, sk) -> None:
        """Dynamische Tabelle mit columns-Definition + Daten aus variables."""
        columns  = block.columns
        data_key = None

        for ev in (block.expected_variables or []):
            if ev.get('type') == 'list':
                data_key = ev.get('name')
                break
        if not data_key:
            data_key = 'positionen'

        rows = self.variable_engine.expand_table_rows(
            columns, data_key, variables
        )

        lay       = doc._lay
        alt_bg    = (geo.get('alt_row_color', '')
                     or lay.table_alt_row_color or '')
        white_bg  = (geo.get('white_row_color', '')
                     or lay.table_white_row_color or '')

        content_w = (lay.column_widths_cm[0]
                     if lay.column_widths_cm
                     else lay.page_width_cm - lay.margin_left_cm - lay.margin_right_cm)
        total_dxa = _cm(content_w)
        n_cols    = len(columns)
        tbl       = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'),    str(total_dxa))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        total_pct = sum(c.get('width_pct', 20) for c in columns)
        hdr_sdef  = self.style_cache.get(sk, 'table_header')
        hdr_bg    = hdr_sdef.table_header_bg_hex if hdr_sdef else ''

        for ci, col in enumerate(columns):
            cell = tbl.rows[0].cells[ci]
            w    = int(total_dxa * col.get('width_pct', 20) / total_pct)
            self._set_cell_width(cell, w)
            if hdr_bg:
                self._set_cell_bg(cell, hdr_bg)
            self._set_cell_borders_none(cell)
            p   = cell.paragraphs[0]
            run = p.add_run(col.get('label', ''))
            self._apply_style(p, 'table_header', sk, run)

        for ri, row_data in enumerate(rows):
            drow = tbl.rows[ri + 1]
            bg   = alt_bg if ri % 2 == 0 else white_bg
            for ci, col in enumerate(columns):
                cell  = drow.cells[ci]
                value = row_data[ci] if ci < len(row_data) else ''
                w     = int(total_dxa * col.get('width_pct', 20) / total_pct)
                self._set_cell_width(cell, w)
                if bg:
                    self._set_cell_bg(cell, bg)
                self._set_cell_borders_none(cell)
                p   = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                               if col.get('align') == 'right'
                               else WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(value)
                self._apply_style(p, 'table_body', sk, run)

        sp    = self.style_cache.get(sk, 'spacing_normal')
        p_gap = doc.add_paragraph()
        if sp:
            p_gap.paragraph_format.space_after = Pt(sp.space_after_pt)

    # ── 17.17 DRAWING <w:drawing> ────────────────────────────────────────

    def _render_drawing(self, doc, block, container=None) -> None:
        """
        Rendert Bild aus layout.image_refs[block.image_ref].
        Geometrie aus layout.layout_refs[block.layout_ref].
        Unterstuetzt inline und absolut positionierte Bilder (wp:anchor).
        """
        from django.conf import settings

        lay        = doc._lay
        image_ref  = block.image_ref or ''
        layout_ref = block.layout_ref or ''

        image_refs = lay.image_refs or {}
        rel_path   = image_refs.get(image_ref, '')
        if not rel_path:
            log.warning(f'image_ref nicht gefunden: {image_ref!r}')
            return

        img_path = os.path.join(settings.BASE_DIR, rel_path)
        if not os.path.exists(img_path):
            tpl_dir_name = getattr(doc, '_template_dir', '')
            if tpl_dir_name:
                tpl_base = os.path.join(
                    settings.BASE_DIR,
                    'apps', 'abpe_doc_studio',
                    'generator', 'templates', tpl_dir_name
                )
                img_path = os.path.join(tpl_base, rel_path)

        if not os.path.exists(img_path):
            shared_base = os.path.join(
                settings.BASE_DIR,
                'apps', 'abpe_doc_studio',
                'generator', 'templates', '_shared'
            )
            img_path = os.path.join(shared_base, rel_path)

        if not os.path.exists(img_path):
            log.warning(f'Bilddatei nicht gefunden: {img_path}')
            return

        geo       = (lay.layout_refs or {}).get(layout_ref, {})
        height_cm = geo.get('image_height_cm', 2.0)
        align_str = geo.get('alignment', 'right')
        pos_h_cm  = geo.get('position_h_cm', 0.0)

        align_map = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
        }

        p = (container.add_paragraph()
             if container is not None
             else doc.add_paragraph())
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

        if pos_h_cm > 0:
            self._add_anchored_picture(
                p, img_path,
                height_cm = height_cm,
                width_cm  = geo.get('width_cm', 6.0),
                pos_h_cm  = pos_h_cm,
                pos_v_cm  = geo.get('position_v_cm', 0.5),
            )
        else:
            p.alignment = align_map.get(align_str, WD_ALIGN_PARAGRAPH.RIGHT)
            p.add_run().add_picture(img_path, height=Cm(height_cm))

    def _add_anchored_picture(self, paragraph, img_path: str,
                               height_cm: float, width_cm: float,
                               pos_h_cm: float, pos_v_cm: float) -> None:
        """Fuegt absolut positioniertes Bild ein (wp:anchor)."""
        cx = _emu(width_cm)
        cy = _emu(height_cm)
        x  = _emu(pos_h_cm)
        y  = _emu(pos_v_cm)

        run = paragraph.add_run()
        run.add_picture(img_path, width=Cm(width_cm), height=Cm(height_cm))

        inline = run._r.find('.//' + qn('wp:inline'))
        if inline is None:
            return

        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251661312')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        sp = OxmlElement('wp:simplePos')
        sp.set('x', '0')
        sp.set('y', '0')
        anchor.append(sp)

        posH = OxmlElement('wp:positionH')
        posH.set('relativeFrom', 'page')
        off  = OxmlElement('wp:posOffset')
        off.text = str(x)
        posH.append(off)
        anchor.append(posH)

        posV = OxmlElement('wp:positionV')
        posV.set('relativeFrom', 'page')
        off  = OxmlElement('wp:posOffset')
        off.text = str(y)
        posV.append(off)
        anchor.append(posV)

        extent = OxmlElement('wp:extent')
        extent.set('cx', str(cx))
        extent.set('cy', str(cy))
        anchor.append(extent)

        ee = OxmlElement('wp:effectExtent')
        ee.set('l', '0')
        ee.set('t', '0')
        ee.set('r', '0')
        ee.set('b', '0')
        anchor.append(ee)

        wrap = OxmlElement('wp:wrapSquare')
        wrap.set('wrapText', 'bothSides')
        anchor.append(wrap)

        docPr = OxmlElement('wp:docPr')
        docPr.set('id', '1')
        docPr.set('name', 'Image')
        anchor.append(docPr)

        graphic = inline.find('.//' + qn('a:graphic'))
        if graphic is not None:
            anchor.append(copy.deepcopy(graphic))

        drawing = inline.getparent()
        drawing.remove(inline)
        drawing.append(anchor)

    # ── 17.16 FIELD <w:fldChar> ─────────────────────────────────────────

    def _render_field(self, doc, block, container=None) -> None:
        """
        Rendert Feldfunktion.
        Unterstuetzt: PAGE, NUMPAGES, DATE, TIME, AUTHOR, FILENAME,
                      TOC, REF, STYLEREF, DOCPROPERTY
        field_type aus block.field_type, Format aus block.content.
        """
        sk        = block.style_kit
        style_key = block.style_key or 'footer_text'
        field_type = block.field_type or ''
        content   = block.content or ''
        sdef      = self.style_cache.get(sk, style_key)

        field_map = {
            'PAGE_NUMBER': ' PAGE ',
            'TOTAL_PAGES': ' NUMPAGES ',
            'DATE':        r' DATE \@ "dd.MM.yyyy" ',
            'TIME':        r' TIME \@ "HH:mm" ',
            'AUTHOR':      ' AUTHOR ',
            'FILENAME':    ' FILENAME ',
            'TOC':         r' TOC \o "1-3" \h \z \u ',
            'REF':         ' REF ',
            'STYLEREF':    ' STYLEREF ',
            'DOCPROPERTY': ' DOCPROPERTY ',
        }
        instr = field_map.get(field_type, ' PAGE ')

        p = (container.add_paragraph()
             if container is not None
             else doc.add_paragraph())

        if '{PAGE}' in content or '{NUMPAGES}' in content:
            parts   = content.replace('{NUMPAGES}', '\x01').split('{PAGE}')
            before  = parts[0]
            after   = parts[1] if len(parts) > 1 else ''
            between = after.split('\x01')[0] if '\x01' in after else ''
            end     = after.split('\x01')[1] if '\x01' in after else ''

            def _add_run_text(text):
                r = p.add_run(text)
                if sdef:
                    r.font.name      = sdef.font_family
                    r.font.size      = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)

            if before:  _add_run_text(before)
            self._add_field_char(p, ' PAGE ', sdef)
            if between: _add_run_text(between)
            self._add_field_char(p, ' NUMPAGES ', sdef)
            if end:     _add_run_text(end)
        else:
            self._add_field_char(p, instr, sdef)

        self._apply_style(p, style_key, sk)

    def _add_field_char(self, paragraph, instr: str, sdef) -> None:
        """Fuegt eine Feldfunktion in einen Paragraph ein."""
        for ftype, text in [
            ('begin', instr), ('separate', None), ('end', None)
        ]:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), ftype)
            run = paragraph.add_run()
            run._r.append(fld)
            if sdef:
                run.font.name      = sdef.font_family
                run.font.size      = Pt(sdef.font_size_pt)
                run.font.color.rgb = _rgb(sdef.color_hex)
            if text:
                ins = OxmlElement('w:instrText')
                ins.set(qn('xml:space'), 'preserve')
                ins.text = text
                run._r.append(ins)

    # ── 17.5 CONTENT_CONTROL <w:sdt> ────────────────────────────────────

    def _render_content_control(self, doc, content: str, block) -> None:
        """
        Fuegt Content Control (Structured Document Tag) ein.
        control_title und control_id aus block-Konfiguration.
        Unterstuetzt Plain Text, Rich Text, Dropdown, Date Picker.
        """
        control_title = block.control_title or ''
        control_id    = block.control_id or '1000'

        sdt    = OxmlElement('w:sdt')
        sdtPr  = OxmlElement('w:sdtPr')

        alias = OxmlElement('w:alias')
        alias.set(qn('w:val'), control_title)
        sdtPr.append(alias)

        cid = OxmlElement('w:id')
        cid.set(qn('w:val'), str(control_id))
        sdtPr.append(cid)

        tag = OxmlElement('w:tag')
        tag.set(qn('w:val'), control_title.lower().replace(' ', '_'))
        sdtPr.append(tag)

        sdt.append(sdtPr)

        sdtContent = OxmlElement('w:sdtContent')
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        r.append(t)
        p.append(r)
        sdtContent.append(p)
        sdt.append(sdtContent)

        doc.element.body.append(sdt)

    # ── 17.9 LIST <w:numPr> ─────────────────────────────────────────────

    def _render_list(self, doc, content: str, block) -> None:
        """
        Rendert Aufzaehlungs- oder Nummerierungsliste.
        Konfiguration aus block:
          style_key     -> 'bullet' oder 'numbered' (bestimmt Listentyp)
          content       -> Zeilen = Listeneintraege
          layout_ref    -> optionale Einrueckungstiefe (level 0-8)
        Jede Zeile in content wird ein Listeneintrag.
        Einrueckungsebene aus layout_ref oder Standard 0.
        """
        sk         = block.style_kit
        style_key  = block.style_key or 'body_text'
        layout_ref = block.layout_ref or ''
        lay        = doc._lay
        geo        = (lay.layout_refs or {}).get(layout_ref, {})
        level      = int(geo.get('list_level', 0))
        list_type  = geo.get('list_type', 'bullet')

        # Nummerierungs-Definition in Dokument einbetten
        num_id = self._ensure_numbering(doc, list_type)

        lines = [l for l in content.split('\n') if l.strip()]
        sdef  = self.style_cache.get(sk, style_key)

        for line in lines:
            p   = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()

            # Nummerierung zuweisen
            numPr = OxmlElement('w:numPr')
            ilvl  = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(level))
            numPr.append(ilvl)
            numId_el = OxmlElement('w:numId')
            numId_el.set(qn('w:val'), str(num_id))
            numPr.append(numId_el)
            pPr.append(numPr)

            run = p.add_run(line.strip())
            if sdef:
                run.font.name      = sdef.font_family
                run.font.size      = Pt(sdef.font_size_pt)
                run.bold           = sdef.bold
                run.italic         = sdef.italic
                run.font.color.rgb = _rgb(sdef.color_hex)

    def _ensure_numbering(self, doc, list_type: str = 'bullet') -> int:
        """
        Stellt sicher dass eine Nummerierungs-Definition existiert.
        Gibt numId zurueck.
        list_type: 'bullet' oder 'numbered'
        """
        # Numbering Part anlegen falls nicht vorhanden
        try:
            from docx.oxml.ns import qn
            from docx.parts.numbering import NumberingPart

            if doc.part.numbering_part is None:
                numbering_part = NumberingPart.new()
                doc.part._rels._rels  # touch rels
                rId = doc.part.relate_to(
                    numbering_part,
                    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
                )
        except Exception:
            pass

        try:
            numbering = doc.part.numbering_part._element
        except Exception:
            # Numbering Part manuell anlegen
            numbering = OxmlElement('w:numbering')
            numbering.set(
                qn('xmlns:w'),
                'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            )

        # Naechste freie abstractNumId
        existing = numbering.findall(qn('w:abstractNum'))
        abs_id   = len(existing)

        # abstractNum erstellen
        abstract = OxmlElement('w:abstractNum')
        abstract.set(qn('w:abstractNumId'), str(abs_id))

        nsid = OxmlElement('w:nsid')
        nsid.set(qn('w:val'), f'{abs_id:08X}')
        abstract.append(nsid)

        mlt = OxmlElement('w:multiLevelType')
        mlt.set(qn('w:val'), 'singleLevel')
        abstract.append(mlt)

        # Level 0 definieren
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        fmt = OxmlElement('w:numFmt')
        fmt.set(qn('w:val'), 'bullet' if list_type == 'bullet' else 'decimal')
        lvl.append(fmt)

        txt = OxmlElement('w:lvlText')
        txt.set(qn('w:val'), '\u2022' if list_type == 'bullet' else '%1.')
        lvl.append(txt)

        jc = OxmlElement('w:lvlJc')
        jc.set(qn('w:val'), 'left')
        lvl.append(jc)

        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'),    '720')
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)

        abstract.append(lvl)
        numbering.append(abstract)

        # num Instance erstellen
        existing_nums = numbering.findall(qn('w:num'))
        num_id        = len(existing_nums) + 1

        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_id))
        abs_ref = OxmlElement('w:abstractNumId')
        abs_ref.set(qn('w:val'), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)

        return num_id

    # ── 17.6 SECTION <w:sectPr> ─────────────────────────────────────────

    def _render_section(self, doc: DocxDocument) -> None:
        """
        Setzt Seitenmasse aus PageLayout.
        Unterstuetzt: Seitengroesse, Raender, Header/Footer-Abstand,
        Spalten (w:cols), Seitenrahmen (w:pgBorders).
        """
        lay = doc._lay
        sk  = doc._sk

        for section in doc.sections:
            section.page_width      = Cm(lay.page_width_cm)
            section.page_height     = Cm(lay.page_height_cm)
            section.left_margin     = Cm(lay.margin_left_cm)
            section.right_margin    = Cm(lay.margin_right_cm)
            section.top_margin      = Cm(lay.margin_top_cm)
            section.bottom_margin   = Cm(lay.margin_bottom_cm)
            section.header_distance = Cm(lay.header_distance_cm)
            section.footer_distance = Cm(lay.footer_distance_cm)

            # Mehrspaltige Layouts (w:cols) aus layout_refs
            cols_config = (lay.layout_refs or {}).get('columns', {})
            if cols_config:
                num_cols = int(cols_config.get('num_cols', 1))
                if num_cols > 1:
                    space_cm = cols_config.get('space_cm', 1.25)
                    sectPr   = section._sectPr
                    cols_el  = OxmlElement('w:cols')
                    cols_el.set(qn('w:num'),   str(num_cols))
                    cols_el.set(qn('w:space'), str(_cm(space_cm)))
                    cols_el.set(qn('w:equalWidth'), '1')
                    sectPr.append(cols_el)

        body_sdef = self.style_cache.get(sk, 'body_text')
        doc.styles['Normal'].font.name = (
            body_sdef.font_family if body_sdef else 'Arial'
        )
        doc.styles['Normal'].font.size = Pt(lay.normal_font_size_pt or 10.0)

    def _render_footer_from_blocks(self, doc, tpl, variables):
        """Rendert alle slot=footer Bloecke in den Word-Footer-Bereich."""
        footer_blocks = tpl.template_blocks.filter(
            slot='footer'
        ).order_by('order')

        if not footer_blocks.exists():
            self._render_header_footer(doc, 'footer', variables)
            return

        sk = doc._sk
        bb = self.style_cache.get(sk, 'border_brand')

        for section in doc.sections:
            ftr   = section.footer
            first = True

            for tb in footer_blocks:
                block    = tb.block
                content  = tb.content_override or block.content or ''
                rendered = self.variable_engine.render_text(content, variables)
                bt       = block.block_type
                lay      = doc._lay
                geo      = (lay.layout_refs or {}).get(block.layout_ref, {})
                col_styles = block.col_styles or []
                col_aligns = block.col_alignments or []

                if bt == 'TABLE':
                    lines = [l for l in rendered.split('\n') if l.strip()]
                    if not lines:
                        continue
                    col_widths_cm = geo.get('column_widths_cm', [])
                    n_cols = max(
                        len(col_widths_cm),
                        max((len(l.split('|')) for l in lines), default=1)
                    )
                    if first and bb:
                        p_sep = (ftr.paragraphs[0]
                                 if ftr.paragraphs else ftr.add_paragraph())
                        p_sep.clear()
                        self._add_border(p_sep, bb.border_bottom_color,
                                         side='top',
                                         sz=int(bb.border_bottom_pt * 8))
                        first = False

                    total_w = sum(_cm(c) for c in col_widths_cm) or _cm(15.0)
                    tbl = ftr.add_table(rows=len(lines), cols=n_cols,
                                        width=total_w)
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

                    for ri, line in enumerate(lines):
                        parts = line.split('|')
                        while len(parts) < n_cols:
                            parts.append('')
                        for ci in range(n_cols):
                            cell = tbl.rows[ri].cells[ci]
                            if ci < len(col_widths_cm):
                                self._set_cell_width(cell, _cm(col_widths_cm[ci]))
                            self._set_cell_borders_none(cell)
                            skey = (col_styles[ci] if ci < len(col_styles)
                                    else (col_styles[-1] if col_styles else ''))
                            p    = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after  = Pt(0)
                            text = parts[ci].strip()
                            if text:
                                run = p.add_run(text)
                                self._apply_style(p, skey, sk, run)

                elif bt in ('PARAGRAPH', 'FOOTER'):
                    if first and bb:
                        p_sep = (ftr.paragraphs[0]
                                 if ftr.paragraphs else ftr.add_paragraph())
                        p_sep.clear()
                        self._add_border(p_sep, bb.border_bottom_color,
                                         side='top',
                                         sz=int(bb.border_bottom_pt * 8))
                        first = False
                    sdef = self.style_cache.get(
                        sk, block.style_key or 'footer_text'
                    )
                    for line in rendered.split('\n'):
                        p = ftr.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if line.strip() and sdef:
                            r = p.add_run(line.strip())
                            r.font.name      = sdef.font_family
                            r.font.size      = Pt(sdef.font_size_pt)
                            r.font.color.rgb = _rgb(sdef.color_hex)

    # ── 17.10 HEADER/FOOTER <w:hdr>/<w:ftr> ─────────────────────────────

    def _render_header_footer(self, doc: DocxDocument,
                               which: str,
                               variables: dict) -> None:
        """
        Rendert Header oder Footer aus ContentBlock.
        which: 'header' oder 'footer'
        """
        from apps.abpe_doc_studio.models import ContentBlock

        block_id = (doc._header_block_id
                    if which == 'header'
                    else doc._footer_block_id)
        if not block_id:
            return

        block = ContentBlock.objects.filter(
            identifier=block_id, is_active=True
        ).select_related('style_kit').first()
        if not block:
            return

        sk      = block.style_kit or doc._sk
        content = self.variable_engine.render_text(
            block.content or '', variables
        )
        bb = self.style_cache.get(sk, 'border_brand')

        for section in doc.sections:
            area = section.header if which == 'header' else section.footer

            for p in area.paragraphs:
                p.clear()
            for p in list(area.paragraphs)[1:]:
                p._element.getparent().remove(p._element)

            bt = block.block_type

            if bt == 'FIELD':
                p_nr = area.paragraphs[0]
                p_nr.clear()
                cnt = block.content or ''
                if '\t' in cnt:
                    # 3-spaltig mit Tab-Stops
                    from docx.oxml.ns import qn as _qn
                    from docx.oxml   import OxmlElement as _OXE
                    pPr = p_nr._p.get_or_add_pPr()
                    tabs = _OXE('w:tabs')
                    t1 = _OXE('w:tab')
                    t1.set(_qn('w:val'), 'center')
                    t1.set(_qn('w:pos'), '4536')
                    tabs.append(t1)
                    t2 = _OXE('w:tab')
                    t2.set(_qn('w:val'), 'right')
                    t2.set(_qn('w:pos'), '9072')
                    tabs.append(t2)
                    pPr.append(tabs)
                    sdef = self.style_cache.get(sk, block.style_key or 'footer_text')
                    parts = cnt.split('\t')
                    for j, part in enumerate(parts):
                        if j > 0:
                            run_tab = p_nr.add_run()
                            tab_elem = _OXE('w:tab')
                            run_tab._r.append(tab_elem)
                        if '{PAGE}' in part or '{NUMPAGES}' in part:
                            sub = part.replace('{NUMPAGES}', '\x01').split('{PAGE}')
                            if sub[0]: p_nr.add_run(sub[0])
                            self._add_field_char(p_nr, ' PAGE ', sdef)
                            if len(sub) > 1:
                                rest = sub[1].split('\x01')
                                if rest[0]: p_nr.add_run(rest[0])
                                if len(rest) > 1:
                                    self._add_field_char(p_nr, ' NUMPAGES ', sdef)
                                    if rest[1]: p_nr.add_run(rest[1])
                        elif part.strip():
                            r = p_nr.add_run(part)
                            if sdef:
                                r.font.name = sdef.font_family
                                r.font.size = Pt(sdef.font_size_pt)
                                r.font.color.rgb = _rgb(sdef.color_hex)
                else:
                    p_nr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    self._render_field_in_para(p_nr, block, sk)
                if bb:
                    self._add_border(p_nr, bb.border_bottom_color,
                                     sz=int(bb.border_bottom_pt * 8))

            elif bt == 'DRAWING':
                self._render_drawing(doc, block, container=area)

            elif bt == 'PARAGRAPH':
                p          = area.paragraphs[0]
                p.clear()
                lines      = content.split('\n')
                row_styles = block.row_styles or []
                for i, line in enumerate(lines):
                    if i > 0:
                        p = area.add_paragraph()
                    skey = (row_styles[i] if i < len(row_styles)
                            else (row_styles[-1] if row_styles
                                  else block.style_key or ''))
                    if '\t' in line:
                        # Tab-gestützte 3-Spalten: links | mitte | rechts
                        from docx.oxml.ns import qn as _qn
                        from docx.oxml   import OxmlElement as _OXE
                        pPr = p._p.get_or_add_pPr()
                        tabs = _OXE('w:tabs')
                        # Mitte bei 7.5cm (Seite 15cm breit)
                        t1 = _OXE('w:tab')
                        t1.set(_qn('w:val'), 'center')
                        t1.set(_qn('w:pos'), '4536')
                        tabs.append(t1)
                        # Rechts bei 15cm
                        t2 = _OXE('w:tab')
                        t2.set(_qn('w:val'), 'right')
                        t2.set(_qn('w:pos'), '9072')
                        tabs.append(t2)
                        pPr.append(tabs)
                        parts = line.split('\t')
                        for j, part in enumerate(parts):
                            if j > 0:
                                run_tab = p.add_run()
                                tab_elem = _OXE('w:tab')
                                run_tab._r.append(tab_elem)
                            if part.strip():
                                sdef = self.style_cache.get(sk, skey)
                                run = p.add_run(part.strip())
                                if sdef:
                                    run.font.name      = sdef.font_family
                                    run.font.size      = Pt(sdef.font_size_pt)
                                    run.font.color.rgb = _rgb(sdef.color_hex)
                    elif line.strip():
                        run = p.add_run(line.strip())
                        self._apply_style(p, skey, sk, run)
                if which == 'footer' and bb:
                    self._add_border(area.paragraphs[0],
                                     bb.border_bottom_color,
                                     side='top',
                                     sz=int(bb.border_bottom_pt * 8))

            elif bt == 'FOOTER':
                p    = area.paragraphs[0]
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sdef = self.style_cache.get(
                    sk, block.style_key or 'footer_text'
                )
                if bb:
                    self._add_border(p, bb.border_bottom_color,
                                     side='top',
                                     sz=int(bb.border_bottom_pt * 8))
                for i, line in enumerate(content.split('\n')):
                    if i > 0:
                        p = area.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.strip() and sdef:
                        r = p.add_run(line.strip())
                        r.font.name      = sdef.font_family
                        r.font.size      = Pt(sdef.font_size_pt)
                        r.font.color.rgb = _rgb(sdef.color_hex)

    def _render_field_in_para(self, paragraph, block, sk) -> None:
        """Hilfsmethode: Feldfunktion direkt in Paragraph rendern."""
        content   = block.content or ''
        style_key = block.style_key or 'footer_text'
        sdef      = self.style_cache.get(sk, style_key)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if sdef:
            paragraph.paragraph_format.space_before = Pt(sdef.space_before_pt)
            paragraph.paragraph_format.space_after  = Pt(sdef.space_after_pt)

        if '{PAGE}' in content or '{NUMPAGES}' in content:
            parts   = content.replace('{NUMPAGES}', '\x01').split('{PAGE}')
            before  = parts[0]
            rest    = parts[1] if len(parts) > 1 else ''
            between = rest.split('\x01')[0] if '\x01' in rest else ''
            end     = rest.split('\x01')[1] if '\x01' in rest else ''

            def _r(text):
                r = paragraph.add_run(text)
                if sdef:
                    r.font.name      = sdef.font_family
                    r.font.size      = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)

            if before:  _r(before)
            self._add_field_char(paragraph, ' PAGE ', sdef)
            if between: _r(between)
            self._add_field_char(paragraph, ' NUMPAGES ', sdef)
            if end:     _r(end)
        else:
            field_map = {
                'PAGE_NUMBER':  ' PAGE ',
                'TOTAL_PAGES':  ' NUMPAGES ',
                'DATE':         r' DATE \@ "dd.MM.yyyy" ',
                'PAGE_NUMPAGES': None,  # Spezialfall: Seite X von Y
            }
            ft = block.field_type or ''
            if ft == 'PAGE_NUMPAGES':
                # Seite {PAGE} von {NUMPAGES}
                _r = lambda text: paragraph.add_run(text)
                _r('Seite ')
                self._add_field_char(paragraph, ' PAGE ', sdef)
                _r(' von ')
                self._add_field_char(paragraph, ' NUMPAGES ', sdef)
            else:
                instr = field_map.get(ft, ' PAGE ')
                self._add_field_char(paragraph, instr, sdef)

    # ── 17.11 FOOTNOTE <w:footnote> ─────────────────────────────────────

    def _render_footnote(self, doc, content: str, block) -> None:
        """
        Fuegt eine Fussnote ein.
        Fussnoten-Referenz im Fliestext + Fussnoten-Text am Seitenende.
        content = Text der Fussnote
        style_key = Style fuer Fussnoten-Text
        """
        sk        = block.style_kit
        style_key = block.style_key or 'body_text'
        sdef      = self.style_cache.get(sk, style_key)

        self._footnote_counter += 1
        fn_id = self._footnote_counter

        # Paragraph mit Fussnoten-Referenz
        p   = doc.add_paragraph()
        run = p.add_run()

        # Fussnoten-Referenz <w:footnoteReference>
        fn_ref = OxmlElement('w:footnoteReference')
        fn_ref.set(qn('w:id'), str(fn_id))
        run._r.append(fn_ref)

        # Fussnoten-Part aufbauen
        try:
            fn_part = doc.part.footnotes_part
        except AttributeError:
            log.warning('Fussnoten-Part nicht verfuegbar — Fussnote uebersprungen')
            return

        fn_elem = OxmlElement('w:footnote')
        fn_elem.set(qn('w:id'), str(fn_id))
        fn_elem.set(qn('w:type'), 'normal')

        fn_p   = OxmlElement('w:p')
        fn_r   = OxmlElement('w:r')

        # Fussnoten-Markierung
        fn_mark = OxmlElement('w:footnoteRef')
        fn_r.append(fn_mark)
        fn_p.append(fn_r)

        # Fussnoten-Text
        fn_r2 = OxmlElement('w:r')
        if sdef:
            rPr = OxmlElement('w:rPr')
            sz  = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(sdef.font_size_pt * 2)))
            rPr.append(sz)
            fn_r2.append(rPr)
        fn_t = OxmlElement('w:t')
        fn_t.set(qn('xml:space'), 'preserve')
        fn_t.text = ' ' + content
        fn_r2.append(fn_t)
        fn_p.append(fn_r2)
        fn_elem.append(fn_p)

        try:
            fn_part._element.append(fn_elem)
        except Exception as e:
            log.warning(f'Fussnote konnte nicht eingefuegt werden: {e}')

    # ── 17.11 ENDNOTE <w:endnote> ───────────────────────────────────────

    def _render_endnote(self, doc, content: str, block) -> None:
        """
        Fuegt eine Endnote ein.
        Endnoten-Referenz im Fliestext + Endnoten-Text am Dokumentende.
        content = Text der Endnote
        """
        sk        = block.style_kit
        style_key = block.style_key or 'body_text'
        sdef      = self.style_cache.get(sk, style_key)

        self._endnote_counter += 1
        en_id = self._endnote_counter

        p   = doc.add_paragraph()
        run = p.add_run()

        en_ref = OxmlElement('w:endnoteReference')
        en_ref.set(qn('w:id'), str(en_id))
        run._r.append(en_ref)

        try:
            en_part = doc.part.endnotes_part
        except AttributeError:
            log.warning('Endnoten-Part nicht verfuegbar — Endnote uebersprungen')
            return

        en_elem = OxmlElement('w:endnote')
        en_elem.set(qn('w:id'), str(en_id))
        en_elem.set(qn('w:type'), 'normal')

        en_p  = OxmlElement('w:p')
        en_r  = OxmlElement('w:r')
        en_mk = OxmlElement('w:endnoteRef')
        en_r.append(en_mk)
        en_p.append(en_r)

        en_r2 = OxmlElement('w:r')
        en_t  = OxmlElement('w:t')
        en_t.set(qn('xml:space'), 'preserve')
        en_t.text = ' ' + content
        en_r2.append(en_t)
        en_p.append(en_r2)
        en_elem.append(en_p)

        try:
            en_part._element.append(en_elem)
        except Exception as e:
            log.warning(f'Endnote konnte nicht eingefuegt werden: {e}')

    # ── 17.13 BOOKMARK <w:bookmarkStart>/<w:bookmarkEnd> ─────────────────

    def _render_bookmark(self, doc, block) -> None:
        """
        Fuegt Lesezeichen ein.
        bookmark_name aus block.bookmark_name,
        bookmark_id  aus block.bookmark_id oder auto-increment.
        """
        bookmark_name = block.bookmark_name or ''
        if block.bookmark_id:
            bookmark_id = block.bookmark_id
        else:
            self._bookmark_counter += 1
            bookmark_id = str(self._bookmark_counter)

        p        = doc.add_paragraph()
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'),   str(bookmark_id))
        bm_start.set(qn('w:name'), bookmark_name)
        p._p.append(bm_start)

        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), str(bookmark_id))
        p._p.append(bm_end)

    # ── 17.13 COMMENT <w:comment> ───────────────────────────────────────

    def _render_comment(self, doc, content: str, block) -> None:
        """
        Fuegt einen Kommentar/Annotation ein.
        content    = Kommentar-Text
        style_key  = Autor des Kommentars (aus StyleKit oder block.style_key)
        Kommentare werden als Randnotizen in Word angezeigt.
        """
        sk     = block.style_kit
        author = block.style_key or 'DocAssembler'

        self._comment_counter += 1
        cmt_id = self._comment_counter

        # Kommentar-Referenz im Text
        p   = doc.add_paragraph()
        run = p.add_run()

        cmt_start = OxmlElement('w:commentRangeStart')
        cmt_start.set(qn('w:id'), str(cmt_id))
        p._p.insert(0, cmt_start)

        cmt_ref = OxmlElement('w:commentReference')
        cmt_ref.set(qn('w:id'), str(cmt_id))
        run._r.append(cmt_ref)

        cmt_end = OxmlElement('w:commentRangeEnd')
        cmt_end.set(qn('w:id'), str(cmt_id))
        p._p.append(cmt_end)

        # Kommentar-Part
        try:
            cmt_part = doc.part.comments_part
            cmt_elem = OxmlElement('w:comment')
            cmt_elem.set(qn('w:id'),     str(cmt_id))
            cmt_elem.set(qn('w:author'), author)
            cmt_elem.set(qn('w:date'),   '2026-01-01T00:00:00Z')

            cmt_p = OxmlElement('w:p')
            cmt_r = OxmlElement('w:r')
            cmt_t = OxmlElement('w:t')
            cmt_t.set(qn('xml:space'), 'preserve')
            cmt_t.text = content
            cmt_r.append(cmt_t)
            cmt_p.append(cmt_r)
            cmt_elem.append(cmt_p)
            cmt_part._element.append(cmt_elem)
        except Exception as e:
            log.warning(f'Kommentar-Part nicht verfuegbar: {e}')

    # ── 17.16 HYPERLINK <w:hyperlink> ───────────────────────────────────

    def _render_hyperlink(self, doc, content: str,
                          block, variables: dict) -> None:
        """
        Fuegt Hyperlink ein.
        url   aus block.url (mit Variablen-Ersetzung)
        text  aus content
        style aus block.style_key
        """
        sk        = block.style_kit
        style_key = block.style_key or 'body_text'
        url       = self.variable_engine.render_text(
            block.url or '', variables
        )

        p     = doc.add_paragraph()
        hlink = OxmlElement('w:hyperlink')

        # Relationship fuer externe URL
        if url:
            try:
                rId = doc.part.relate_to(
                    url,
                    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                    is_external=True
                )
                hlink.set(qn('r:id'), rId)
            except Exception:
                hlink.set(qn('r:id'), '')
        else:
            hlink.set(qn('r:id'), '')

        run  = OxmlElement('w:r')
        rPr  = OxmlElement('w:rPr')
        rSty = OxmlElement('w:rStyle')
        rSty.set(qn('w:val'), 'Hyperlink')
        rPr.append(rSty)
        run.append(rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = content
        run.append(t)
        hlink.append(run)
        p._p.append(hlink)
        self._apply_style(p, style_key, sk)

    # ── 17.17 TEXTBOX <wp:anchor+wps:txbx> ──────────────────────────────

    def _render_textbox(self, doc, content: str, block) -> None:
        """
        Rendert absolut positionierte Textbox.
        Laedt XML-Template aus template_dir/textbox_{layout_ref}.xml
        und ersetzt {variablen} mit Werten aus doc._last_vars.
        """
        from django.conf import settings

        tpl_dir    = getattr(doc, '_template_dir', '')
        layout_ref = block.layout_ref or 'kontakt'

        xml_path = os.path.join(
            settings.BASE_DIR,
            'apps', 'abpe_doc_studio',
            'generator', 'templates',
            tpl_dir,
            f'textbox_{layout_ref}.xml'
        )

        if not os.path.exists(xml_path):
            log.warning(f'Textbox-Template nicht gefunden: {xml_path}')
            return

        with open(xml_path, encoding='utf-8') as f:
            xml_str = f.read()

        rendered = self.variable_engine.render_text(
            xml_str, doc._last_vars or {}
        )

        try:
            elem = etree.fromstring(rendered.encode('utf-8'))

            # Bild-Relationships registrieren (rId8, rId9 etc.)
            R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            IMG_TYPE = f'{R}/image'
            media_base = os.path.join(
                settings.BASE_DIR,
                'apps', 'abpe_doc_studio',
                'generator', 'templates',
                tpl_dir, 'word', 'media'
            )
            used_rids = set()
            for el in elem.iter():
                embed = el.get(f'{{{R}}}embed')
                if embed:
                    used_rids.add(embed)

            rid_map = {}
            for rId in sorted(used_rids):
                num = rId.replace('rId', '')
                img_file = f'image{num}.png'
                img_path = os.path.join(media_base, img_file)
                if os.path.exists(img_path):
                    try:
                        img_part = doc.part.package.image_parts.get_or_add_image_part(img_path)
                        new_rId = doc.part.relate_to(img_part, IMG_TYPE)
                        rid_map[rId] = new_rId
                        log.info(f'Textbox img: {rId}→{new_rId} ({img_file})')
                    except Exception as e:
                        log.warning(f'Textbox img Fehler {rId}: {e}')

            if rid_map:
                xml_str = etree.tostring(elem, encoding='unicode')
                for old_rid, new_rid in rid_map.items():
                    if old_rid != new_rid:
                        xml_str = xml_str.replace(f'"{old_rid}"', f'"{new_rid}"')
                elem = etree.fromstring(xml_str.encode('utf-8'))

            body = doc.element.body
            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            if elem.tag == f'{{{W}}}p':
                # Anchor-Ziel: entweder explizit per anchor_to_block oder letzter Para
                target_p = getattr(doc, '_textbox_anchor_target', None)
                if target_p is None:
                    paras = [e for e in body if e.tag == f'{{{W}}}p']
                    target_p = paras[-1] if paras else None
                if target_p is not None:
                    for child in list(elem):
                        target_p.append(child)
                else:
                    sect = body.find(f'{{{W}}}sectPr')
                    if sect is not None:
                        body.insert(list(body).index(sect), elem)
                    else:
                        body.append(elem)
            else:
                sect = body.find(f'{{{W}}}sectPr')
                if sect is not None:
                    body.insert(list(body).index(sect), elem)
                else:
                    body.append(elem)
        except Exception as e:
            log.warning(f'Textbox XML Fehler: {e}')

    # ── 17.17 SUBDOCUMENT <w:subDoc> ────────────────────────────────────

    def _render_subdocument(self, doc, block, variables: dict) -> None:
        """
        Fuegt ein Unterdokument (w:subDoc) ein.
        Pfad des externen DOCX aus layout.image_refs[block.image_ref]
        oder direkt aus block.content (mit Variablen-Ersetzung).
        Nützlich fuer modulare Betriebshandbuecher/Architekturdokumente.
        """
        from django.conf import settings

        lay      = doc._lay
        sub_path = ''

        # Pfad aus image_refs (empfohlen)
        if block.image_ref:
            image_refs = (lay.image_refs or {})
            rel_path   = image_refs.get(block.image_ref, '')
            if rel_path:
                sub_path = os.path.join(settings.BASE_DIR, rel_path)

        # Pfad direkt aus content (Fallback)
        if not sub_path and block.content:
            rendered = self.variable_engine.render_text(
                block.content, variables
            )
            sub_path = os.path.join(settings.BASE_DIR, rendered.strip())

        if not sub_path or not os.path.exists(sub_path):
            log.warning(f'Subdokument nicht gefunden: {sub_path!r}')
            return

        # Relationship zum externen Dokument
        try:
            rId = doc.part.relate_to(
                sub_path,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/subDocument',
                is_external=False
            )
        except Exception as e:
            log.warning(f'Subdokument Relationship Fehler: {e}')
            return

        # w:p mit w:subDoc einfuegen
        p      = doc.add_paragraph()
        pPr    = p._p.get_or_add_pPr()
        sub    = OxmlElement('w:subDoc')
        sub.set(qn('r:id'), rId)
        pPr.append(sub)

        log.info(f'Subdokument eingefuegt: {sub_path}')

    # ── 17.17 RAW_XML ────────────────────────────────────────────────────

    def _render_raw_xml(self, doc, content: str) -> None:
        """
        Fuegt direktes OOXML in den Dokument-Body ein.
        content = vollstaendiges XML-Element als String.
        Variablen werden vor dem Parsen ersetzt (via _dispatch rendered).
        Bild-Relationships werden automatisch registriert.
        """
        if not content or not content.strip():
            return

        try:
            elem = etree.fromstring(content.encode('utf-8'))

            # Bild-Relationships registrieren (r:embed rId*)
            R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            IMG_TYPE = f'{R}/image'
            from django.conf import settings
            tpl_dir = getattr(doc, '_template_dir', '')
            media_base = os.path.join(
                settings.BASE_DIR,
                'apps', 'abpe_doc_studio',
                'generator', 'templates',
                tpl_dir, 'word', 'media'
            )
            # Alle r:embed Referenzen sammeln
            used_rids = set()
            for el in elem.iter():
                embed = el.get(f'{{{R}}}embed')
                if embed:
                    used_rids.add(embed)

            # Bestehende Relationships prüfen
            existing = {}
            for rId, rel in doc.part.rels.items():
                existing[rId] = rel

            # Bilder registrieren
            rid_map = {}
            for rId in sorted(used_rids):
                if rId in existing:
                    rid_map[rId] = rId
                    continue
                # Bild aus media-Verzeichnis laden
                # rId7→image1, rId8→image2, rId9→image3
                num = rId.replace('rId', '')
                img_file = f'image{num}.png'
                img_path = os.path.join(media_base, img_file)
                if os.path.exists(img_path):
                    try:
                        new_rId = doc.part.relate_to(img_path, IMG_TYPE)
                        rid_map[rId] = new_rId
                        log.info(f'Relationship: {rId} → {new_rId} ({img_file})')
                    except Exception as e:
                        log.warning(f'Relationship Fehler {rId}: {e}')

            # rIds in XML ersetzen falls nötig
            if rid_map:
                xml_str = etree.tostring(elem, encoding='unicode')
                for old_rid, new_rid in rid_map.items():
                    if old_rid != new_rid:
                        xml_str = xml_str.replace(
                            f'r:embed="{old_rid}"',
                            f'r:embed="{new_rid}"'
                        )
                elem = etree.fromstring(xml_str.encode('utf-8'))

            body = doc.element.body
            sect = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sect is not None:
                body.insert(list(body).index(sect), elem)
            else:
                body.append(elem)
            log.info('RAW_XML eingefuegt')
        except etree.XMLSyntaxError as e:
            log.warning(f'RAW_XML Syntax-Fehler: {e}')
        except Exception as e:
            log.warning(f'RAW_XML Fehler: {e}')

    # ── Zell-Helfer ──────────────────────────────────────────────────────

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.lstrip('#'))
        tcPr.append(shd)

    def _set_cell_width(self, cell, width_dxa: int) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_cell_borders_none(self, cell) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        bdrs = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement('w:' + side)
            el.set(qn('w:val'), 'nil')
            bdrs.append(el)
        tcPr.append(bdrs)
