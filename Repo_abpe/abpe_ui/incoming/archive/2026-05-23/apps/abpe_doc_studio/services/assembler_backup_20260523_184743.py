"""
services/assembler.py
=====================
DocAssembler — DOCX-Generierung + PDF via LibreOffice.

Zuständig NUR für:
  - DOCX bauen mit python-docx
  - PDF konvertieren via LibreOffice
  - DocLog anlegen
  - Dateien speichern via exporter.py

NICHT zuständig für:
  - HTML-Preview (→ assembly_preview.py)
  - Variablen aus Kontext laden (→ context_loader.py)

Aufruf:
    assembler = DocAssembler()
    result = assembler.generate(
        template_identifier = 'sub_dienstvertrag',
        variables           = {'an_firma': 'ACME GmbH', ...},
        context_ref         = 'ANF-2026-0042',
        scope               = 'contract',
        engine              = 'BOTH',
        user                = request.user,
    )
"""
import io
import logging

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml    import OxmlElement

from .context_loader    import ContextLoader
from .variable_engine   import VariableEngine
from .exporter          import DocExporter
from .layout_constants  import (
    FONT_NAME, COLOR_BRAND, COLOR_TEXT, COLOR_GRAY,
    COLOR_WHITE, COLOR_ALT_ROW, COLOR_SEPARATOR, COLOR_SIG_LINE,
    TABLE_WIDTH_DXA, LABEL_VALUE_LABEL_DXA, LABEL_VALUE_VALUE_DXA,
    INV_HEADER_LEFT_DXA, INV_HEADER_RIGHT_DXA,
    TOTAL_LEFT_DXA, TOTAL_LABEL_DXA, TOTAL_VALUE_DXA,
    SIG_LEFT_DXA, SIG_MID_DXA, SIG_RIGHT_DXA,
)

log = logging.getLogger('abpe_doc_studio.assembler')


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


class DocAssembler:

    def __init__(self):
        self.context_loader  = ContextLoader()
        self.variable_engine = VariableEngine()
        self.exporter        = DocExporter()

    # ── Öffentliche API ────────────────────────────────────────────────────

    def generate(self, template_identifier: str,
                  variables: dict = None,
                  context_ref: str = '',
                  scope: str = '',
                  engine: str = 'BOTH',
                  user=None) -> dict:
        """
        Vollständige Generierung: DOCX bauen + speichern + DocLog.

        Returns:
            {
                'success':        True,
                'log_id':         '...',
                'file_path_docx': '...',
                'file_path_pdf':  '...',
                'file_size':       45123,
            }
        """
        from apps.abpe_doc_studio.models import (
            DocTemplate, DocLog, DocStatus, LogStatus
        )

        tpl = DocTemplate.objects.filter(
            identifier=template_identifier,
            status=DocStatus.ACTIVE
        ).select_related('layout', 'style_kit').prefetch_related(
            'template_blocks__block__style_kit'
        ).first()

        if not tpl:
            raise ValueError(
                f"Template '{template_identifier}' nicht gefunden oder inaktiv"
            )

        all_vars = self.context_loader.load(
            scope           = scope or tpl.scope,
            context_ref     = context_ref,
            extra_variables = variables or {},
        )

        doc = self._build_docx(tpl, all_vars)

        result       = {'success': True}
        docx_info    = self.exporter.save_docx(
            doc, template_identifier,
            scope=scope or tpl.scope, context_ref=context_ref,
        )
        result['file_path_docx'] = docx_info['file_path']
        result['file_size']      = docx_info['file_size']

        if engine in ('PDF', 'BOTH'):
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    docx_info['file_path']
                )
                result['file_path_pdf'] = pdf_info['file_path']
            except Exception as e:
                log.warning(f'PDF-Konvertierung fehlgeschlagen: {e}')
                result['file_path_pdf'] = ''
                result['pdf_error']     = str(e)

        doc_log = DocLog.objects.create(
            template         = tpl,
            template_version = tpl.active_version,
            engine_used      = engine,
            file_path_docx   = result.get('file_path_docx', ''),
            file_path_pdf    = result.get('file_path_pdf', ''),
            file_size_bytes  = result.get('file_size', 0),
            context_ref      = context_ref,
            scope            = scope or tpl.scope,
            variables_used   = {
                k: v for k, v in all_vars.items()
                if not isinstance(v, list)
            },
            status      = LogStatus.OK,
            generated_by= user,
        )
        result['log_id'] = str(doc_log.log_id)
        log.info(
            f'Generiert: {template_identifier} → '
            f'{result.get("file_path_docx","")} [log={doc_log.log_id}]'
        )
        return result

    def render_to_bytes(self, template, variables: dict = None,
                         engine: str = 'DOCX') -> bytes:
        """
        Rendert zu Bytes ohne zu speichern — für Download.
        template kann DocTemplate-Objekt oder Identifier-String sein.
        """
        from apps.abpe_doc_studio.models import DocTemplate

        if isinstance(template, str):
            template = DocTemplate.objects.filter(
                identifier=template
            ).select_related('layout', 'style_kit').prefetch_related(
                'template_blocks__block'
            ).first()
            if not template:
                raise ValueError('Template nicht gefunden')

        all_vars = variables or {}
        doc      = self._build_docx(template, all_vars)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        if engine == 'PDF':
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(buf.getvalue())
                tmp_path = tmp.name
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    tmp_path, output_dir=os.path.dirname(tmp_path)
                )
                return self.exporter.read_as_bytes(pdf_info['file_path'])
            finally:
                os.unlink(tmp_path)
                pdf_path = tmp_path.replace('.docx', '.pdf')
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)

        return buf.getvalue()

    # ── DOCX bauen ────────────────────────────────────────────────────────

    def _build_docx(self, tpl, all_vars: dict) -> DocxDocument:
        """Baut ein python-docx Dokument aus Template + Variablen."""
        doc = DocxDocument()
        self._apply_page_layout(doc, tpl.layout)

        blocks = tpl.template_blocks.order_by('slot', 'order')
        for tb in blocks:
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue
            if tb.page_break_before and doc.paragraphs:
                doc.add_page_break()
            self._render_block(doc, tb, all_vars)

        self._apply_footer(doc, tpl.layout, all_vars)
        return doc

    # ── Seitenlayout ───────────────────────────────────────────────────────

    def _apply_page_layout(self, doc: DocxDocument, layout) -> None:
        for section in doc.sections:
            section.page_width      = Cm(layout.page_width_cm)
            section.page_height     = Cm(layout.page_height_cm)
            section.left_margin     = Cm(layout.margin_left_cm)
            section.right_margin    = Cm(layout.margin_right_cm)
            section.top_margin      = Cm(layout.margin_top_cm)
            section.bottom_margin   = Cm(layout.margin_bottom_cm)
            section.header_distance = Cm(layout.header_distance_cm)

        if layout.show_page_numbers:
            self._add_page_numbers(doc, layout)

        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal'].font.size = Pt(10)

    def _add_page_numbers(self, doc: DocxDocument, layout) -> None:
        for section in doc.sections:
            hdr  = section.header
            para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.space_after = Pt(4)

            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),   'single')
            bot.set(qn('w:sz'),    '4')
            bot.set(qn('w:space'), '4')
            bot.set(qn('w:color'), COLOR_BRAND)
            pBdr.append(bot)
            pPr.append(pBdr)

            r = para.add_run('Seite ')
            r.font.name = FONT_NAME
            r.font.size = Pt(8)
            r.font.color.rgb = _rgb(COLOR_BRAND)

            for ftype, instr in [
                ('begin', ' PAGE '), ('separate', None), ('end', None)
            ]:
                fld = OxmlElement('w:fldChar')
                fld.set(qn('w:fldCharType'), ftype)
                run = para.add_run()
                run._r.append(fld)
                run.font.size = Pt(8)
                run.font.color.rgb = _rgb(COLOR_BRAND)
                if instr:
                    ins = OxmlElement('w:instrText')
                    ins.text = instr
                    run._r.append(ins)

            r2 = para.add_run(' von ')
            r2.font.name = FONT_NAME
            r2.font.size = Pt(8)
            r2.font.color.rgb = _rgb(COLOR_BRAND)

            for ftype, instr in [
                ('begin', ' NUMPAGES '), ('separate', None), ('end', None)
            ]:
                fld = OxmlElement('w:fldChar')
                fld.set(qn('w:fldCharType'), ftype)
                run = para.add_run()
                run._r.append(fld)
                run.font.size = Pt(8)
                run.font.color.rgb = _rgb(COLOR_BRAND)
                if instr:
                    ins = OxmlElement('w:instrText')
                    ins.text = instr
                    run._r.append(ins)

    def _apply_footer(self, doc: DocxDocument, layout,
                       variables: dict) -> None:
        for section in doc.sections:
            ftr  = section.footer
            para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(4)

            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top  = OxmlElement('w:top')
            top.set(qn('w:val'),   'single')
            top.set(qn('w:sz'),    '4')
            top.set(qn('w:space'), '4')
            top.set(qn('w:color'), COLOR_BRAND)
            pBdr.append(top)
            pPr.append(pBdr)

            ag  = variables
            imp = (
                f"{ag.get('ag_firma','abcona e. K.')} / "
                f"Inhaber: {ag.get('ag_name','Angelo Malaguarnera')} | "
                f"{ag.get('ag_strasse','Bornhohl 26')} | "
                f"{ag.get('ag_plz_ort','61449 Steinbach')} / Ts."
            )
            r = para.add_run(imp)
            r.font.name = FONT_NAME
            r.font.size = Pt(7)
            r.font.color.rgb = _rgb(COLOR_TEXT)

            para2 = ftr.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            jur = (
                f"{ag.get('ag_amtsgericht','Amtsgericht Bad Homburg v. d. H.')} | "
                f"{ag.get('ag_hra','HRA Nr. 3662')} | "
                f"USt-IdNr.: {ag.get('ag_ustid','DE813519516')} / "
                f"{ag.get('ag_steuernr','03-844-32179')}"
            )
            r2 = para2.add_run(jur)
            r2.font.name = FONT_NAME
            r2.font.size = Pt(7)
            r2.font.color.rgb = _rgb(COLOR_TEXT)

    # ── Block-Dispatch ────────────────────────────────────────────────────

    def _render_block(self, doc: DocxDocument, tb, variables: dict) -> None:
        block            = tb.block
        content          = tb.content_override or block.content or ''
        rendered_content = self.variable_engine.render_text(content, variables)
        bt               = block.block_type

        dispatch = {
            'LOGO':        lambda: self._render_logo(doc, variables),
            'DOC_TITLE':   lambda: self._render_doc_title(doc, rendered_content),
            'SECTION_HEAD':lambda: self._render_section_heading(doc, rendered_content),
            'CLAUSE':      lambda: self._render_clause(doc, rendered_content),
            'PARAGRAPH':   lambda: self._render_paragraph(doc, rendered_content),
            'PARTY_BLOCK': lambda: self._render_party_block(doc, rendered_content),
            'LABEL_VALUE': lambda: self._render_label_value(doc, rendered_content),
            'TIME_TABLE':  lambda: self._render_time_table(doc, block, variables),
            'AP_TABLE':    lambda: self._render_ap_table(doc, block, variables),
            'TOTAL_BLOCK': lambda: self._render_total_block(doc, variables),
            'SIGNATURE':   lambda: self._render_signature(doc, variables),
            'SEPARATOR':   lambda: self._render_separator(doc),
            'PAGE_BREAK':  lambda: doc.add_page_break(),
            'INV_HEADER':  lambda: self._render_inv_header(doc, variables),
            'INV_META':    lambda: self._render_inv_meta(doc, variables),
            'INV_SUBJECT': lambda: self._render_inv_subject(doc, rendered_content, variables),
            'CLOSING':     lambda: self._render_closing(doc, variables),
        }

        fn = dispatch.get(bt)
        if fn:
            fn()
        elif rendered_content:
            p = doc.add_paragraph(rendered_content)
            p.paragraph_format.space_after = Pt(6)

    # ── Block-Typen DOCX ──────────────────────────────────────────────────

    def _render_logo(self, doc, variables: dict) -> None:
        import os
        from django.conf import settings
        logo_paths = [
            os.path.join(settings.BASE_DIR, 'data', 'cv', 'adds', 'logo_abcona.png'),
            os.path.join(settings.BASE_DIR, 'apps', 'abpe_ui', 'static',
                         'abpe_ui', 'img', 'logo_abcona.png'),
        ]
        logo_path = next((p for p in logo_paths if os.path.exists(p)), None)

        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(20)
        p.paragraph_format.space_before = Pt(0)

        if logo_path:
            run = p.add_run()
            run.add_picture(logo_path, height=Cm(2.0))
        else:
            r = p.add_run('abcona e. K.')
            r.font.name = FONT_NAME
            r.font.size = Pt(16)
            r.bold      = True
            r.font.color.rgb = _rgb(COLOR_BRAND)

    def _render_doc_title(self, doc, content: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(16)
        r = p.add_run(content)
        r.font.name = FONT_NAME
        r.font.size = Pt(16)
        r.bold      = True
        r.font.color.rgb = _rgb(COLOR_TEXT)

    def _render_section_heading(self, doc, content: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(content)
        r.font.name = FONT_NAME
        r.font.size = Pt(12)
        r.bold      = True
        r.font.color.rgb = _rgb(COLOR_BRAND)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), COLOR_BRAND)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _render_clause(self, doc, content: str) -> None:
        if not content:
            return
        for line in content.split('\n'):
            if not line.strip():
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(line.strip())
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.font.color.rgb = _rgb(COLOR_TEXT)

    def _render_paragraph(self, doc, content: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(content)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb(COLOR_TEXT)

    def _render_party_block(self, doc, content: str) -> None:
        lines = content.split('\n') if content else []
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1.5)
            r = p.add_run(line)
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.bold      = True
            r.italic    = i < 2
            r.font.color.rgb = _rgb(COLOR_TEXT)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _render_label_value(self, doc, content: str) -> None:
        if not content:
            return
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc, vc = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
        for cell in (lc, vc):
            self._set_cell_borders_none(cell)
        self._set_cell_width(lc, LABEL_VALUE_LABEL_DXA)
        self._set_cell_width(vc, LABEL_VALUE_VALUE_DXA)

        if '|' in content:
            label, value = content.split('|', 1)
        else:
            label, value = content, ''

        rl = lc.paragraphs[0].add_run(label.strip())
        rl.font.name = FONT_NAME
        rl.font.size = Pt(10)
        rl.bold      = True
        rl.font.color.rgb = _rgb(COLOR_BRAND)

        rv = vc.paragraphs[0].add_run(value.strip())
        rv.font.name = FONT_NAME
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(COLOR_TEXT)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _render_time_table(self, doc, block, variables: dict) -> None:
        columns = block.columns or []
        rows    = self.variable_engine.expand_table_rows(
            columns, 'positionen', variables
        )
        self._build_invoice_table(doc, columns, rows)

    def _render_ap_table(self, doc, block, variables: dict) -> None:
        columns = block.columns or []
        rows    = self.variable_engine.expand_table_rows(
            columns, 'arbeitspakete', variables
        )
        self._build_invoice_table(doc, columns, rows)

    def _build_invoice_table(self, doc, columns: list,
                               rows: list) -> None:
        if not columns:
            return
        n_cols = len(columns)
        tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'),    str(TABLE_WIDTH_DXA))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        total_pct = sum(c.get('width_pct', 20) for c in columns)
        hrow      = tbl.rows[0]
        for ci, col in enumerate(columns):
            cell = hrow.cells[ci]
            w    = int(TABLE_WIDTH_DXA * col.get('width_pct', 20) / total_pct)
            self._set_cell_width(cell, w)
            self._set_cell_bg(cell, COLOR_BRAND)
            self._set_cell_borders_none(cell)
            r = cell.paragraphs[0].add_run(col.get('label', ''))
            r.font.name = FONT_NAME
            r.font.size = Pt(9)
            r.bold      = True
            r.font.color.rgb = _rgb(COLOR_WHITE)

        for ri, row_data in enumerate(rows):
            drow = tbl.rows[ri + 1]
            bg   = COLOR_ALT_ROW if ri % 2 == 0 else COLOR_WHITE
            for ci, col in enumerate(columns):
                cell  = drow.cells[ci]
                value = row_data[ci] if ci < len(row_data) else ''
                w     = int(TABLE_WIDTH_DXA * col.get('width_pct', 20) / total_pct)
                self._set_cell_width(cell, w)
                self._set_cell_bg(cell, bg)
                self._set_cell_borders_none(cell)
                p     = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                               if col.get('align') == 'right'
                               else WD_ALIGN_PARAGRAPH.LEFT)
                r = p.add_run(value)
                r.font.name = FONT_NAME
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(COLOR_TEXT)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _render_total_block(self, doc, variables: dict) -> None:
        v    = variables
        rows = [
            ('Summe',
             v.get('summe_netto', '')),
            (f"zzgl. Umsatzsteuer {v.get('mwst_satz','19')} %",
             v.get('mwst_euro', '')),
            ('Gesamtbetrag',
             v.get('gesamtbetrag', '')),
        ]
        tbl = doc.add_table(rows=len(rows), cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'),    str(TABLE_WIDTH_DXA))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        for ri, (label, value) in enumerate(rows):
            is_total = ri == len(rows) - 1
            bg       = COLOR_BRAND if is_total else COLOR_WHITE
            clr      = COLOR_WHITE if is_total else COLOR_TEXT
            for ci, (w, text) in enumerate([
                (TOTAL_LEFT_DXA,  ''),
                (TOTAL_LABEL_DXA, label),
                (TOTAL_VALUE_DXA, value),
            ]):
                cell = tbl.rows[ri].cells[ci]
                self._set_cell_width(cell, w)
                self._set_cell_bg(cell, bg)
                self._set_cell_borders_none(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(text)
                r.font.name = FONT_NAME
                r.font.size = Pt(10 if is_total else 9)
                r.bold      = is_total
                r.font.color.rgb = _rgb(clr)

        doc.add_paragraph().paragraph_format.space_after = Pt(16)

    def _render_signature(self, doc, variables: dict) -> None:
        v      = variables
        ag_ort = v.get('ag_ort', 'Steinbach')
        an_ort = v.get('ort_an', '')
        datum  = v.get('datum_heute', '')

        tbl = doc.add_table(rows=3, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for ci, (w, text) in enumerate([
            (SIG_LEFT_DXA,  f'{ag_ort}, den {datum}'),
            (SIG_MID_DXA,   ''),
            (SIG_RIGHT_DXA, f'{an_ort}, den {datum}'),
        ]):
            cell = tbl.rows[0].cells[ci]
            self._set_cell_width(cell, w)
            self._set_cell_borders_none(cell)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.bold      = True
            r.font.color.rgb = _rgb(COLOR_TEXT)
            if ci != 1:
                pPr  = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bot  = OxmlElement('w:bottom')
                bot.set(qn('w:val'),   'single')
                bot.set(qn('w:sz'),    '4')
                bot.set(qn('w:space'), '4')
                bot.set(qn('w:color'), COLOR_SIG_LINE)
                pBdr.append(bot)
                pPr.append(pBdr)

        for ci in range(3):
            cell = tbl.rows[1].cells[ci]
            self._set_cell_borders_none(cell)
            cell.paragraphs[0].paragraph_format.space_before = Pt(50)

        for ci, (w, label, hint) in enumerate([
            (SIG_LEFT_DXA,  'Auftraggeber AG',  'Stempel/Unterschrift'),
            (SIG_MID_DXA,   '',                  ''),
            (SIG_RIGHT_DXA, 'Auftragnehmer AN', 'Unterschrift'),
        ]):
            cell = tbl.rows[2].cells[ci]
            self._set_cell_width(cell, w)
            self._set_cell_borders_none(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if label:
                r = p.add_run(label)
                r.font.name = FONT_NAME
                r.font.size = Pt(10)
                r.bold      = True
                r.font.color.rgb = _rgb(COLOR_TEXT)
            if hint:
                p2 = cell.add_paragraph()
                r2 = p2.add_run(hint)
                r2.font.name = FONT_NAME
                r2.font.size = Pt(9)
                r2.font.color.rgb = _rgb(COLOR_GRAY)

    def _render_separator(self, doc) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), COLOR_SEPARATOR)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _render_inv_header(self, doc, variables: dict) -> None:
        v    = variables
        tbl  = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc, rc = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
        self._set_cell_width(lc, INV_HEADER_LEFT_DXA)
        self._set_cell_width(rc, INV_HEADER_RIGHT_DXA)
        self._set_cell_borders_none(lc)
        self._set_cell_borders_none(rc)

        firma = v.get('empfaenger_firma', '')
        addr  = v.get('empfaenger_adresse', '').strip()
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        rl = lp.add_run(firma)
        rl.font.name = FONT_NAME
        rl.font.size = Pt(10)
        rl.bold      = True
        rl.font.color.rgb = _rgb(COLOR_TEXT)
        for line in addr.split('\n'):
            if line.strip():
                p2 = lc.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                r2 = p2.add_run(line.strip())
                r2.font.name = FONT_NAME
                r2.font.size = Pt(10)
                r2.font.color.rgb = _rgb(COLOR_TEXT)

        contact = [
            (v.get('ag_firma', 'abcona e.K.'), True,  True,  10),
            ('active business consulting agency', False, True, 9),
            (f"Tel: {v.get('ag_tel','')}", False, False, 9),
            (f"Fax: {v.get('ag_fax','')}", False, False, 9),
            (f"Mail: {v.get('ag_email','')}", False, False, 9),
            (f"http://{v.get('ag_web','')}", False, False, 9),
        ]
        first = True
        for text, bold, italic, size in contact:
            p = rc.paragraphs[0] if first else rc.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.name = FONT_NAME
            r.font.size = Pt(size)
            r.bold      = bold
            r.italic    = italic
            r.font.color.rgb = _rgb(COLOR_BRAND if bold else COLOR_TEXT)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def _render_inv_meta(self, doc, variables: dict) -> None:
        v    = variables
        p    = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(f"Rg.-Nr.: {v.get('rg_nummer','')}")
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
        r.bold      = True
        r.font.color.rgb = _rgb(COLOR_BRAND)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), COLOR_BRAND)
        pBdr.append(bot)
        pPr.append(pBdr)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(v.get('rg_datum', ''))
        r2.font.name = FONT_NAME
        r2.font.size = Pt(10)
        r2.bold      = True
        r2.font.color.rgb = _rgb(COLOR_BRAND)

    def _render_inv_subject(self, doc, content: str,
                              variables: dict) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        betreff = content or variables.get('betreff', '')
        r = p.add_run(betreff)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.bold      = True
        r.font.color.rgb = _rgb(COLOR_BRAND)

    def _render_closing(self, doc, variables: dict) -> None:
        v = variables
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run('Mit freundlichen Grüssen')
        r.font.name = FONT_NAME
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(COLOR_TEXT)

        ziel = v.get('zahlungsziel_text', '30 Tage netto')
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(40)
        r2 = p2.add_run(f'Zahlungsziel: {ziel}')
        r2.font.name = FONT_NAME
        r2.font.size = Pt(9)
        r2.bold      = True
        r2.font.color.rgb = _rgb(COLOR_TEXT)

        for text, bold, italic in [
            ('active business consulting agency', False, True),
            ('abcona e.K.',                       True,  False),
            ('Angelo Malaguarnera · Inhaber',      False, False),
        ]:
            px = doc.add_paragraph()
            rx = px.add_run(text)
            rx.font.name = FONT_NAME
            rx.font.size = Pt(9)
            rx.bold      = bold
            rx.italic    = italic
            rx.font.color.rgb = _rgb(COLOR_TEXT)

    # ── Zell-Helfer ───────────────────────────────────────────────────────

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.lstrip('#'))
        tcPr.append(shd)

    def _set_cell_width(self, cell, width_dxa: int) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_cell_borders_none(self, cell) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        bdrs = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement('w:' + side)
            el.set(qn('w:val'), 'nil')
            bdrs.append(el)
        tcPr.append(bdrs)
