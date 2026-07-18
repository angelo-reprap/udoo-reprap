"""
services/assembler.py
=====================
DocAssembler — Herzstück des Doc Studios.

Orchestriert alle 4 Ebenen:
  1. PageLayout laden → Seitenformat
  2. StyleKit laden   → Formatierungen
  3. Blöcke in Reihenfolge rendern
  4. Variablen füllen

Dann: .docx speichern + optional PDF konvertieren + DocLog anlegen.

Aufruf:
    assembler = DocAssembler()
    result = assembler.generate(
        template_identifier = 'sub_dienstvertrag',
        variables           = {'an_firma': 'ACME GmbH', ...},
        context_ref         = 'ANF-2026-0042',
        scope               = 'contract',
        engine              = 'BOTH',
        user                = request.user,
    )
"""
import io
import logging
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml    import OxmlElement

from .context_loader  import ContextLoader
from .variable_engine import VariableEngine
from .exporter        import DocExporter

log = logging.getLogger('abpe_doc_studio.assembler')


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


class DocAssembler:

    def __init__(self):
        self.context_loader   = ContextLoader()
        self.variable_engine  = VariableEngine()
        self.exporter         = DocExporter()

    # ── Öffentliche API ────────────────────────────────────────────────────

    def generate(self, template_identifier: str,
                  variables: dict = None,
                  context_ref: str = '',
                  scope: str = '',
                  engine: str = 'BOTH',
                  user=None) -> dict:
        """
        Vollständige Generierung: rendern + speichern + DocLog.

        Returns:
            {
                'success':       True,
                'log_id':        '...',
                'file_path_docx': '...',
                'file_path_pdf':  '...',
                'file_size':      45123,
            }
        """
        from apps.abpe_doc_studio.models import (
            DocTemplate, DocLog, DocStatus, LogStatus
        )

        tpl = DocTemplate.objects.filter(
            identifier=template_identifier,
            status=DocStatus.ACTIVE
        ).select_related('layout', 'style_kit').prefetch_related(
            'template_blocks__block__style_kit'
        ).first()

        if not tpl:
            raise ValueError(
                f"Template '{template_identifier}' nicht gefunden oder inaktiv"
            )

        # Variablen aus Kontext laden + extra_variables ergänzen
        all_vars = self.context_loader.load(
            scope       = scope or tpl.scope,
            context_ref = context_ref,
            extra_variables = variables or {},
        )

        # Dokument bauen
        doc = DocxDocument()
        self._apply_page_layout(doc, tpl.layout)

        # Blöcke in Reihenfolge rendern
        blocks = tpl.template_blocks.order_by('slot', 'order')
        for tb in blocks:
            # Bedingung prüfen
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue

            if tb.page_break_before and doc.paragraphs:
                doc.add_page_break()

            self._render_block(doc, tb, all_vars)

        # Footer auf jeder Seite (via Header-Objekt)
        self._apply_footer(doc, tpl.layout, all_vars)

        # Speichern
        result = {'success': True}
        docx_info = self.exporter.save_docx(
            doc,
            template_identifier,
            scope       = scope or tpl.scope,
            context_ref = context_ref,
        )
        result['file_path_docx'] = docx_info['file_path']
        result['file_size']      = docx_info['file_size']

        # PDF konvertieren
        if engine in ('PDF', 'BOTH'):
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    docx_info['file_path']
                )
                result['file_path_pdf'] = pdf_info['file_path']
            except Exception as e:
                log.warning(f'PDF-Konvertierung fehlgeschlagen: {e}')
                result['file_path_pdf'] = ''
                result['pdf_error']     = str(e)

        # DocLog anlegen
        doc_log = DocLog.objects.create(
            template         = tpl,
            template_version = tpl.active_version,
            engine_used      = engine,
            file_path_docx   = result.get('file_path_docx', ''),
            file_path_pdf    = result.get('file_path_pdf', ''),
            file_size_bytes  = result.get('file_size', 0),
            context_ref      = context_ref,
            scope            = scope or tpl.scope,
            variables_used   = {
                k: v for k, v in all_vars.items()
                if not isinstance(v, list)
            },
            status           = LogStatus.OK,
            generated_by     = user,
        )
        result['log_id'] = str(doc_log.log_id)

        log.info(
            f'Generiert: {template_identifier} → '
            f'{result.get("file_path_docx", "")} '
            f'[log={doc_log.log_id}]'
        )
        return result

    def render_to_bytes(self, template, variables: dict = None,
                         engine: str = 'DOCX') -> bytes:
        """
        Rendert zu Bytes ohne zu speichern — für Preview und Download.
        template kann DocTemplate-Objekt oder Identifier-String sein.
        """
        from apps.abpe_doc_studio.models import DocTemplate, DocStatus

        if isinstance(template, str):
            template = DocTemplate.objects.filter(
                identifier=template
            ).select_related('layout', 'style_kit').prefetch_related(
                'template_blocks__block'
            ).first()
            if not template:
                raise ValueError(f"Template nicht gefunden")

        all_vars = variables or {}

        doc = DocxDocument()
        self._apply_page_layout(doc, template.layout)

        blocks = template.template_blocks.order_by('slot', 'order')
        for tb in blocks:
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue
            if tb.page_break_before and doc.paragraphs:
                doc.add_page_break()
            self._render_block(doc, tb, all_vars)

        self._apply_footer(doc, template.layout, all_vars)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        if engine == 'PDF':
            # Temp-Datei für LibreOffice
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(buf.getvalue())
                tmp_path = tmp.name
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    tmp_path, output_dir=os.path.dirname(tmp_path)
                )
                return self.exporter.read_as_bytes(pdf_info['file_path'])
            finally:
                os.unlink(tmp_path)
                pdf_path = tmp_path.replace('.docx', '.pdf')
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)

        return buf.getvalue()

    def preview_html(self, template, variables: dict = None) -> str:
        """
        Rendert eine HTML-Vorschau mit echter DIN-A4 Seitenaufteilung.
        Blöcke werden serverseitig auf Seiten verteilt anhand ihrer
        geschätzten Höhe in mm.
        """
        from apps.abpe_doc_studio.models import DocTemplate
        if isinstance(template, str):
            template = DocTemplate.objects.filter(
                identifier=template
            ).select_related('layout', 'style_kit').prefetch_related(
                'template_blocks__block'
            ).first()

        all_vars = variables or {}
        lay = template.layout

        # ── Layout-Werte aus DB ────────────────────────────────────────
        import json as _j
        if lay:
            page_h_mm   = float(lay.page_height_cm)  * 10   # 297mm
            margin_t_mm = float(lay.margin_top_cm)   * 10   # 42mm
            margin_b_mm = float(lay.margin_bottom_cm)* 10   # 52mm
            lay_data = {
                'w': float(lay.page_width_cm),
                'h': float(lay.page_height_cm),
                't': float(lay.margin_top_cm),
                'b': float(lay.margin_bottom_cm),
                'l': float(lay.margin_left_cm),
                'r': float(lay.margin_right_cm),
            }
        else:
            page_h_mm   = 297.0
            margin_t_mm = 42.0
            margin_b_mm = 52.0
            lay_data = {'w':21.0,'h':29.7,'t':4.2,'b':5.2,'l':3.0,'r':3.0}

        lay_json = _j.dumps(lay_data)

        # Nutzbare Inhaltshöhe pro Seite in mm
        # 20mm Reserve für Fußzeile-Abstand
        max_h_mm = page_h_mm - margin_t_mm - margin_b_mm - 20.0

        # ── Fußzeile HTML (auf jeder Seite) ───────────────────────────
        FOOTER = (
            '<div class="ds-page-footer">' +
            '<div style="display:grid;grid-template-columns:1fr 1fr 1fr 1fr;' +
            'font-size:6px;color:#555;border-top:1px solid #163258;padding-top:4px;">' +
            '<div><b style="color:#1a1a1a;">Deutsche Bank</b><br>' +
            'IBAN: DE39 5007 0024 0447 3625 00<br>BIC: DEUTDEDBFRA</div>' +
            '<div><b style="color:#1a1a1a;">Volksbank Mittelhessen</b><br>' +
            'IBAN: DE60 5139 0000 0038 3591 00<br>BIC: GENODE51OBU</div>' +
            '<div><b style="color:#1a1a1a;">Amtsgericht</b><br>' +
            'Bad Homburg v.d.H.<br>HRA Nr. 3662</div>' +
            '<div><b style="color:#1a1a1a;">USt-IdNr.</b><br>' +
            'DE 813519516<br>03-844-32179</div>' +
            '</div></div>'
        )

        # ── Blöcke auf Seiten aufteilen ────────────────────────────────
        blocks   = template.template_blocks.order_by('slot', 'order')
        parts    = [f'<div class="ds-doc-preview-wrap" data-layout="{lay_json}">']
        cur_mm   = 0.0
        page_num = 1

        parts.append(f'<div class="ds-preview-page" data-page="{page_num}">')

        for tb in blocks:
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue

            block_mm = self._estimate_height_mm(tb)

            # Seitenumbruch: wenn Block nicht mehr passt UND
            # bereits genug Inhalt auf der Seite (mind. 20mm)
            if cur_mm + block_mm > max_h_mm and cur_mm > 20.0:
                parts.append(FOOTER)
                parts.append('</div>')   # Seite schließen
                page_num += 1
                cur_mm    = 0.0
                parts.append(
                    f'<div class="ds-preview-page" data-page="{page_num}">'
                )

            parts.append(self._render_block_html(tb, all_vars, tb.pk))
            cur_mm += block_mm

        # Letzte Seite schließen
        parts.append(FOOTER)
        parts.append('</div>')   # letzte Seite
        parts.append('</div>')   # wrap
        return '\n'.join(parts)

    def _estimate_height_mm(self, tb) -> float:
        """
        Schätzt die Höhe eines Blocks in mm.
        Grundlage: DIN A4, Arial 10pt, ~62 Zeichen pro Zeile, 5mm pro Zeile.
        Stimmt gut mit dem echten DOCX-Output überein.
        """
        bt      = tb.block.block_type
        content = tb.content_override or tb.block.content or ''

        # ── Fixe Höhen ─────────────────────────────────────────────────
        FIXED = {
            'LOGO':         40.0,   # Logo + Kontaktblock + Trennlinie
            'DOC_TITLE':    14.0,   # Titel 16pt + Abstand
            'SECTION_HEAD': 12.0,   # §-Heading 12pt + Unterlinie + Abstand
            'SIGNATURE':    35.0,   # Signatur-Tabelle 3 Zeilen
            'PAGE_BREAK':    0.0,   # Erzwingt neue Seite → sofort umbrechen
            'SEPARATOR':     4.0,   # Trennlinie
            'FOOTER':       18.0,   # Bankdaten-Fußzeile
            'INV_META':     16.0,   # Rg-Nr. + Datum
            'INV_SUBJECT':  10.0,   # Betreff-Zeile
            'CLOSING':      40.0,   # Grußformel + Unterschrift
            'TOTAL_BLOCK':  20.0,   # Summen-Tabelle
        }
        if bt in FIXED:
            # PAGE_BREAK: sofort Seitenumbruch erzwingen
            if bt == 'PAGE_BREAK':
                return 9999.0
            return FIXED[bt]

        # ── PARTY_BLOCK: Zeilen zählen ─────────────────────────────────
        if bt == 'PARTY_BLOCK':
            lines = len([l for l in content.split('\n') if l.strip()])
            return max(16.0, lines * 5.5 + 4.0)

        # ── CLAUSE: Überschrift + Fließtext ────────────────────────────
        if bt == 'CLAUSE':
            lines_total = 0.0
            for line in content.split('\n'):
                chars = len(line.strip())
                if chars == 0:
                    lines_total += 0.4   # Leerzeile = kleiner Abstand
                else:
                    # 62 Zeichen pro Zeile bei 10pt Arial, Blocksatz
                    lines_total += max(1.0, chars / 62.0)
            # 5mm pro Zeile + 10mm Overhead (Heading + Abstand)
            return max(14.0, lines_total * 5.0 + 10.0)

        # ── PARAGRAPH: Fließtext ───────────────────────────────────────
        if bt == 'PARAGRAPH':
            chars = len(content)
            lines = max(1.0, chars / 65.0)
            return lines * 5.0 + 5.0

        # ── LABEL_VALUE: 1 Zeile Tabelle ──────────────────────────────
        if bt == 'LABEL_VALUE':
            return 8.0

        # ── Tabellen ───────────────────────────────────────────────────
        if bt in ('TIME_TABLE', 'AP_TABLE'):
            return 50.0   # Schätzwert inkl. Header

        # ── INV_HEADER: Rechnungskopf 2-spaltig ───────────────────────
        if bt == 'INV_HEADER':
            return 35.0

        # ── Fallback: proportional zur Zeichenzahl ─────────────────────
        chars = len(content)
        return max(8.0, chars / 65.0 * 5.0 + 6.0)


    # ── HTML-Vorschau ──────────────────────────────────────────────────────

    def _render_block_html(self, tb, variables: dict, tb_pk: int = None) -> str:
        """HTML-Vorschau eines Blocks — mit data-tb-id für den Editor."""
        block   = tb.block
        content = self.variable_engine.render_text(
            tb.content_override or block.content or '', variables
        )
        bt   = block.block_type
        pid  = tb_pk or 0
        ord_ = getattr(tb, 'order', 0)

        def wrap(inner):
            if not pid:
                return inner or ''
            attrs = f'data-tb-id="{pid}" data-block-type="{bt}" data-order="{ord_}"'
            return f'<div class="ds-preview-block" {attrs}>{inner or ""}</div>'

        if bt == 'DOC_TITLE':
            return wrap(f'<h1 style="font-size:20px;color:#163258;margin:0 0 8px">{content}</h1>')
        elif bt == 'SECTION_HEAD':
            return wrap(f'<h2 style="font-size:13px;color:#163258;border-bottom:2px solid #163258;padding-bottom:3px;margin:8px 0 3px">{content}</h2>')
        elif bt == 'CLAUSE':
            lines = content.split('\n')
            hdr  = lines[0] if lines else ''
            body = ''.join(f'<p style="text-align:justify;margin:3px 0;font-size:10px">{l}</p>' for l in lines[1:] if l.strip())
            inner = (f'<div style="font-size:12px;font-weight:700;color:#163258;border-bottom:1px solid #163258;padding-bottom:2px;margin-bottom:3px">{hdr}</div>' + body)
            return wrap(inner)
        elif bt == 'PARAGRAPH':
            paras = ''.join(f'<p style="text-align:justify;margin:3px 0;font-size:10px">{l}</p>' for l in content.split('\n') if l.strip())
            return wrap(paras)
        elif bt == 'PARTY_BLOCK':
            html = ''.join(f'<div style="font-size:9px;font-weight:600;color:#1a1a1a;margin-left:16px;line-height:1.8">{l}</div>' for l in content.split('\n') if l.strip())
            return wrap(html)
        elif bt == 'LABEL_VALUE':
            parts = content.split('|', 1)
            if len(parts) == 2:
                inner = (f'<div style="display:flex;gap:8px;margin:3px 0;font-size:9px"><span style="font-weight:700;color:#163258;min-width:140px">{parts[0]}</span><span>{parts[1]}</span></div>')
            else:
                inner = f'<p style="margin:3px 0;font-size:9px">{content}</p>'
            return wrap(inner)
        elif bt == 'PAGE_BREAK':
            return wrap('<hr style="border:2px dashed #ccc;margin:16px 0">')
        elif bt in ('TIME_TABLE', 'AP_TABLE'):
            return wrap('<div style="background:#f5f5f5;padding:8px;color:#888;font-size:9px;border-radius:3px">[Tabelle]</div>')
        elif bt == 'SIGNATURE':
            inner = ('<div style="display:flex;justify-content:space-between;margin-top:24px;border-top:1px solid #ccc;padding-top:6px"><div style="font-size:8px"><b>Auftraggeber AG</b><br><span style="color:#aaa">Stempel/Unterschrift</span></div><div style="font-size:8px"><b>Auftragnehmer AN</b><br><span style="color:#aaa">Unterschrift</span></div></div>')
            return wrap(inner)
        elif bt == 'LOGO':
            logo_html = (
                '<div style="display:flex;justify-content:space-between;'
                'align-items:flex-start;margin-bottom:16px;border-bottom:'
                '1px solid #163258;padding-bottom:8px;">'
                '<div style="font-size:7px;color:#163258;">'
                'abcona e.K. &middot; Bornhohl 26 &nbsp; 61449 Steinbach/Ts.</div>'
                '<div style="text-align:right;">'
                '<img src="/static/abpe_ui/img/logo_abcona.png" '
                'style="height:40px;width:auto;display:block;margin-bottom:4px;"><br>'
                '<div style="font-size:7px;color:#163258;line-height:1.8;">'
                '<i>active business consulting agency</i><br>'
                '<b>abcona e.K.</b><br>'
                'Tel: +49 (0)6171 8867 00<br>'
                'Fax: +49 (0)6171 8867 09<br>'
                'Mail: office@abcona.de<br>'
                'http://www.abcona.de</div></div></div>'
            )
            return wrap(logo_html)
        elif bt == 'FOOTER':
            footer_html = (
                '<div style="display:grid;grid-template-columns:1fr 1fr 1fr 1fr;'
                'gap:0;font-size:6.5px;color:#555;border-top:1px solid #163258;'
                'padding-top:5px;margin-top:8px;">'
                '<div><b style="color:#1a1a1a;">Deutsche Bank</b><br>'
                'IBAN: DE39 5007 0024 0447 3625 00<br>BIC: DEUTDEDBFRA</div>'
                '<div><b style="color:#1a1a1a;">Volksbank Mittelhessen</b><br>'
                'IBAN: DE60 5139 0000 0038 3591 00<br>BIC: GENODE51OBU</div>'
                '<div><b style="color:#1a1a1a;">Amtsgericht</b><br>'
                'Bad Homburg v.d.H.<br>HRA Nr. 3662</div>'
                '<div><b style="color:#1a1a1a;">USt-IdNr.</b><br>'
                'DE 813519516<br>03-844-32179</div>'
                '</div>'
            )
            return wrap(footer_html)
        elif content:
            return wrap(f'<p style="margin:3px 0;font-size:10px">{content}</p>')
        return wrap('')


    # ── Zell-Helfer ────────────────────────────────────────────────────────

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        from docx.oxml import OxmlElement
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.lstrip('#'))
        tcPr.append(shd)

    def _set_cell_width(self, cell, width_dxa: int) -> None:
        from docx.oxml import OxmlElement
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_cell_borders_none(self, cell) -> None:
        from docx.oxml import OxmlElement
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        bdrs = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement('w:' + side)
            el.set(qn('w:val'), 'nil')
            bdrs.append(el)
        tcPr.append(bdrs)
