"""
bin/docx_extractor.py
=====================
Liest ein DOCX und extrahiert:
  - Seitenmaße → layout.json
  - Styles (Font, Size, Color, Bold, Alignment) → styles.json
  - Textboxen → blocks.json (block_type=TEXTBOX)
  - Tabellen → blocks.json (block_type=TABLE)
  - Paragraphen → blocks.json (block_type=PARAGRAPH)
  - Bilder → image_refs in layout.json

Aufruf:
    python apps/abpe_doc_studio/bin/docx_extractor.py \
        --docx data/cv/adds/templates/docx/briefpapier.docx \
        --out  apps/abpe_doc_studio/generator/templates/bp_1/ \
        --name bp1

    # Nur analysieren ohne Schreiben:
    python apps/abpe_doc_studio/bin/docx_extractor.py \
        --docx data/cv/adds/templates/docx/briefpapier.docx \
        --analyze
"""
import os
import sys
import json
import argparse
from collections import OrderedDict

# Django setup
def setup_django():
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(
        os.path.dirname(os.path.abspath(__file__))
    )))
    sys.path.insert(0, BASE_DIR)
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'abpe_backend.settings')
    import django
    django.setup()

if __name__ == '__main__':
    setup_django()

from docx import Document
from docx.oxml.ns import qn


# ── Hilfsfunktionen ───────────────────────────────────────────────────────

def emu_to_cm(emu: int) -> float:
    return round(emu / 914400 * 2.54, 2)

def dxa_to_cm(dxa: int) -> float:
    return round(dxa / 567, 2)

def twips_to_pt(twips: int) -> float:
    return round(twips / 20, 1)

def get_text(elem) -> str:
    # Unterstuetzt sowohl lxml-Element als auch python-docx Paragraph
    if hasattr(elem, '_p'):
        elem = elem._p
    return ''.join(t.text or '' for t in elem.findall('.//' + qn('w:t')))

def get_run_props(rPr) -> dict:
    """Extrahiert Run-Properties."""
    if rPr is None:
        return {}
    props = {}
    sz = rPr.find(qn('w:sz'))
    if sz is not None:
        props['font_size_pt'] = int(sz.get(qn('w:val'), '20')) / 2
    b = rPr.find(qn('w:b'))
    if b is not None:
        props['bold'] = True
    i = rPr.find(qn('w:i'))
    if i is not None:
        props['italic'] = True
    u = rPr.find(qn('w:u'))
    if u is not None and u.get(qn('w:val'), 'none') != 'none':
        props['underline'] = True
    color = rPr.find(qn('w:color'))
    if color is not None:
        val = color.get(qn('w:val'), '')
        if val and val != 'auto':
            props['color_hex'] = val
    font = rPr.find(qn('w:rFonts'))
    if font is not None:
        name = font.get(qn('w:ascii')) or font.get(qn('w:hAnsi'))
        if name:
            props['font_family'] = name
    return props

def get_para_props(pPr) -> dict:
    """Extrahiert Paragraph-Properties."""
    if pPr is None:
        return {}
    props = {}
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        props['alignment'] = jc.get(qn('w:val'), 'left')
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        before = spacing.get(qn('w:before'))
        after  = spacing.get(qn('w:after'))
        if before:
            props['space_before_pt'] = twips_to_pt(int(before))
        if after:
            props['space_after_pt']  = twips_to_pt(int(after))
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        left = ind.get(qn('w:left'))
        if left:
            props['indent_left_cm'] = dxa_to_cm(int(left))
    # Rahmen
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is not None:
        bot = pBdr.find(qn('w:bottom'))
        if bot is not None:
            props['border_bottom'] = True
            props['border_bottom_color'] = bot.get(qn('w:color'), '163258')
            sz = bot.get(qn('w:sz'), '4')
            props['border_bottom_pt'] = round(int(sz) / 8, 1)
    return props


# ── Extraktor ─────────────────────────────────────────────────────────────

class DocxExtractor:

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc       = Document(docx_path)
        self.body      = self.doc.element.body
        self.name      = os.path.splitext(os.path.basename(docx_path))[0]

        self.layout  = {}
        self.styles  = []
        self.blocks  = []
        self._style_map = {}  # key → style_key

    def extract(self) -> dict:
        self._extract_layout()
        self._extract_header_footer()
        self._extract_textboxes()
        self._extract_tables()
        self._extract_paragraphs()
        return {
            'layout': self.layout,
            'styles': self.styles,
            'blocks': self.blocks,
        }

    def _extract_layout(self):
        """Seitenmaße aus Section."""
        for section in self.doc.sections:
            self.layout = {
                'identifier':          f'layout_{self.name}',
                'name':                f'Layout {self.name}',
                'page_width_cm':       round(section.page_width.cm, 1),
                'page_height_cm':      round(section.page_height.cm, 1),
                'margin_left_cm':      round(section.left_margin.cm, 1),
                'margin_right_cm':     round(section.right_margin.cm, 1),
                'margin_top_cm':       round(section.top_margin.cm, 1),
                'margin_bottom_cm':    round(section.bottom_margin.cm, 1),
                'header_distance_cm':  round(section.header_distance.cm, 1),
                'footer_distance_cm':  round(section.footer_distance.cm, 1),
                'orientation':         'landscape' if section.page_width > section.page_height else 'portrait',
                'normal_font_size_pt': 10.0,
                'layout_refs':         {},
                'image_refs':          {},
            }
            break

    def _extract_header_footer(self):
        """Header + Footer Paragraphen."""
        for section in self.doc.sections:
            # Header
            hdr_lines = []
            hdr_styles = []
            for p in section.header.paragraphs:
                text = get_text(p._p)
                if text.strip():
                    hdr_lines.append(text.strip())
                    rPr = p._p.find('.//' + qn('w:rPr'))
                    pPr = p._p.find(qn('w:pPr'))
                    rp  = get_run_props(rPr)
                    pp  = get_para_props(pPr)
                    sk  = self._get_or_create_style(rp, pp, f'header_line')
                    hdr_styles.append(sk)

            if hdr_lines:
                self.blocks.append({
                    'identifier':  f'{self.name}_header',
                    'name':        f'{self.name} Header',
                    'block_type':  'HEADER',
                    'style_key':   hdr_styles[0] if hdr_styles else 'footer_text',
                    'content':     '\n'.join(hdr_lines),
                    'row_styles':  hdr_styles,
                    'is_active':   True,
                })

            # Footer
            ftr_lines  = []
            ftr_styles = []
            for p in section.footer.paragraphs:
                text = get_text(p._p)
                if text.strip():
                    ftr_lines.append(text.strip())
                    rPr = p._p.find('.//' + qn('w:rPr'))
                    pPr = p._p.find(qn('w:pPr'))
                    rp  = get_run_props(rPr)
                    pp  = get_para_props(pPr)
                    sk  = self._get_or_create_style(rp, pp, 'footer_line')
                    ftr_styles.append(sk)

            if ftr_lines:
                self.blocks.append({
                    'identifier': f'{self.name}_footer',
                    'name':       f'{self.name} Footer',
                    'block_type': 'FOOTER',
                    'style_key':  ftr_styles[0] if ftr_styles else 'footer_text',
                    'content':    '\n'.join(ftr_lines),
                    'row_styles': ftr_styles,
                    'is_active':  True,
                })
            break

    def _extract_textboxes(self):
        """Textboxen aus <w:txbxContent>."""
        seen = set()
        for i, txbx in enumerate(self.body.findall('.//' + qn('w:txbxContent'))):
            lines      = []
            row_styles = []
            content_key = []

            for p in txbx.findall('.//' + qn('w:p')):
                text = get_text(p)
                content_key.append(text)
                rPr  = p.find('.//' + qn('w:rPr'))
                pPr  = p.find(qn('w:pPr'))
                rp   = get_run_props(rPr)
                pp   = get_para_props(pPr)
                sk   = self._get_or_create_style(rp, pp, f'txbx_{i}')
                if text.strip():
                    lines.append(text.strip())
                    row_styles.append(sk)

            # Duplikate überspringen
            key = '|'.join(content_key)
            if key in seen:
                continue
            seen.add(key)

            # Position aus übergeordnetem Drawing
            anchor = None
            parent = txbx.getparent()
            while parent is not None:
                xfrm = parent.find('.//' + qn('a:xfrm'))
                if xfrm is not None:
                    off = xfrm.find(qn('a:off'))
                    ext = xfrm.find(qn('a:ext'))
                    if off is not None and ext is not None:
                        anchor = {
                            'x_cm': emu_to_cm(int(off.get('x',0))),
                            'y_cm': emu_to_cm(int(off.get('y',0))),
                            'w_cm': emu_to_cm(int(ext.get('cx',0))),
                            'h_cm': emu_to_cm(int(ext.get('cy',0))),
                        }
                    break
                parent = parent.getparent()

            block = {
                'identifier': f'{self.name}_textbox_{i}',
                'name':       f'{self.name} Textbox {i}',
                'block_type': 'PARAGRAPH',
                'style_key':  row_styles[0] if row_styles else '',
                'content':    '\n'.join(lines),
                'row_styles': row_styles,
                'is_active':  True,
            }
            if anchor:
                block['_position'] = anchor

            self.blocks.append(block)

    def _extract_tables(self):
        """Tabellen."""
        for i, tbl in enumerate(self.doc.tables):
            lines      = []
            col_styles = []
            col_widths = []

            # Spaltenbreiten aus erster Zeile
            if tbl.rows:
                for cell in tbl.rows[0].cells:
                    tc  = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is not None:
                            w_type = tcW.get(qn('w:type'), 'dxa')
                            w_val  = int(tcW.get(qn('w:w'), 0))
                            if w_type == 'dxa':
                                col_widths.append(dxa_to_cm(w_val))
                            elif w_type == 'pct':
                                col_widths.append(round(w_val / 5000 * 15, 1))

            # Inhalt
            for row in tbl.rows:
                parts = []
                for j, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    parts.append(text)
                    if i == 0:  # Ersten Zeile für col_styles
                        rPr = cell._tc.find('.//' + qn('w:rPr'))
                        pPr = cell._tc.find('.//' + qn('w:pPr'))
                        rp  = get_run_props(rPr)
                        pp  = get_para_props(pPr)
                        sk  = self._get_or_create_style(rp, pp, f'tbl_{i}_col{j}')
                        col_styles.append(sk)
                lines.append('|'.join(parts))

            # layout_ref erstellen
            ref_key = f'table_{i}'
            self.layout['layout_refs'][ref_key] = {
                'column_widths_cm': col_widths
            }

            self.blocks.append({
                'identifier':   f'{self.name}_table_{i}',
                'name':         f'{self.name} Tabelle {i}',
                'block_type':   'TABLE',
                'style_key':    col_styles[0] if col_styles else 'body_text',
                'content':      '\n'.join(lines),
                'col_styles':   col_styles,
                'col_alignments': ['left'] * len(col_widths),
                'layout_ref':   ref_key,
                'is_active':    True,
            })

    def _extract_paragraphs(self):
        """Alle Paragraphen im Body."""
        for i, p in enumerate(self.doc.paragraphs):
            text = get_text(p)
            if not text.strip():
                continue

            rPr = p._p.find('.//' + qn('w:rPr'))
            pPr = p._p.find(qn('w:pPr'))
            rp  = get_run_props(rPr)
            pp  = get_para_props(pPr)
            sk  = self._get_or_create_style(rp, pp, f'para_{i}')

            self.blocks.append({
                'identifier': f'{self.name}_p{i}',
                'name':       f'{self.name} Paragraph {i}',
                'block_type': 'PARAGRAPH',
                'style_key':  sk,
                'content':    text.strip(),
                'row_styles': [sk],
                'is_active':  True,
            })

    def _get_or_create_style(self, rp: dict, pp: dict, hint: str) -> str:
        """
        Findet existierenden Style oder erstellt neuen.
        Gibt style_key zurueck.
        """
        # Fingerprint
        props = {**rp, **pp}
        fp = json.dumps(props, sort_keys=True)

        if fp in self._style_map:
            return self._style_map[fp]

        # Bekannte Styles mappen
        font_size = props.get('font_size_pt', 10.0)
        bold      = props.get('bold', False)
        color     = props.get('color_hex', '1A1A1A')
        align     = props.get('alignment', 'left')

        # Heuristik fuer bekannte Styles
        if font_size >= 14 and bold:
            sk = 'doc_title'
        elif font_size >= 12 and bold:
            sk = 'section_head'
        elif font_size <= 7:
            sk = 'footer_text'
        elif font_size <= 8 and color == '163258':
            sk = 'logo_contact_title'
        elif font_size <= 6:
            sk = 'logo_contact_tagline'
        elif font_size <= 7 and color not in ('163258',):
            sk = 'logo_contact_line'
        elif bold and color == '163258':
            sk = 'label_blue'
        elif bold:
            sk = 'party_bold'
        elif align == 'right':
            sk = 'body_text_right'
        else:
            sk = 'body_text'

        # Style-Definition erstellen
        style_def = {
            'style_key':      sk,
            'style_type':     'TEXT',
            'name':           sk.replace('_', ' ').title(),
            'font_family':    props.get('font_family', 'Arial'),
            'font_size_pt':   font_size,
            'bold':           bold,
            'italic':         props.get('italic', False),
            'underline':      props.get('underline', False),
            'color_hex':      color,
            'alignment':      align,
            'space_before_pt':props.get('space_before_pt', 0.0),
            'space_after_pt': props.get('space_after_pt', 4.0),
            'border_bottom':  props.get('border_bottom', False),
        }
        if props.get('border_bottom'):
            style_def['border_bottom_color'] = props.get('border_bottom_color', '163258')
            style_def['border_bottom_pt']    = props.get('border_bottom_pt', 0.5)

        # Nur hinzufuegen wenn noch nicht vorhanden
        existing_keys = {s['style_key'] for s in self.styles}
        if sk not in existing_keys:
            self.styles.append(style_def)

        self._style_map[fp] = sk
        return sk

    def print_analysis(self):
        """Gibt Analyse-Report aus."""
        result = self.extract()

        print('\n' + '='*60)
        print(f'ANALYSE: {self.docx_path}')
        print('='*60)

        print('\n── LAYOUT ──')
        for k, v in result['layout'].items():
            if k not in ('layout_refs', 'image_refs'):
                print(f'  {k}: {v}')

        print('\n── STYLES ──')
        for s in result['styles']:
            print(f"  {s['style_key']:25s} sz={s['font_size_pt']}pt"
                  f" bold={s['bold']} color=#{s['color_hex']}"
                  f" align={s['alignment']}")

        print('\n── BLOECKE ──')
        for b in result['blocks']:
            pos = b.get('_position', '')
            pos_str = f" pos=({pos['x_cm']:.1f},{pos['y_cm']:.1f})cm size=({pos['w_cm']:.1f}x{pos['h_cm']:.1f})cm" if pos else ''
            print(f"  [{b['block_type']:12s}] {b['identifier']:30s}{pos_str}")
            if b.get('content'):
                for j, line in enumerate(b['content'].split('\n')[:3]):
                    sk = b['row_styles'][j] if j < len(b.get('row_styles',[])) else ''
                    print(f"    Zeile {j}: {line!r:40s} → {sk}")
                if b['content'].count('\n') > 2:
                    print(f"    ... + {b['content'].count(chr(10))-2} weitere Zeilen")

        print(f'\nTotal: {len(result["styles"])} Styles, {len(result["blocks"])} Bloecke')

    def write_to_template(self, out_dir: str, prefix: str = ''):
        """Schreibt Ergebnis in Template-Verzeichnis."""
        result = self.extract()
        os.makedirs(out_dir, exist_ok=True)

        # layout.json
        lay_path = os.path.join(out_dir, 'layout.json')
        if os.path.exists(lay_path):
            with open(lay_path) as f:
                existing = json.load(f)
            # Merge — vorhandene Werte behalten, neue ergaenzen
            for k, v in result['layout'].items():
                if k not in existing:
                    existing[k] = v
            result['layout'] = existing
        with open(lay_path, 'w', encoding='utf-8') as f:
            json.dump(result['layout'], f, indent=4, ensure_ascii=False)
        print(f'  OK: {lay_path}')

        # styles.json — merged
        styles_path = os.path.join(out_dir, 'styles.json')
        if os.path.exists(styles_path):
            with open(styles_path) as f:
                existing = json.load(f)
            existing_keys = {s['style_key'] for s in existing.get('definitions', [])}
            for s in result['styles']:
                if s['style_key'] not in existing_keys:
                    existing.setdefault('definitions', []).append(s)
                    print(f"    + Style: {s['style_key']}")
            result_styles = existing
        else:
            result_styles = {
                'identifier': 'extracted_styles',
                'name':       f'Extrahiert aus {self.name}',
                'definitions': result['styles']
            }
        with open(styles_path, 'w', encoding='utf-8') as f:
            json.dump(result_styles, f, indent=4, ensure_ascii=False)
        print(f'  OK: {styles_path}')

        # blocks.json — merged
        blocks_path = os.path.join(out_dir, 'blocks.json')
        if os.path.exists(blocks_path):
            with open(blocks_path) as f:
                existing = json.load(f)
            existing_ids = {b['identifier'] for b in existing}
            added = 0
            for b in result['blocks']:
                bid = f"{prefix}_{b['identifier']}" if prefix else b['identifier']
                b['identifier'] = bid
                if bid not in existing_ids:
                    existing.append(b)
                    added += 1
            result_blocks = existing
            print(f'    + {added} neue Bloecke')
        else:
            result_blocks = result['blocks']
        with open(blocks_path, 'w', encoding='utf-8') as f:
            json.dump(result_blocks, f, indent=4, ensure_ascii=False)
        print(f'  OK: {blocks_path}')


def main():
    parser = argparse.ArgumentParser(description='DOCX Extraktor')
    parser.add_argument('--docx',    required=True, help='Pfad zum DOCX')
    parser.add_argument('--out',     help='Output-Verzeichnis (Template-Dir)')
    parser.add_argument('--name',    default='', help='Prefix fuer Block-Identifier')
    parser.add_argument('--analyze', action='store_true', help='Nur analysieren')
    args = parser.parse_args()

    ext = DocxExtractor(args.docx)

    if args.analyze or not args.out:
        ext.print_analysis()
    else:
        print(f'\nExtrahiere: {args.docx} → {args.out}')
        ext.write_to_template(args.out, prefix=args.name)
        print('Fertig.')


if __name__ == '__main__':
    main()
