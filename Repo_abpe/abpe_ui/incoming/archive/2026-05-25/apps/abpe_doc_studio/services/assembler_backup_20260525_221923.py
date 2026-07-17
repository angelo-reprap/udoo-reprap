"""
services/assembler.py
=====================
Generischer OOXML-Assembler.

Implementiert die 10 WordprocessingML-Grundbausteine:
  1. _render_paragraph()       <w:p>
  2. _render_table()           <w:tbl>
  3. _render_drawing()         <w:drawing>
  4. _render_field()           <w:fldChar>
  5. _render_page_break()      <w:br type="page">
  6. _render_content_control() <w:sdt>
  7. _render_section()         <w:sectPr>  → via _apply_page_layout()
  8. _render_header_footer()   <w:hdr>/<w:ftr>
  9. _render_hyperlink()       <w:hyperlink>
  10. _render_bookmark()       <w:bookmarkStart>/<w:bookmarkEnd>

Keine hardcoded Werte. Alles aus Konfiguration:
  block.content      → Text
  block.row_styles   → Style pro Zeile
  block.col_styles   → Style pro Spalte
  block.layout_ref   → Geometrie aus layout.layout_refs
  block.image_ref    → Bildpfad aus layout.image_refs
  StyleDefinition    → Formatierung
  PageLayout         → Seitenmaße + layout_refs + image_refs
"""
import io
import logging
import os

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml    import OxmlElement
from lxml import etree

# Fehlende Namespaces fuer python-docx registrieren
_extra_ns = {
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
}
from docx.oxml.ns import nsmap as _nsmap
for _pfx, _uri in _extra_ns.items():
    if _pfx not in _nsmap:
        _nsmap[_pfx] = _uri

from .context_loader  import ContextLoader
from .variable_engine import VariableEngine
from .exporter        import DocExporter

log = logging.getLogger('abpe_doc_studio.assembler')

DXA = 567  # 1cm = 567 DXA (Word-interne Einheit)


def _rgb(hex_str: str) -> RGBColor:
    """Hex-String zu RGBColor."""
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def _cm(cm: float) -> int:
    """Zentimeter zu DXA."""
    return int(cm * DXA)


class StyleCache:
    """Cached StyleDefinitions pro StyleKit."""

    def __init__(self):
        self._cache = {}

    def load(self, style_kit) -> dict:
        kid = style_kit.pk if style_kit else None
        if kid not in self._cache:
            self._cache[kid] = (
                {s.style_key: s for s in style_kit.definitions.all()}
                if style_kit else {}
            )
        return self._cache[kid]

    def get(self, style_kit, style_key: str):
        if not style_key or not style_kit:
            return None
        return self.load(style_kit).get(style_key)


class DocAssembler:

    def __init__(self):
        self.context_loader  = ContextLoader()
        self.variable_engine = VariableEngine()
        self.exporter        = DocExporter()
        self.style_cache     = StyleCache()

    # ── Oeffentliche API ──────────────────────────────────────────────────

    def generate(self, template_identifier: str,
                  variables: dict = None,
                  context_ref: str = '',
                  scope: str = '',
                  engine: str = 'BOTH',
                  user=None) -> dict:

        from apps.abpe_doc_studio.models import (
            DocTemplate, DocLog, DocStatus, LogStatus
        )
        tpl = DocTemplate.objects.filter(
            identifier=template_identifier,
            status=DocStatus.ACTIVE
        ).select_related('layout', 'style_kit').prefetch_related(
            'template_blocks__block__style_kit__definitions'
        ).first()

        if not tpl:
            raise ValueError(f"Template '{template_identifier}' nicht gefunden")

        all_vars = self.context_loader.load(
            scope           = scope or tpl.scope,
            context_ref     = context_ref,
            extra_variables = variables or {},
        )

        doc       = self._build_docx(tpl, all_vars)
        result    = {'success': True}
        docx_info = self.exporter.save_docx(
            doc, template_identifier,
            scope=scope or tpl.scope, context_ref=context_ref,
        )
        result['file_path_docx'] = docx_info['file_path']
        result['file_size']      = docx_info['file_size']

        if engine in ('PDF', 'BOTH'):
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    docx_info['file_path']
                )
                result['file_path_pdf'] = pdf_info['file_path']
            except Exception as e:
                log.warning(f'PDF fehlgeschlagen: {e}')
                result['file_path_pdf'] = ''
                result['pdf_error']     = str(e)

        doc_log = DocLog.objects.create(
            template         = tpl,
            template_version = tpl.active_version,
            engine_used      = engine,
            file_path_docx   = result.get('file_path_docx', ''),
            file_path_pdf    = result.get('file_path_pdf', ''),
            file_size_bytes  = result.get('file_size', 0),
            context_ref      = context_ref,
            scope            = scope or tpl.scope,
            variables_used   = {
                k: v for k, v in all_vars.items()
                if not isinstance(v, list)
            },
            status       = LogStatus.OK,
            generated_by = user,
        )
        result['log_id'] = str(doc_log.log_id)
        log.info(f'Generiert: {template_identifier} [{doc_log.log_id}]')
        return result

    def render_to_bytes(self, template, variables: dict = None,
                         engine: str = 'DOCX') -> bytes:
        from apps.abpe_doc_studio.models import DocTemplate
        if isinstance(template, str):
            template = DocTemplate.objects.filter(
                identifier=template
            ).select_related('layout', 'style_kit').prefetch_related(
                'template_blocks__block__style_kit__definitions'
            ).first()
            if not template:
                raise ValueError('Template nicht gefunden')

        doc = self._build_docx(template, variables or {})
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        if engine == 'PDF':
            import tempfile
            with tempfile.NamedTemporaryFile(
                suffix='.docx', delete=False
            ) as tmp:
                tmp.write(buf.getvalue())
                tmp_path = tmp.name
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    tmp_path, output_dir=os.path.dirname(tmp_path)
                )
                return self.exporter.read_as_bytes(pdf_info['file_path'])
            finally:
                os.unlink(tmp_path)
                pdf_path = tmp_path.replace('.docx', '.pdf')
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)

        return buf.getvalue()

    # ── DOCX bauen ────────────────────────────────────────────────────────

    def _build_docx(self, tpl, all_vars: dict) -> DocxDocument:
        """
        Baut DOCX aus Template-Konfiguration.
        Liest alles aus DB — kein hardcoded Wert.
        """
        doc = DocxDocument()
        doc._sk  = tpl.style_kit
        doc._lay = tpl.layout
        doc._logo_block_id   = tpl.logo_block_id
        doc._header_block_id = tpl.header_block_id
        doc._footer_block_id = tpl.footer_block_id
        doc._template_id  = tpl.identifier
        doc._template_dir = tpl.template_dir or ''
        doc._last_vars    = {}

        # 7. SECTION — Seitenmaße aus PageLayout
        self._render_section(doc)

        # 8. HEADER — aus header_block_id
        self._render_header_footer(doc, 'header', all_vars)

        doc._last_vars = all_vars

        doc._last_vars = all_vars

        # Body-Bloecke — nur slot=body
        for tb in tpl.template_blocks.order_by('slot', 'order'):
            if tb.slot != 'body':
                continue
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue
            if tb.page_break_before and doc.paragraphs:
                doc.add_page_break()
            self._dispatch(doc, tb, all_vars)

        # 8. FOOTER — slot=footer Bloecke in Word-Footer-Bereich rendern
        self._render_footer_from_blocks(doc, tpl, all_vars)

        return doc

    # ── _apply_style: EINZIGE Formatierungsquelle ─────────────────────────

    def _apply_style(self, paragraph, style_key: str,
                      style_kit, run=None) -> None:
        """
        Liest StyleDefinition aus DB.
        Einzige Stelle fuer Formatierung — kein Pt/bold/color sonst.
        """
        sdef = self.style_cache.get(style_kit, style_key)
        if not sdef:
            return

        if sdef.space_before_pt:
            paragraph.paragraph_format.space_before = Pt(sdef.space_before_pt)
        if sdef.space_after_pt:
            paragraph.paragraph_format.space_after  = Pt(sdef.space_after_pt)
        if sdef.indent_left_cm:
            paragraph.paragraph_format.left_indent  = Cm(sdef.indent_left_cm)

        align_map = {
            'left':    WD_ALIGN_PARAGRAPH.LEFT,
            'right':   WD_ALIGN_PARAGRAPH.RIGHT,
            'center':  WD_ALIGN_PARAGRAPH.CENTER,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if sdef.alignment in align_map:
            paragraph.alignment = align_map[sdef.alignment]

        if sdef.border_bottom:
            self._add_border(paragraph, sdef.border_bottom_color,
                             sz=int(sdef.border_bottom_pt * 8),
                             style=sdef.border_bottom_style or 'single')

        if run is not None:
            run.font.name      = sdef.font_family
            run.font.size      = Pt(sdef.font_size_pt)
            run.bold           = sdef.bold
            run.italic         = sdef.italic
            run.underline      = sdef.underline
            run.font.color.rgb = _rgb(sdef.color_hex)

    def _add_border(self, paragraph, color: str,
                     side: str = 'bottom',
                     sz: int = 4,
                     style: str = 'single') -> None:
        pPr  = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        el   = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   style)
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color.lstrip('#'))
        pBdr.append(el)
        pPr.append(pBdr)

    # ── Block-Dispatch ────────────────────────────────────────────────────

    def _dispatch(self, doc, tb, variables: dict) -> None:
        """Verteilt auf die 10 OOXML-Grundbausteine."""
        block   = tb.block
        bt      = block.block_type
        content = tb.content_override or block.content or ''
        rendered = self.variable_engine.render_text(content, variables)

        dispatch = {
            'PARAGRAPH':       lambda: self._render_paragraph(
                                   doc, rendered, block, variables),
            'TABLE':           lambda: self._render_table(
                                   doc, rendered, block, variables),
            'DRAWING':         lambda: self._render_drawing(
                                   doc, block),
            'FIELD':           lambda: self._render_field(
                                   doc, block),
            'PAGE_BREAK':      lambda: self._render_page_break(doc),
            'CONTENT_CONTROL': lambda: self._render_content_control(
                                   doc, rendered, block),
            'HEADER':          lambda: None,  # in _build_docx behandelt
            'FOOTER':          lambda: None,  # in _build_docx behandelt
            'HYPERLINK':       lambda: self._render_hyperlink(
                                   doc, rendered, block, variables),
            'BOOKMARK':        lambda: self._render_bookmark(
                                   doc, block),
            'TEXTBOX':         lambda: self._render_textbox(
                                   doc, rendered, block),
        }
        fn = dispatch.get(bt)
        if fn:
            fn()
        else:
            log.warning(f'Unbekannter block_type: {bt}')

    # ── 1. PARAGRAPH <w:p> ────────────────────────────────────────────────

    def _render_paragraph(self, doc, content: str,
                           block, variables: dict,
                           container=None) -> None:
        """
        Rendert Paragraph-Block.
        row_styles[i] bestimmt Style fuer Zeile i.
        Letzter Style in row_styles wird fuer alle weiteren Zeilen wiederholt.
        """
        sk         = block.style_kit
        row_styles = block.row_styles or []
        fallback   = block.style_key or ''

        lines = content.split('\n') if content else []
        target = container or doc

        for i, line in enumerate(lines):
            # Style fuer diese Zeile
            if row_styles:
                skey = row_styles[i] if i < len(row_styles) else row_styles[-1]
            else:
                skey = fallback

            if container is not None:
                p = container.add_paragraph()
            else:
                p = doc.add_paragraph()

            if line.strip():
                run = p.add_run(line.strip())
                self._apply_style(p, skey, sk, run)
            else:
                # Leerzeile
                sdef = self.style_cache.get(sk, skey)
                if sdef and sdef.space_after_pt:
                    p.paragraph_format.space_after = Pt(
                        sdef.space_after_pt / 2
                    )

    # ── 2. TABLE <w:tbl> ─────────────────────────────────────────────────

    def _render_table(self, doc, content: str,
                       block, variables: dict) -> None:
        """
        Rendert Table-Block.

        Zwei Modi:
        A) Dynamische Tabelle: block.columns + Daten aus variables
           (fuer TIME_TABLE / AP_TABLE)
        B) Statische Tabelle: block.content als 'sp1|sp2|sp3' pro Zeile
           (fuer LABEL_VALUE, SIGNATURE, TOTAL_BLOCK, INV_HEADER etc.)
        """
        sk         = block.style_kit
        lay        = doc._lay
        layout_ref = block.layout_ref or ''
        geo        = lay.layout_refs.get(layout_ref, {}) if lay.layout_refs else {}
        col_styles = block.col_styles or []
        col_aligns = block.col_alignments or []
        row_bdr    = block.row_border_bottom or []
        bb_sdef    = self.style_cache.get(sk, 'border_signature')

        if block.columns:
            # Modus A: Dynamische Tabelle mit columns-Definition
            self._render_dynamic_table(doc, block, variables, geo, sk)
        else:
            # Modus B: Statische Tabelle aus content
            self._render_static_table(
                doc, content, sk, geo,
                col_styles, col_aligns, row_bdr, bb_sdef
            )

    def _render_static_table(self, doc, content: str,
                               sk, geo: dict,
                               col_styles: list,
                               col_aligns: list,
                               row_bdr: list,
                               bb_sdef) -> None:
        """
        Statische Tabelle aus 'sp1|sp2|sp3' pro Zeile.
        Spaltenbreiten aus geo['column_widths_cm'].
        """
        lines = [l for l in content.split('\n') if l.strip()]
        if not lines:
            return

        col_widths_cm = geo.get('column_widths_cm', [])
        n_cols = max(
            len(col_widths_cm),
            max((len(l.split('|')) for l in lines), default=1)
        )

        tbl = doc.add_table(rows=len(lines), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for ri, line in enumerate(lines):
            parts = line.split('|')
            # Parts auf n_cols auffuellen
            while len(parts) < n_cols:
                parts.append('')

            for ci in range(n_cols):
                cell = tbl.rows[ri].cells[ci]

                # Breite setzen
                if ci < len(col_widths_cm):
                    self._set_cell_width(cell, _cm(col_widths_cm[ci]))
                self._set_cell_borders_none(cell)

                # Stil
                skey = col_styles[ci] if ci < len(col_styles) else (
                    col_styles[-1] if col_styles else ''
                )
                # Ausrichtung
                align_str = col_aligns[ci] if ci < len(col_aligns) else 'left'

                p   = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)

                align_map = {
                    'left':   WD_ALIGN_PARAGRAPH.LEFT,
                    'right':  WD_ALIGN_PARAGRAPH.RIGHT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                }
                p.alignment = align_map.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)

                text = parts[ci].strip()
                if text:
                    run = p.add_run(text)
                    self._apply_style(p, skey, sk, run)

                # Unterlinie bei bestimmten Zeilen
                if ri in row_bdr and ci != 1 and bb_sdef:
                    self._add_border(p, bb_sdef.border_bottom_color,
                                     sz=int(bb_sdef.border_bottom_pt * 8))

        # Abstand nach Tabelle
        sp = self.style_cache.get(sk, 'spacing_normal')
        p_gap = doc.add_paragraph()
        if sp:
            p_gap.paragraph_format.space_after = Pt(sp.space_after_pt)

    def _render_dynamic_table(self, doc, block,
                               variables: dict,
                               geo: dict, sk) -> None:
        """
        Dynamische Tabelle mit columns-Definition + Daten aus variables.
        Header aus columns[].label, Zeilen aus variables-Liste.
        """
        columns    = block.columns
        data_key   = None

        # Daten-Schluessel aus expected_variables ermitteln
        for ev in (block.expected_variables or []):
            if ev.get('type') == 'list':
                data_key = ev.get('name')
                break
        if not data_key:
            # Fallback: ersten Schluessel aus columns suchen
            data_key = 'positionen'

        rows = self.variable_engine.expand_table_rows(
            columns, data_key, variables
        )

        lay        = doc._lay
        col_widths = geo.get('column_widths_cm', [])
        alt_bg     = geo.get('alt_row_color', '') or lay.table_alt_row_color or ''
        white_bg   = geo.get('white_row_color','') or lay.table_white_row_color or ''

        # Tabellenbreite aus layout
        content_w = (lay.column_widths_cm[0]
                     if lay.column_widths_cm
                     else lay.page_width_cm - lay.margin_left_cm - lay.margin_right_cm)
        total_dxa = _cm(content_w)

        n_cols    = len(columns)
        tbl       = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Tabellenbreite setzen
        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'),    str(total_dxa))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        total_pct = sum(c.get('width_pct', 20) for c in columns)
        hdr_sdef  = self.style_cache.get(sk, 'table_header')
        hdr_bg    = hdr_sdef.table_header_bg_hex if hdr_sdef else ''

        # Header-Zeile
        for ci, col in enumerate(columns):
            cell = tbl.rows[0].cells[ci]
            w    = int(total_dxa * col.get('width_pct', 20) / total_pct)
            self._set_cell_width(cell, w)
            if hdr_bg:
                self._set_cell_bg(cell, hdr_bg)
            self._set_cell_borders_none(cell)
            p   = cell.paragraphs[0]
            run = p.add_run(col.get('label', ''))
            self._apply_style(p, 'table_header', sk, run)

        # Daten-Zeilen
        for ri, row_data in enumerate(rows):
            drow = tbl.rows[ri + 1]
            bg   = alt_bg if ri % 2 == 0 else white_bg
            for ci, col in enumerate(columns):
                cell  = drow.cells[ci]
                value = row_data[ci] if ci < len(row_data) else ''
                w     = int(total_dxa * col.get('width_pct', 20) / total_pct)
                self._set_cell_width(cell, w)
                if bg:
                    self._set_cell_bg(cell, bg)
                self._set_cell_borders_none(cell)
                p   = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                               if col.get('align') == 'right'
                               else WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(value)
                self._apply_style(p, 'table_body', sk, run)

        sp = self.style_cache.get(sk, 'spacing_normal')
        p_gap = doc.add_paragraph()
        if sp:
            p_gap.paragraph_format.space_after = Pt(sp.space_after_pt)

    # ── 3. DRAWING <w:drawing> ────────────────────────────────────────────

    def _render_drawing(self, doc, block,
                         container=None) -> None:
        """
        Rendert Bild aus layout.image_refs[block.image_ref].
        Geometrie aus layout.layout_refs[block.layout_ref].
        """
        from django.conf import settings

        lay        = doc._lay
        image_ref  = block.image_ref or ''
        layout_ref = block.layout_ref or ''

        # Bildpfad aus layout.image_refs
        image_refs = lay.image_refs or {}
        rel_path   = image_refs.get(image_ref, '')
        if not rel_path:
            log.warning(f'image_ref nicht gefunden: {image_ref!r}')
            return

        # Pfad relativ zu BASE_DIR oder relativ zum Template-Verzeichnis
        img_path = os.path.join(settings.BASE_DIR, rel_path)
        if not os.path.exists(img_path):
            # Fallback: relativ zum Template-Verzeichnis suchen
            # Template-Identifier → Verzeichnisname aus DB
            tpl_dir_name = getattr(doc, '_template_dir', '')
            if tpl_dir_name:
                tpl_base = os.path.join(
                    settings.BASE_DIR,
                    'apps', 'abpe_doc_studio',
                    'generator', 'templates', tpl_dir_name
                )
                img_path = os.path.join(tpl_base, rel_path)

        if not os.path.exists(img_path):
            # Fallback 2: _shared suchen
            shared_base = os.path.join(
                settings.BASE_DIR,
                'apps', 'abpe_doc_studio',
                'generator', 'templates', '_shared'
            )
            img_path = os.path.join(shared_base, rel_path)

        if not os.path.exists(img_path):
            log.warning(f'Bilddatei nicht gefunden: {img_path}')
            return

        # Geometrie aus layout.layout_refs
        geo        = (lay.layout_refs or {}).get(layout_ref, {})
        height_cm  = geo.get('image_height_cm', 2.0)
        align_str  = geo.get('alignment', 'right')

        align_map = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
        }

        if container is not None:
            p = container.add_paragraph()
        else:
            p = doc.add_paragraph()

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

        # Anchor-Position aus layout_ref
        pos_h_cm = geo.get('position_h_cm', 0.0)

        if pos_h_cm > 0:
            # Absolut positioniertes Drawing (wp:anchor) wie im Original
            self._add_anchored_picture(p, img_path,
                                       height_cm=height_cm,
                                       width_cm=geo.get('width_cm', 6.0),
                                       pos_h_cm=pos_h_cm,
                                       pos_v_cm=0.5)
        else:
            # Inline Drawing
            p.alignment = align_map.get(align_str, WD_ALIGN_PARAGRAPH.RIGHT)
            p.add_run().add_picture(img_path, height=Cm(height_cm))

    # ── 4. FIELD <w:fldChar> ─────────────────────────────────────────────

    def _render_field(self, doc, block,
                       container=None) -> None:
        """
        Rendert Feldfunktion (PAGE_NUMBER, TOTAL_PAGES, DATE etc.).
        Format aus block.content, Style aus block.style_key.
        """
        sk         = block.style_kit
        style_key  = block.style_key or 'footer_text'
        field_type = block.field_type or ''
        content    = block.content or ''
        sdef       = self.style_cache.get(sk, style_key)

        # Feldinstruktionen je Typ
        field_map = {
            'PAGE_NUMBER': ' PAGE ',
            'TOTAL_PAGES': ' NUMPAGES ',
            'DATE':        r' DATE \@ "dd.MM.yyyy" ',
            'TIME':        r' TIME \@ "HH:mm" ',
            'AUTHOR':      ' AUTHOR ',
            'FILENAME':    ' FILENAME ',
        }
        instr = field_map.get(field_type, ' PAGE ')

        if container is not None:
            p = container.add_paragraph()
        else:
            p = doc.add_paragraph()

        # Format-String parsen: "Seite {PAGE} von {NUMPAGES}"
        # Oder einfach Feldfunktion direkt
        if '{PAGE}' in content or '{NUMPAGES}' in content:
            parts = content.replace('{NUMPAGES}', '\x01').split('{PAGE}')
            before = parts[0]
            after  = parts[1] if len(parts) > 1 else ''
            between = after.split('\x01')[0] if '\x01' in after else ''
            end     = after.split('\x01')[1] if '\x01' in after else ''

            if before:
                r = p.add_run(before)
                if sdef:
                    r.font.name = sdef.font_family
                    r.font.size = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)
            self._add_field_char(p, ' PAGE ', sdef)
            if between:
                r = p.add_run(between)
                if sdef:
                    r.font.name = sdef.font_family
                    r.font.size = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)
            self._add_field_char(p, ' NUMPAGES ', sdef)
            if end:
                r = p.add_run(end)
                if sdef:
                    r.font.name = sdef.font_family
                    r.font.size = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)
        else:
            self._add_field_char(p, instr, sdef)

        self._apply_style(p, style_key, sk)

    def _add_field_char(self, paragraph, instr: str, sdef) -> None:
        """Fuegt eine Feldfunktion in einen Paragraph ein."""
        for ftype, text in [
            ('begin', instr), ('separate', None), ('end', None)
        ]:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), ftype)
            run = paragraph.add_run()
            run._r.append(fld)
            if sdef:
                run.font.name = sdef.font_family
                run.font.size = Pt(sdef.font_size_pt)
                run.font.color.rgb = _rgb(sdef.color_hex)
            if text:
                ins = OxmlElement('w:instrText')
                ins.text = text
                run._r.append(ins)

    # ── 5. PAGE_BREAK <w:br> ─────────────────────────────────────────────

    def _render_page_break(self, doc) -> None:
        """Fuegt Seitenumbruch ein."""
        doc.add_page_break()

    # ── 6. CONTENT_CONTROL <w:sdt> ───────────────────────────────────────

    def _render_content_control(self, doc, content: str,
                                  block) -> None:
        """
        Fuegt Content Control ein.
        control_title und control_id aus block-Konfiguration.
        """
        sk            = block.style_kit
        style_key     = block.style_key or 'body_text'
        control_title = block.control_title or ''
        control_id    = block.control_id or '1000'

        sdt    = OxmlElement('w:sdt')
        sdtPr  = OxmlElement('w:sdtPr')

        # Alias (Titel)
        alias = OxmlElement('w:alias')
        alias.set(qn('w:val'), control_title)
        sdtPr.append(alias)

        # ID
        cid = OxmlElement('w:id')
        cid.set(qn('w:val'), str(control_id))
        sdtPr.append(cid)

        sdt.append(sdtPr)

        sdtContent = OxmlElement('w:sdtContent')
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = content
        r.append(t)
        p.append(r)
        sdtContent.append(p)
        sdt.append(sdtContent)

        doc.element.body.append(sdt)

    # ── 7. SECTION <w:sectPr> ────────────────────────────────────────────

    def _render_section(self, doc: DocxDocument) -> None:
        """
        Setzt Seitenmaße aus PageLayout.
        Normal-Style aus layout.normal_font_size_pt + body_text StyleDef.
        """
        lay = doc._lay
        sk  = doc._sk

        for section in doc.sections:
            section.page_width      = Cm(lay.page_width_cm)
            section.page_height     = Cm(lay.page_height_cm)
            section.left_margin     = Cm(lay.margin_left_cm)
            section.right_margin    = Cm(lay.margin_right_cm)
            section.top_margin      = Cm(lay.margin_top_cm)
            section.bottom_margin   = Cm(lay.margin_bottom_cm)
            section.header_distance = Cm(lay.header_distance_cm)
            section.footer_distance = Cm(lay.footer_distance_cm)

        # Normal-Style aus body_text StyleDefinition
        body_sdef = self.style_cache.get(sk, 'body_text')
        doc.styles['Normal'].font.name = (
            body_sdef.font_family if body_sdef else 'Arial'
        )
        doc.styles['Normal'].font.size = Pt(lay.normal_font_size_pt or 10.0)

    def _render_footer_from_blocks(self, doc, tpl, variables):
        """
        Rendert alle slot=footer Bloecke in den Word-Footer-Bereich.
        Jeder Block wird als Tabelle oder Paragraph in section.footer gerendert.
        """
        footer_blocks = tpl.template_blocks.filter(
            slot='footer'
        ).order_by('order')

        if not footer_blocks.exists():
            # Fallback: footer_block_id
            self._render_header_footer(doc, 'footer', variables)
            return

        sk  = doc._sk
        bb  = self.style_cache.get(sk, 'border_brand')

        for section in doc.sections:
            ftr   = section.footer
            first = True

            for tb in footer_blocks:
                block    = tb.block
                content  = tb.content_override or block.content or ''
                rendered = self.variable_engine.render_text(content, variables)
                bt       = block.block_type
                lay      = doc._lay
                geo      = (lay.layout_refs or {}).get(block.layout_ref, {})
                col_styles = block.col_styles or []
                col_aligns = block.col_alignments or []

                if bt == 'TABLE':
                    lines = [l for l in rendered.split('\n') if l.strip()]
                    if not lines:
                        continue
                    col_widths_cm = geo.get('column_widths_cm', [])
                    n_cols = max(
                        len(col_widths_cm),
                        max((len(l.split('|')) for l in lines), default=1)
                    )
                    # Erste Tabelle: Trennlinie oben
                    if first and bb:
                        p_sep = (ftr.paragraphs[0]
                                 if ftr.paragraphs else ftr.add_paragraph())
                        p_sep.clear()
                        self._add_border(p_sep, bb.border_bottom_color,
                                         side='top',
                                         sz=int(bb.border_bottom_pt * 8))
                        first = False

                    total_w = sum(_cm(c) for c in col_widths_cm) or _cm(15.0)
                    tbl = ftr.add_table(rows=len(lines), cols=n_cols, width=total_w)
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

                    for ri, line in enumerate(lines):
                        parts = line.split('|')
                        while len(parts) < n_cols:
                            parts.append('')
                        for ci in range(n_cols):
                            cell = tbl.rows[ri].cells[ci]
                            if ci < len(col_widths_cm):
                                self._set_cell_width(cell, _cm(col_widths_cm[ci]))
                            self._set_cell_borders_none(cell)
                            skey = (col_styles[ci] if ci < len(col_styles)
                                    else (col_styles[-1] if col_styles else ''))
                            p    = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after  = Pt(0)
                            text = parts[ci].strip()
                            if text:
                                run = p.add_run(text)
                                self._apply_style(p, skey, sk, run)

                elif bt in ('PARAGRAPH', 'FOOTER'):
                    if first and bb:
                        p_sep = (ftr.paragraphs[0]
                                 if ftr.paragraphs else ftr.add_paragraph())
                        p_sep.clear()
                        self._add_border(p_sep, bb.border_bottom_color,
                                         side='top',
                                         sz=int(bb.border_bottom_pt * 8))
                        first = False
                    sdef = self.style_cache.get(sk, block.style_key or 'footer_text')
                    for line in rendered.split('\n'):
                        p = ftr.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if line.strip() and sdef:
                            r = p.add_run(line.strip())
                            r.font.name = sdef.font_family
                            r.font.size = Pt(sdef.font_size_pt)
                            r.font.color.rgb = _rgb(sdef.color_hex)

    # ── 8. HEADER/FOOTER <w:hdr>/<w:ftr> ────────────────────────────────

    def _render_header_footer(self, doc: DocxDocument,
                               which: str,
                               variables: dict) -> None:
        """
        Rendert Header oder Footer aus ContentBlock.
        which: 'header' oder 'footer'

        Block-Identifier aus doc._header_block_id / doc._footer_block_id.
        Der Block kann PARAGRAPH, FIELD, DRAWING oder TABLE sein.
        """
        from apps.abpe_doc_studio.models import ContentBlock

        block_id = (doc._header_block_id
                    if which == 'header'
                    else doc._footer_block_id)

        if not block_id:
            return

        block = ContentBlock.objects.filter(
            identifier=block_id, is_active=True
        ).select_related('style_kit').first()

        if not block:
            return

        sk      = block.style_kit or doc._sk
        content = self.variable_engine.render_text(
            block.content or '', variables
        )
        bb      = self.style_cache.get(sk, 'border_brand')

        for section in doc.sections:
            area = section.header if which == 'header' else section.footer

            # Bestehende Paragraphen leeren
            for p in area.paragraphs:
                p.clear()
            for p in list(area.paragraphs)[1:]:
                p._element.getparent().remove(p._element)

            # Block-Typ bestimmt Rendering
            bt = block.block_type

            if bt == 'FIELD':
                # Feldfunktion im Header (Seitennummer)
                p_nr = area.paragraphs[0]
                p_nr.clear()
                p_nr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._render_field_in_para(p_nr, block, sk)
                if bb:
                    self._add_border(p_nr, bb.border_bottom_color,
                                     sz=int(bb.border_bottom_pt * 8))

            elif bt == 'PARAGRAPH':
                # Text im Header/Footer
                p = area.paragraphs[0]
                p.clear()
                lines = content.split('\n')
                row_styles = block.row_styles or []
                for i, line in enumerate(lines):
                    if i > 0:
                        p = area.add_paragraph()
                    skey = (row_styles[i] if i < len(row_styles)
                            else (row_styles[-1] if row_styles
                                  else block.style_key or ''))
                    if line.strip():
                        run = p.add_run(line.strip())
                        self._apply_style(p, skey, sk, run)
                if which == 'footer' and bb:
                    first_p = area.paragraphs[0]
                    self._add_border(first_p, bb.border_bottom_color,
                                     side='top',
                                     sz=int(bb.border_bottom_pt * 8))

            elif bt == 'FOOTER':
                # Strukturierter Footer: Zeilen aus content
                p = area.paragraphs[0]
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sdef = self.style_cache.get(sk, block.style_key or 'footer_text')
                if bb:
                    self._add_border(p, bb.border_bottom_color,
                                     side='top',
                                     sz=int(bb.border_bottom_pt * 8))
                lines = content.split('\n')
                for i, line in enumerate(lines):
                    if i > 0:
                        p = area.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.strip() and sdef:
                        r = p.add_run(line.strip())
                        r.font.name = sdef.font_family
                        r.font.size = Pt(sdef.font_size_pt)
                        r.font.color.rgb = _rgb(sdef.color_hex)

    def _render_field_in_para(self, paragraph, block, sk) -> None:
        """Hilfsmethode: Feldfunktion direkt in Paragraph rendern."""
        content    = block.content or ''
        style_key  = block.style_key or 'footer_text'
        sdef       = self.style_cache.get(sk, style_key)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if sdef:
            paragraph.paragraph_format.space_before = Pt(sdef.space_before_pt)
            paragraph.paragraph_format.space_after  = Pt(sdef.space_after_pt)

        if '{PAGE}' in content or '{NUMPAGES}' in content:
            parts   = content.replace('{NUMPAGES}', '\x01').split('{PAGE}')
            before  = parts[0]
            rest    = parts[1] if len(parts) > 1 else ''
            between = rest.split('\x01')[0] if '\x01' in rest else ''
            end     = rest.split('\x01')[1] if '\x01' in rest else ''

            def _r(text):
                r = paragraph.add_run(text)
                if sdef:
                    r.font.name = sdef.font_family
                    r.font.size = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)

            if before:  _r(before)
            self._add_field_char(paragraph, ' PAGE ', sdef)
            if between: _r(between)
            self._add_field_char(paragraph, ' NUMPAGES ', sdef)
            if end:     _r(end)
        else:
            field_map = {
                'PAGE_NUMBER': ' PAGE ',
                'TOTAL_PAGES': ' NUMPAGES ',
                'DATE':        r' DATE \@ "dd.MM.yyyy" ',
            }
            instr = field_map.get(block.field_type or '', ' PAGE ')
            self._add_field_char(paragraph, instr, sdef)

    # ── 9. HYPERLINK <w:hyperlink> ───────────────────────────────────────

    def _render_hyperlink(self, doc, content: str,
                           block, variables: dict) -> None:
        """
        Fuegt Hyperlink ein.
        url aus block.url, Anzeigetext aus block.content.
        Style aus block.style_key.
        """
        sk        = block.style_kit
        style_key = block.style_key or 'hyperlink'
        url       = self.variable_engine.render_text(
            block.url or '', variables
        )

        p    = doc.add_paragraph()
        hlink = OxmlElement('w:hyperlink')
        hlink.set(qn('r:id'), '')  # Relationship wird normalerweise gesetzt

        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        run.append(rPr)

        t = OxmlElement('w:t')
        t.text = content
        run.append(t)
        hlink.append(run)

        p._p.append(hlink)
        self._apply_style(p, style_key, sk)

    # ── 10. BOOKMARK <w:bookmarkStart>/<w:bookmarkEnd> ────────────────────

    def _render_bookmark(self, doc, block) -> None:
        """
        Fuegt Lesezeichen ein.
        bookmark_name und bookmark_id aus block-Konfiguration.
        """
        bookmark_name = block.bookmark_name or ''
        bookmark_id   = block.bookmark_id or '0'

        p  = doc.add_paragraph()
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'),   str(bookmark_id))
        bm_start.set(qn('w:name'), bookmark_name)
        p._p.append(bm_start)

        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), str(bookmark_id))
        p._p.append(bm_end)

    def _render_textbox(self, doc, content: str, block) -> None:
        """
        Rendert absolut positionierte Textbox.
        Laedt XML-Template aus template_dir/textbox_{layout_ref}.xml
        und ersetzt {variablen} mit Werten aus variables.
        """
        import os, copy
        from django.conf import settings
        from lxml import etree

        lay        = doc._lay
        tpl_dir    = getattr(doc, '_template_dir', '')
        layout_ref = block.layout_ref or 'kontakt'

        # XML-Template laden
        xml_path = os.path.join(
            settings.BASE_DIR,
            'apps', 'abpe_doc_studio',
            'generator', 'templates',
            tpl_dir,
            f'textbox_{layout_ref}.xml'
        )

        if not os.path.exists(xml_path):
            log.warning(f'Textbox-Template nicht gefunden: {xml_path}')
            return

        with open(xml_path, encoding='utf-8') as f:
            xml_str = f.read()

        # Variablen aus block.content ersetzen
        rendered = self.variable_engine.render_text(
            xml_str, self._last_vars or {}
        )

        # XML parsen und in Body einfügen
        try:
            elem = etree.fromstring(rendered.encode('utf-8'))
            doc.element.body.append(elem)
        except Exception as e:
            log.warning(f'Textbox XML Fehler: {e}')

    def _add_anchored_picture(self, paragraph, img_path: str,
                               height_cm: float, width_cm: float,
                               pos_h_cm: float, pos_v_cm: float) -> None:
        """
        Fuegt ein absolut positioniertes Bild ein (wp:anchor).
        Das Bild schwebt an fester Position unabhaengig vom Textfluss.
        Entspricht dem Original-DOCX Logo-Verhalten.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
        import os

        # EMU Konversion
        EMU = 914400 / 2.54
        cx  = int(width_cm  * EMU)
        cy  = int(height_cm * EMU)
        x   = int(pos_h_cm  * EMU)
        y   = int(pos_v_cm  * EMU)

        # Bild in Dokument einbetten
        doc = paragraph._p.getroottree().getroot()
        # python-docx Weg: inline pic erstellen dann zu anchor konvertieren
        run = paragraph.add_run()
        from docx.shared import Inches
        pic = run.add_picture(img_path, width=Cm(width_cm), height=Cm(height_cm))

        # Inline → Anchor konvertieren
        inline = run._r.find('.//' + qn('wp:inline'))
        if inline is None:
            return

        # Anchor Element erstellen
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251661312')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        # simplePos
        sp = OxmlElement('wp:simplePos')
        sp.set('x', '0'); sp.set('y', '0')
        anchor.append(sp)

        # Horizontale Position: absolut von Seite
        posH = OxmlElement('wp:positionH')
        posH.set('relativeFrom', 'page')
        posOffset_h = OxmlElement('wp:posOffset')
        posOffset_h.text = str(x)
        posH.append(posOffset_h)
        anchor.append(posH)

        # Vertikale Position: absolut von Seite
        posV = OxmlElement('wp:positionV')
        posV.set('relativeFrom', 'page')
        posOffset_v = OxmlElement('wp:posOffset')
        posOffset_v.text = str(y)
        posV.append(posOffset_v)
        anchor.append(posV)

        # Groesse
        extent = OxmlElement('wp:extent')
        extent.set('cx', str(cx))
        extent.set('cy', str(cy))
        anchor.append(extent)

        # effectExtent
        ee = OxmlElement('wp:effectExtent')
        ee.set('l','0'); ee.set('t','0'); ee.set('r','0'); ee.set('b','0')
        anchor.append(ee)

        # Textumbruch
        wrap = OxmlElement('wp:wrapSquare')
        wrap.set('wrapText', 'bothSides')
        anchor.append(wrap)

        # docPr
        docPr = OxmlElement('wp:docPr')
        docPr.set('id', '1')
        docPr.set('name', 'Logo')
        anchor.append(docPr)

        # Graphic aus inline kopieren
        graphic = inline.find('.//' + qn('a:graphic'))
        if graphic is not None:
            import copy
            anchor.append(copy.deepcopy(graphic))

        # cNvGraphicFramePr
        cNv = OxmlElement('wp:cNvGraphicFramePr')
        anchor.insert(list(anchor).index(anchor.find(qn('a:graphic'))), cNv)

        # Inline durch Anchor ersetzen
        drawing = inline.getparent()
        drawing.remove(inline)
        drawing.append(anchor)

    # ── Zell-Helfer ───────────────────────────────────────────────────────

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.lstrip('#'))
        tcPr.append(shd)

    def _set_cell_width(self, cell, width_dxa: int) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_cell_borders_none(self, cell) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        bdrs = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement('w:' + side)
            el.set(qn('w:val'), 'nil')
            bdrs.append(el)
        tcPr.append(bdrs)
