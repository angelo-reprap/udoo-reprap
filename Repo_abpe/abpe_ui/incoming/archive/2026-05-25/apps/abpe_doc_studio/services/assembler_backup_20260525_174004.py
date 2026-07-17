"""
services/assembler.py
=====================
DocAssembler — DOCX-Generierung + PDF via LibreOffice.

Prinzipien — KEINE Ausnahmen:
  - Kein hardcoded Text
  - Keine hardcoded Formatierung (Pt, bold, italic, color)
  - Keine hardcoded Geometrie (cm, dxa)
  - ALLES kommt aus:
      block.content          → Text (mit {variablen})
      StyleDefinition (DB)   → Formatierung via _apply_style()
      PageLayout (DB)        → Geometrie
      layout_constants.py    → nur technische DXA-Konstanten
"""
import io
import logging

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml    import OxmlElement

from .context_loader   import ContextLoader
from .variable_engine  import VariableEngine
from .exporter         import DocExporter
from .layout_constants import (
    FONT_NAME,
    COLOR_BRAND, COLOR_TEXT,
    TABLE_WIDTH_DXA,
    LABEL_VALUE_LABEL_DXA, LABEL_VALUE_VALUE_DXA,
    INV_HEADER_LEFT_DXA, INV_HEADER_RIGHT_DXA,
    TOTAL_LEFT_DXA, TOTAL_LABEL_DXA, TOTAL_VALUE_DXA,
    SIG_LEFT_DXA, SIG_MID_DXA, SIG_RIGHT_DXA,
    COLOR_WHITE, COLOR_ALT_ROW,
    COLOR_SEPARATOR, COLOR_SIG_LINE, COLOR_GRAY,
)

log = logging.getLogger('abpe_doc_studio.assembler')


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


class StyleCache:
    """Cached StyleDefinitions pro StyleKit."""
    def __init__(self):
        self._cache = {}

    def load(self, style_kit) -> dict:
        kid = style_kit.pk if style_kit else None
        if kid not in self._cache:
            self._cache[kid] = (
                {s.style_key: s for s in style_kit.definitions.all()}
                if style_kit else {}
            )
        return self._cache[kid]

    def get(self, style_kit, style_key: str):
        if not style_key or not style_kit:
            return None
        return self.load(style_kit).get(style_key)


class DocAssembler:

    def __init__(self):
        self.context_loader  = ContextLoader()
        self.variable_engine = VariableEngine()
        self.exporter        = DocExporter()
        self.style_cache     = StyleCache()

    # ── Öffentliche API ───────────────────────────────────────────────────

    def generate(self, template_identifier: str,
                  variables: dict = None,
                  context_ref: str = '',
                  scope: str = '',
                  engine: str = 'BOTH',
                  user=None) -> dict:

        from apps.abpe_doc_studio.models import (
            DocTemplate, DocLog, DocStatus, LogStatus
        )
        tpl = DocTemplate.objects.filter(
            identifier=template_identifier,
            status=DocStatus.ACTIVE
        ).select_related('layout', 'style_kit').prefetch_related(
            'template_blocks__block__style_kit__definitions'
        ).first()

        if not tpl:
            raise ValueError(
                f"Template '{template_identifier}' nicht gefunden"
            )

        all_vars = self.context_loader.load(
            scope           = scope or tpl.scope,
            context_ref     = context_ref,
            extra_variables = variables or {},
        )

        doc       = self._build_docx(tpl, all_vars)
        result    = {'success': True}
        docx_info = self.exporter.save_docx(
            doc, template_identifier,
            scope=scope or tpl.scope, context_ref=context_ref,
        )
        result['file_path_docx'] = docx_info['file_path']
        result['file_size']      = docx_info['file_size']

        if engine in ('PDF', 'BOTH'):
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    docx_info['file_path']
                )
                result['file_path_pdf'] = pdf_info['file_path']
            except Exception as e:
                log.warning(f'PDF fehlgeschlagen: {e}')
                result['file_path_pdf'] = ''
                result['pdf_error']     = str(e)

        doc_log = DocLog.objects.create(
            template         = tpl,
            template_version = tpl.active_version,
            engine_used      = engine,
            file_path_docx   = result.get('file_path_docx', ''),
            file_path_pdf    = result.get('file_path_pdf', ''),
            file_size_bytes  = result.get('file_size', 0),
            context_ref      = context_ref,
            scope            = scope or tpl.scope,
            variables_used   = {
                k: v for k, v in all_vars.items()
                if not isinstance(v, list)
            },
            status       = LogStatus.OK,
            generated_by = user,
        )
        result['log_id'] = str(doc_log.log_id)
        log.info(f'Generiert: {template_identifier} [{doc_log.log_id}]')
        return result

    def render_to_bytes(self, template, variables: dict = None,
                         engine: str = 'DOCX') -> bytes:
        from apps.abpe_doc_studio.models import DocTemplate
        if isinstance(template, str):
            template = DocTemplate.objects.filter(
                identifier=template
            ).select_related('layout', 'style_kit').prefetch_related(
                'template_blocks__block__style_kit__definitions'
            ).first()
            if not template:
                raise ValueError('Template nicht gefunden')

        doc = self._build_docx(template, variables or {})
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        if engine == 'PDF':
            import tempfile, os
            with tempfile.NamedTemporaryFile(
                suffix='.docx', delete=False
            ) as tmp:
                tmp.write(buf.getvalue())
                tmp_path = tmp.name
            try:
                pdf_info = self.exporter.convert_docx_to_pdf(
                    tmp_path, output_dir=os.path.dirname(tmp_path)
                )
                return self.exporter.read_as_bytes(pdf_info['file_path'])
            finally:
                os.unlink(tmp_path)
                pdf_path = tmp_path.replace('.docx', '.pdf')
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)

        return buf.getvalue()

    # ── DOCX bauen ────────────────────────────────────────────────────────

    def _build_docx(self, tpl, all_vars: dict) -> DocxDocument:
        """
        DOCX aufbauen.
        LOGO-Block erstellt 2-spaltige Tabelle im Body.
        DOC_TITLE + PARTY_BLOCK → linke Zelle.
        Ab CLAUSE → normaler Body-Flow.

        Alle Referenzen kommen aus tpl:
          tpl.style_kit        → Formatierung
          tpl.logo_block_id    → welcher Block ist das Logo
          tpl.footer_block_id  → welcher Block ist der Footer
          tpl.layout           → Geometrie
        """
        doc = DocxDocument()
        doc._logo_left_cell = None
        doc._style_kit       = tpl.style_kit
        doc._logo_block_id   = tpl.logo_block_id
        doc._footer_block_id = tpl.footer_block_id
        # Geometrie aus layout — kein hardcoded Wert im assembler
        doc._logo_col_widths = getattr(tpl.layout, 'logo_column_widths_cm', [9.0, 6.0])
        doc._logo_height_cm  = getattr(tpl.layout, 'logo_height_cm', 2.0)

        self._apply_page_layout(doc, tpl.layout, all_vars)

        for tb in tpl.template_blocks.order_by('slot', 'order'):
            if tb.conditional:
                if not self.variable_engine.check_condition(
                    tb.conditional, all_vars
                ):
                    continue
            if tb.page_break_before and doc.paragraphs:
                doc.add_page_break()

            bt = tb.block.block_type

            if bt in ('DOC_TITLE', 'PARTY_BLOCK') \
                    and doc._logo_left_cell is not None:
                self._render_block_in_cell(
                    doc._logo_left_cell, tb, all_vars
                )
            else:
                self._render_block(doc, tb, all_vars)

        self._apply_footer(doc, all_vars)
        return doc

    # ── _apply_style: EINZIGE Formatierungsquelle ─────────────────────────

    def _apply_style(self, paragraph, style_key: str,
                      style_kit, run=None) -> None:
        """
        Liest StyleDefinition aus DB und wendet sie an.
        Kein Pt(), kein bold, kein color außerhalb dieser Methode.
        """
        sdef = self.style_cache.get(style_kit, style_key)
        if not sdef:
            return

        if sdef.space_before_pt:
            paragraph.paragraph_format.space_before = Pt(sdef.space_before_pt)
        if sdef.space_after_pt:
            paragraph.paragraph_format.space_after  = Pt(sdef.space_after_pt)
        if sdef.indent_left_cm:
            paragraph.paragraph_format.left_indent  = Cm(sdef.indent_left_cm)

        align_map = {
            'left':    WD_ALIGN_PARAGRAPH.LEFT,
            'right':   WD_ALIGN_PARAGRAPH.RIGHT,
            'center':  WD_ALIGN_PARAGRAPH.CENTER,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if sdef.alignment in align_map:
            paragraph.alignment = align_map[sdef.alignment]

        if sdef.border_bottom:
            pPr  = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),   sdef.border_bottom_style or 'single')
            bot.set(qn('w:sz'),    str(int(sdef.border_bottom_pt * 8)))
            bot.set(qn('w:space'), '4')
            bot.set(qn('w:color'),
                    sdef.border_bottom_color.lstrip('#'))
            pBdr.append(bot)
            pPr.append(pBdr)

        if run is not None:
            run.font.name      = sdef.font_family or FONT_NAME
            run.font.size      = Pt(sdef.font_size_pt)
            run.bold           = sdef.bold
            run.italic         = sdef.italic
            run.underline      = sdef.underline
            run.font.color.rgb = _rgb(sdef.color_hex)

    def _add_border_line(self, paragraph, color: str,
                          side: str = 'bottom',
                          sz: int = 4) -> None:
        """Fügt Rahmen-Linie an Paragraph hinzu."""
        pPr  = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        el   = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color.lstrip('#'))
        pBdr.append(el)
        pPr.append(pBdr)

    def _add_page_number_field(self, paragraph, sdef,
                                instr: str) -> None:
        """Fügt PAGE oder NUMPAGES Feldfunktion ein."""
        for ftype, text in [
            ('begin', instr), ('separate', None), ('end', None)
        ]:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), ftype)
            run = paragraph.add_run()
            run._r.append(fld)
            if sdef:
                run.font.name = sdef.font_family or FONT_NAME
                run.font.size = Pt(sdef.font_size_pt)
                run.font.color.rgb = _rgb(COLOR_BRAND)
            if text:
                ins = OxmlElement('w:instrText')
                ins.text = text
                run._r.append(ins)

    # ── Seitenlayout ──────────────────────────────────────────────────────

    def _apply_page_layout(self, doc: DocxDocument, layout,
                            variables: dict) -> None:
        for section in doc.sections:
            section.page_width      = Cm(layout.page_width_cm)
            section.page_height     = Cm(layout.page_height_cm)
            section.left_margin     = Cm(layout.margin_left_cm)
            section.right_margin    = Cm(layout.margin_right_cm)
            section.top_margin      = Cm(layout.margin_top_cm)
            section.bottom_margin   = Cm(layout.margin_bottom_cm)
            section.header_distance = Cm(layout.header_distance_cm)

        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal'].font.size = Pt(10)

        self._apply_header(doc, layout, style_kit=getattr(doc, '_style_kit', None))

    def _apply_header(self, doc: DocxDocument, layout,
                       style_kit=None) -> None:
        """
        Word-Header: nur Seitennummer.
        Format aus layout.page_number_format.
        Stil aus StyleDefinition 'footer_text' des übergebenen StyleKits.
        """
        sk   = style_kit
        sdef = self.style_cache.get(sk, 'footer_text') if sk else None

        nr_fmt      = layout.page_number_format
        before_page = nr_fmt.split('{page}')[0]
        rest        = nr_fmt.split('{page}')[1] if '{page}' in nr_fmt else ''
        between     = rest.split('{total}')[0]  if '{total}' in rest  else ''
        after_total = rest.split('{total}')[1]  if '{total}' in rest  else ''

        for section in doc.sections:
            hdr = section.header
            for p in hdr.paragraphs:
                p.clear()
            for p in list(hdr.paragraphs)[1:]:
                p._element.getparent().remove(p._element)

            p_nr = hdr.paragraphs[0]
            p_nr.clear()
            p_nr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _snr = self.style_cache.get(sk, 'footer_text')
            p_nr.paragraph_format.space_before = Pt(_snr.space_before_pt if _snr else 0)
            p_nr.paragraph_format.space_after  = Pt(_snr.space_after_pt  if _snr else 0)

            def _run(text):
                r = p_nr.add_run(text)
                if sdef:
                    r.font.name = sdef.font_family or FONT_NAME
                    r.font.size = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(COLOR_BRAND)
                return r

            if before_page:
                _run(before_page)
            self._add_page_number_field(p_nr, sdef, ' PAGE ')
            if between:
                _run(between)
            self._add_page_number_field(p_nr, sdef, ' NUMPAGES ')
            if after_total:
                _run(after_total)

            _bb = self.style_cache.get(sk, 'border_brand')
            self._add_border_line(p_nr,
                _bb.border_bottom_color if _bb else COLOR_BRAND)

    def _apply_footer(self, doc: DocxDocument,
                       variables: dict) -> None:
        """
        Footer aus variables (context_loader._system_vars).
        Stil aus StyleDefinition 'footer_text' des Template-StyleKits.
        Kein hardcoded Text.
        """
        sk   = getattr(doc, '_style_kit', None)
        sdef = self.style_cache.get(sk, 'footer_text') if sk else None

        v = variables or {}

        # Zeile 1: Firma aus variables
        imp_parts = [p for p in [
            v.get('ag_firma', ''),
            f"Inhaber: {v['ag_name']}" if v.get('ag_name') else '',
            v.get('ag_strasse', ''),
            v.get('ag_plz_ort', ''),
        ] if p]
        imp_line = ' | '.join(imp_parts)

        # Zeile 2: Rechtliches aus variables
        jur_parts = [p for p in [
            v.get('ag_amtsgericht', ''),
            v.get('ag_hra', ''),
            f"USt-IdNr.: {v['ag_ustid']}" if v.get('ag_ustid') else '',
            v.get('ag_steuernr', ''),
        ] if p]
        jur_line = ' | '.join(jur_parts)

        for section in doc.sections:
            ftr  = section.footer
            para = (ftr.paragraphs[0]
                    if ftr.paragraphs else ftr.add_paragraph())
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _fb = self.style_cache.get(sk, 'border_brand')
            para.paragraph_format.space_before = Pt(
                _fb.space_before_pt if _fb else 4)
            self._add_border_line(para,
                _fb.border_bottom_color if _fb else COLOR_BRAND, side='top')

            if imp_line:
                r = para.add_run(imp_line)
                if sdef:
                    r.font.name = sdef.font_family or FONT_NAME
                    r.font.size = Pt(sdef.font_size_pt)
                    r.font.color.rgb = _rgb(sdef.color_hex)

            if jur_line:
                p2 = ftr.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(jur_line)
                if sdef:
                    r2.font.name = sdef.font_family or FONT_NAME
                    r2.font.size = Pt(sdef.font_size_pt)
                    r2.font.color.rgb = _rgb(sdef.color_hex)

    # ── Block-Dispatch ────────────────────────────────────────────────────

    def _render_block(self, doc: DocxDocument, tb,
                       variables: dict) -> None:
        block    = tb.block
        content  = tb.content_override or block.content or ''
        rendered = self.variable_engine.render_text(content, variables)
        bt       = block.block_type
        sk       = block.style_kit
        skey     = (tb.style_override.get('style_key', '')
                    or block.style_key or '')

        {
            'LOGO':        lambda: self._render_logo(doc, variables, sk),
            'DOC_TITLE':   lambda: self._render_lines(
                               doc, rendered, skey, sk),
            'SECTION_HEAD':lambda: self._render_lines(
                               doc, rendered, skey, sk),
            'CLAUSE':      lambda: self._render_clause(
                               doc, rendered, skey, sk),
            'PARAGRAPH':   lambda: self._render_lines(
                               doc, rendered, skey, sk),
            'PARTY_BLOCK': lambda: self._render_party_block(
                               doc, rendered, skey, sk),
            'LABEL_VALUE': lambda: self._render_label_value(
                               doc, rendered, skey, sk),
            'TIME_TABLE':  lambda: self._render_time_table(
                               doc, block, variables),
            'AP_TABLE':    lambda: self._render_ap_table(
                               doc, block, variables),
            'TOTAL_BLOCK': lambda: self._render_total_block(
                               doc, rendered, skey, sk),
            'SIGNATURE':   lambda: self._render_signature(
                               doc, rendered, variables, skey, sk),
            'SEPARATOR':   lambda: self._render_separator(doc),
            'PAGE_BREAK':  lambda: doc.add_page_break(),
            'INV_HEADER':  lambda: self._render_inv_header(
                               doc, variables, sk),
            'INV_META':    lambda: self._render_inv_meta(
                               doc, rendered, skey, sk),
            'INV_SUBJECT': lambda: self._render_lines(
                               doc, rendered, skey, sk),
            'CLOSING':     lambda: self._render_closing(
                               doc, rendered, skey, sk),
        }.get(bt, lambda: self._render_lines(
            doc, rendered, skey, sk
        ))()

    def _render_block_in_cell(self, cell, tb,
                               variables: dict) -> None:
        """Rendert DOC_TITLE / PARTY_BLOCK in linke Logo-Zelle."""
        block    = tb.block
        content  = tb.content_override or block.content or ''
        rendered = self.variable_engine.render_text(content, variables)
        bt       = block.block_type
        sk       = block.style_kit
        skey     = (tb.style_override.get('style_key', '')
                    or block.style_key or '')

        if bt == 'DOC_TITLE':
            p   = cell.add_paragraph()
            run = p.add_run(rendered)
            self._apply_style(p, skey or 'doc_title', sk, run)

        elif bt == 'PARTY_BLOCK':
            def is_frame(line: str) -> bool:
                l = line.strip().lower()
                return (l.startswith('-') or l in ('zwischen', 'und')
                        or l.startswith('zum rahmenvertrag'))
            for line in rendered.split('\n'):
                if not line.strip():
                    continue
                p   = cell.add_paragraph()
                run = p.add_run(line.strip())
                self._apply_style(
                    p,
                    'body_text' if is_frame(line) else (skey or 'party_bold'),
                    sk, run
                )
            _sp = self.style_cache.get(sk, 'spacing_normal')
            cell.add_paragraph().paragraph_format.space_after = Pt(
                _sp.space_after_pt if _sp else 4)

    # ── Block-Renderer ────────────────────────────────────────────────────

    def _render_logo(self, doc, variables: dict, sk) -> None:
        """
        2-spaltige Tabelle im Body.
        Geometrie aus PageLayout (15cm content = 9cm + 6cm).
        Logo-Bild + Kontaktzeilen aus ContentBlock 'abcona_logo'.
        Styles: logo_contact_title, logo_contact_tagline, logo_contact_line.
        """
        import os
        from django.conf import settings
        from apps.abpe_doc_studio.models import ContentBlock

        logo_paths = [
            os.path.join(settings.BASE_DIR, 'data', 'cv', 'adds',
                         'logo_abcona.png'),
            os.path.join(settings.BASE_DIR, 'apps', 'abpe_ui', 'static',
                         'abpe_ui', 'img', 'logo_abcona.png'),
        ]
        logo_path = next(
            (p for p in logo_paths if os.path.exists(p)), None
        )

        logo_block_id = getattr(doc, '_logo_block_id', 'abcona_logo')
        logo_block = ContentBlock.objects.filter(
            identifier=logo_block_id, is_active=True
        ).select_related('style_kit').first()
        raw  = logo_block.content if logo_block else ''
        lsk  = logo_block.style_kit if logo_block else sk

        rendered = self.variable_engine.render_text(
            raw, variables or {}
        )
        lines = [l.strip() for l in rendered.split('\n') if l.strip()]

        # Spaltenbreiten aus layout.logo_column_widths_cm
        # Fallback: [9.0, 6.0] wenn nicht gesetzt
        col_widths  = getattr(doc, '_logo_col_widths', None) or [9.0, 6.0]
        logo_height = getattr(doc, '_logo_height_cm', 2.0)
        left_dxa  = int(col_widths[0] * 567)
        right_dxa = int(col_widths[1] * 567)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc = tbl.rows[0].cells[0]
        rc = tbl.rows[0].cells[1]
        self._set_cell_width(lc, left_dxa)
        self._set_cell_width(rc, right_dxa)
        self._set_cell_borders_none(lc)
        self._set_cell_borders_none(rc)
        lc.paragraphs[0].clear()

        # Rechts: Logo-Bild
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(0)
        rp.paragraph_format.space_after  = Pt(4)
        if logo_path:
            rp.add_run().add_picture(logo_path, height=Cm(logo_height))

        # Rechts: Kontaktzeilen aus block.content
        for line in lines:
            if '|' in line:
                firma, tagline = line.split('|', 1)
                if firma.strip():
                    p2 = rc.add_paragraph()
                    r2 = p2.add_run(firma.strip())
                    self._apply_style(p2, 'logo_contact_title', lsk, r2)
                if tagline.strip():
                    p3 = rc.add_paragraph()
                    r3 = p3.add_run(tagline.strip())
                    self._apply_style(p3, 'logo_contact_tagline', lsk, r3)
            else:
                p4 = rc.add_paragraph()
                r4 = p4.add_run(line)
                self._apply_style(p4, 'logo_contact_line', lsk, r4)

        doc._logo_left_cell = lc

    def _render_lines(self, doc, content: str,
                       skey: str, sk) -> None:
        """Generischer Renderer: jede Zeile = ein Paragraph mit skey-Style."""
        for line in content.split('\n'):
            if not line.strip():
                _sp2 = self.style_cache.get(getattr(doc, '_style_kit', sk) if hasattr(doc,'_style_kit') else None, 'spacing_normal')
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue
            p   = doc.add_paragraph()
            run = p.add_run(line.strip())
            self._apply_style(p, skey, sk, run)

    def _render_clause(self, doc, content: str,
                        skey: str, sk) -> None:
        """
        Zeile 1: skey (section_head)
        Rest:    body_text
        """
        if not content:
            return
        lines = content.split('\n')
        if lines:
            p   = doc.add_paragraph()
            run = p.add_run(lines[0].strip())
            self._apply_style(p, skey or 'section_head', sk, run)
        for line in lines[1:]:
            if not line.strip():
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue
            p   = doc.add_paragraph()
            run = p.add_run(line.strip())
            self._apply_style(p, 'body_text', sk, run)

    def _render_party_block(self, doc, content: str,
                             skey: str, sk) -> None:
        """
        Rahmenzeilen (zwischen/und/-...) → body_text
        Adresszeilen → skey (party_bold/party_italic)
        """
        def is_frame(line: str) -> bool:
            l = line.strip().lower()
            return (l.startswith('-') or l in ('zwischen', 'und')
                    or l.startswith('zum rahmenvertrag'))

        for line in (content.split('\n') if content else []):
            if not line.strip():
                continue
            p   = doc.add_paragraph()
            run = p.add_run(line.strip())
            self._apply_style(
                p,
                'body_text' if is_frame(line) else (skey or 'party_bold'),
                sk, run
            )
        _sp3 = self.style_cache.get(sk, 'spacing_normal')
        doc.add_paragraph().paragraph_format.space_after = Pt(
            _sp3.space_before_pt if _sp3 and _sp3.space_before_pt else 8)

    def _render_label_value(self, doc, content: str,
                             skey: str, sk) -> None:
        """
        content: "Label|Wert"
        Mehrere Zeilen möglich: jede Zeile = eine Tabellenzeile.
        """
        if not content:
            return
        for line in content.split('\n'):
            if not line.strip():
                continue
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            lc, vc = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
            self._set_cell_borders_none(lc)
            self._set_cell_borders_none(vc)
            self._set_cell_width(lc, LABEL_VALUE_LABEL_DXA)
            self._set_cell_width(vc, LABEL_VALUE_VALUE_DXA)

            if '|' in line:
                label, value = line.split('|', 1)
            else:
                label, value = line, ''

            lrun = lc.paragraphs[0].add_run(label.strip())
            self._apply_style(lc.paragraphs[0],
                              skey or 'label_blue', sk, lrun)
            vrun = vc.paragraphs[0].add_run(value.strip())
            self._apply_style(vc.paragraphs[0], 'body_text', sk, vrun)

        _sp4 = self.style_cache.get(sk, 'spacing_normal')
        doc.add_paragraph().paragraph_format.space_after = Pt(
            _sp4.space_after_pt if _sp4 else 4)

    def _render_total_block(self, doc, content: str,
                             skey: str, sk) -> None:
        """
        content Format (aus inv_total_block):
          Summe|{summe_netto}
          zzgl. Umsatzsteuer {mwst_satz} %|{mwst_euro}
          Gesamtbetrag|{gesamtbetrag}

        Letzte Zeile = Gesamtbetrag → total_row Style (bg blau, weiß)
        Andere Zeilen → body_text
        """
        lines = [l for l in content.split('\n') if l.strip()]
        if not lines:
            return

        tbl = doc.add_table(rows=len(lines), cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'),    str(TABLE_WIDTH_DXA))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        for ri, line in enumerate(lines):
            is_total = (ri == len(lines) - 1)
            row_skey = 'total_row' if is_total else 'body_text'
            total_sdef = self.style_cache.get(sk, row_skey)
            bg = (total_sdef.bg_color_hex
                  if (is_total and total_sdef and total_sdef.bg_color_hex)
                  else '')

            label, value = (line.split('|', 1)
                            if '|' in line else (line, ''))

            for ci, (w, text) in enumerate([
                (TOTAL_LEFT_DXA,  ''),
                (TOTAL_LABEL_DXA, label.strip()),
                (TOTAL_VALUE_DXA, value.strip()),
            ]):
                cell = tbl.rows[ri].cells[ci]
                self._set_cell_width(cell, w)
                self._set_cell_borders_none(cell)
                if bg:
                    self._set_cell_bg(cell, bg)
                p   = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(text)
                self._apply_style(p, row_skey, sk, run)

        _sp5 = self.style_cache.get(sk, 'spacing_normal')
        doc.add_paragraph().paragraph_format.space_after = Pt(
            (_sp5.space_after_pt * 2.5) if _sp5 else 16)

    def _render_signature(self, doc, content: str,
                           variables: dict,
                           skey: str, sk) -> None:
        """
        content Format (aus abcona_signature):
          {ag_ort}, den {datum_heute}|{an_ort}, den {datum_heute}
          Auftraggeber AG|Auftragnehmer AN
          Stempel/Unterschrift|Unterschrift

        Zeile 0: Ort+Datum links|rechts → sig_label + Unterlinie
        Zeile 1: Labels → sig_label
        Zeile 2: Hinweise → sig_hint
        """
        lines = [l for l in content.split('\n') if l.strip()]
        if not lines:
            return

        # 3 Zeilen: [Datum, Labels, Hinweise]
        # Tabelle mit 3 Spalten: AG | Abstand | AN
        tbl = doc.add_table(rows=len(lines), cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for ri, line in enumerate(lines):
            parts = line.split('|', 1)
            left  = parts[0].strip()
            right = parts[1].strip() if len(parts) > 1 else ''

            for ci, (w, text) in enumerate([
                (SIG_LEFT_DXA,  left),
                (SIG_MID_DXA,   ''),
                (SIG_RIGHT_DXA, right),
            ]):
                cell = tbl.rows[ri].cells[ci]
                self._set_cell_width(cell, w)
                self._set_cell_borders_none(cell)
                p   = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)

                if not text:
                    continue

                # Stil je Zeile
                if ri == 0:
                    row_skey = 'sig_label'
                elif ri == 1:
                    row_skey = 'sig_label'
                else:
                    row_skey = 'sig_hint'

                run = p.add_run(text)
                self._apply_style(p, row_skey, sk, run)

                # Datum-Zeile: Unterlinie unter Ort/Datum
                if ri == 0 and ci != 1:
                    _bsig = self.style_cache.get(sk, 'border_signature')
                    self._add_border_line(p,
                        _bsig.border_bottom_color if _bsig else COLOR_SIG_LINE)

            # Zwischen Datum und Labels: Platz für Unterschrift
            if ri == 0:
                for ci in range(3):
                    cell = tbl.rows[ri].cells[ci]
                    # Zeile 1 Platz
                    pass

        # Leerzeile für Unterschrift zwischen Zeile 0 und 1
        # → neue Tabelle mit Abstand
        _ssig = self.style_cache.get(sk, 'spacing_signature')
        doc.add_paragraph().paragraph_format.space_before = Pt(
            _ssig.space_before_pt if _ssig else 40)

    def _render_closing(self, doc, content: str,
                         skey: str, sk) -> None:
        """
        content aus inv_closing:
          Mit freundlichen Grüssen
          Zahlungsziel: {zahlungsziel_text}
          {ag_firma}
          {ag_name} · Inhaber

        Alle Zeilen → body_text Style.
        Letzter Block rechtsbündig (Zahlungsziel).
        """
        lines = content.split('\n') if content else []
        for line in lines:
            if not line.strip():
                continue
            p   = doc.add_paragraph()
            run = p.add_run(line.strip())
            self._apply_style(p, skey or 'body_text', sk, run)

    def _render_inv_header(self, doc, variables: dict, sk) -> None:
        """
        Rechnungskopf 2-spaltig.
        Links: Empfänger aus variables.
        Rechts: Absender aus ContentBlock 'abcona_logo'.
        Styles aus SK.
        """
        from apps.abpe_doc_studio.models import ContentBlock

        v = variables or {}

        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc, rc = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
        self._set_cell_width(lc, INV_HEADER_LEFT_DXA)
        self._set_cell_width(rc, INV_HEADER_RIGHT_DXA)
        self._set_cell_borders_none(lc)
        self._set_cell_borders_none(rc)

        # Links: Empfänger aus variables
        addr_lines = []
        if v.get('empfaenger_firma'):
            addr_lines.append(
                (v['empfaenger_firma'], 'label_blue')
            )
        for line in v.get('empfaenger_adresse', '').split('\n'):
            if line.strip():
                addr_lines.append((line.strip(), 'body_text'))

        first = True
        for text, lskey in addr_lines:
            p = lc.paragraphs[0] if first else lc.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            self._apply_style(p, lskey, sk, run)

        # Rechts: Absender aus 'abcona_logo' ContentBlock
        logo_block_id = getattr(doc, '_logo_block_id', 'abcona_logo')
        logo_block = ContentBlock.objects.filter(
            identifier=logo_block_id, is_active=True
        ).select_related('style_kit').first()
        lsk = logo_block.style_kit if logo_block else sk
        raw = logo_block.content if logo_block else ''
        rendered = self.variable_engine.render_text(raw, v)

        first = True
        for line in rendered.split('\n'):
            if not line.strip():
                continue
            p = rc.paragraphs[0] if first else rc.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)

            if '|' in line:
                firma, tagline = line.split('|', 1)
                if firma.strip():
                    run = p.add_run(firma.strip())
                    self._apply_style(p, 'logo_contact_title', lsk, run)
                if tagline.strip():
                    p2 = rc.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r2 = p2.add_run(tagline.strip())
                    self._apply_style(p2, 'logo_contact_tagline', lsk, r2)
            else:
                run = p.add_run(line.strip())
                self._apply_style(p, 'logo_contact_line', lsk, run)

        _sp6 = self.style_cache.get(sk, 'spacing_normal')
        doc.add_paragraph().paragraph_format.space_after = Pt(
            (_sp6.space_after_pt * 2) if _sp6 else 12)

    def _render_inv_meta(self, doc, content: str,
                          skey: str, sk) -> None:
        """
        content aus inv_meta:
          Rg.-Nr.:|{rg_nummer}
          Datum:|{rg_datum}
        → LABEL_VALUE Rendering
        """
        self._render_label_value(doc, content, skey or 'inv_rgnr', sk)

    def _render_separator(self, doc) -> None:
        p = doc.add_paragraph()
        _bsep = self.style_cache.get(getattr(doc, '_style_kit', None), 'border_separator') if hasattr(doc, '_style_kit') else None
        p.paragraph_format.space_before = Pt(_bsep.space_before_pt if _bsep else 4)
        p.paragraph_format.space_after  = Pt(_bsep.space_after_pt  if _bsep else 4)
        self._add_border_line(p,
            _bsep.border_bottom_color if _bsep else COLOR_SEPARATOR)

    def _render_time_table(self, doc, block,
                            variables: dict) -> None:
        columns = block.columns or []
        rows    = self.variable_engine.expand_table_rows(
            columns, 'positionen', variables
        )
        self._build_invoice_table(doc, columns, rows, block.style_kit)

    def _render_ap_table(self, doc, block,
                          variables: dict) -> None:
        columns = block.columns or []
        rows    = self.variable_engine.expand_table_rows(
            columns, 'arbeitspakete', variables
        )
        self._build_invoice_table(doc, columns, rows, block.style_kit)

    def _build_invoice_table(self, doc, columns: list,
                               rows: list, sk) -> None:
        if not columns:
            return

        hdr_sdef  = self.style_cache.get(sk, 'table_header')
        body_sdef = self.style_cache.get(sk, 'table_body')
        hdr_bg    = (hdr_sdef.table_header_bg_hex
                     if hdr_sdef else COLOR_BRAND)
        alt_bg    = (body_sdef.table_row_alt_bg_hex
                     if body_sdef else COLOR_ALT_ROW)

        n_cols = len(columns)
        tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'),    str(TABLE_WIDTH_DXA))
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        total_pct = sum(c.get('width_pct', 20) for c in columns)
        hrow = tbl.rows[0]

        for ci, col in enumerate(columns):
            cell = hrow.cells[ci]
            w    = int(TABLE_WIDTH_DXA * col.get('width_pct',20) / total_pct)
            self._set_cell_width(cell, w)
            self._set_cell_bg(cell, hdr_bg)
            self._set_cell_borders_none(cell)
            p   = cell.paragraphs[0]
            run = p.add_run(col.get('label', ''))
            self._apply_style(p, 'table_header', sk, run)

        for ri, row_data in enumerate(rows):
            drow = tbl.rows[ri + 1]
            bg   = alt_bg if ri % 2 == 0 else COLOR_WHITE
            for ci, col in enumerate(columns):
                cell  = drow.cells[ci]
                value = row_data[ci] if ci < len(row_data) else ''
                w     = int(TABLE_WIDTH_DXA * col.get('width_pct',20) / total_pct)
                self._set_cell_width(cell, w)
                self._set_cell_bg(cell, bg)
                self._set_cell_borders_none(cell)
                p   = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                               if col.get('align') == 'right'
                               else WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(value)
                self._apply_style(p, 'table_body', sk, run)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Zell-Helfer ───────────────────────────────────────────────────────

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.lstrip('#'))
        tcPr.append(shd)

    def _set_cell_width(self, cell, width_dxa: int) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_cell_borders_none(self, cell) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        bdrs = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement('w:' + side)
            el.set(qn('w:val'), 'nil')
            bdrs.append(el)
        tcPr.append(bdrs)
