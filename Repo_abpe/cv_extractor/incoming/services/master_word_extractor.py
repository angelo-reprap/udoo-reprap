"""
services/master_word_extractor.py
==================================
Universeller DOCX/DOC-Extractor für die FL-Dokument-Classifier-Pipeline.

Unterstützte Formate:
  .docx  → direkt via python-docx (im RAM)
  .doc   → LibreOffice (exakt wie tasks.py, getestet) → tmp PDF
            → pdf_extractor → Spans im RAM
            Temp-PDF wird nach Extraktion sofort gelöscht.

RAM-Struktur (WordExtractResult):
  .ok            bool           — Extraktion erfolgreich
  .error         str|None       — Fehlermeldung wenn nicht ok
  .filename      str            — Dateiname ohne Pfad
  .source_format str            — 'docx' oder 'doc'
  .page_count    int            — Anzahl Seiten
  .span_count    int            — Anzahl Spans
  .char_count    int            — Gesamtzeichenanzahl (ohne Whitespace)
  .plain_text    str            — Volltext (für Regex/Keyword-Scan)
  .first_500     str            — Erste 500 Zeichen (für schnelle Klassifikation)
  .spans         List[dict]     — Span-Liste kompatibel mit pdf_extractor.py
  .tables        List[dict]     — Tabellen als {headers, rows}
  .headings      List[str]      — Nur Überschriften-Texte (bold+groß)
  .meta          dict           — Rohmetadaten (für Debugging)

Span-Format (kompatibel mit pdf_extractor.py Output):
  {
    'page':   int,
    'y':      int,
    'x':      int,
    'size':   float,
    'bold':   bool,
    'italic': bool,
    'font':   str,
    'text':   str,
    'width':  float,
  }
"""

import logging
import os
import subprocess
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Any

logger = logging.getLogger(__name__)

# ── Konstanten ────────────────────────────────────────────────────────────────

DEFAULT_SIZE         = 12.0
HEADING_SIZE_MIN     = 13.0
Y_STEP               = 20
Y_PAGE_OFFSET        = 1000
MAX_PLAIN_TEXT_CHARS = 50_000


# ── Ergebnis-Dataclass ────────────────────────────────────────────────────────

@dataclass
class WordExtractResult:
    ok:            bool           = False
    error:         Optional[str]  = None
    filename:      str            = ''
    source_format: str            = ''
    page_count:    int            = 0
    span_count:    int            = 0
    char_count:    int            = 0
    plain_text:    str            = ''
    first_500:     str            = ''
    spans:         List[dict]     = field(default_factory=list)
    tables:        List[dict]     = field(default_factory=list)
    headings:      List[str]      = field(default_factory=list)
    meta:          Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        if self.plain_text and not self.first_500:
            self.first_500 = self.plain_text[:500]


# ── Style-Hilfsfunktionen (DOCX) ─────────────────────────────────────────────

def _resolve_size(style) -> float:
    try:
        if style and style.font and style.font.size:
            return round(style.font.size.pt, 1)
        if style and style.base_style:
            return _resolve_size(style.base_style)
    except Exception:
        pass
    return DEFAULT_SIZE


def _resolve_bold(style) -> bool:
    try:
        if style and style.font and style.font.bold is not None:
            return bool(style.font.bold)
        if style and style.base_style:
            return _resolve_bold(style.base_style)
    except Exception:
        pass
    return False


def _resolve_font(style) -> str:
    try:
        if style and style.font and style.font.name:
            return style.font.name
        if style and style.base_style:
            return _resolve_font(style.base_style)
    except Exception:
        pass
    return 'Arial'


def _is_heading_style(style) -> bool:
    if not style:
        return False
    name = (style.name or '').lower()
    return (name.startswith('heading') or
            name.startswith('ueberschrift') or
            name.startswith('überschrift') or
            name.startswith('title'))


def _has_page_break_from_element(para_element) -> bool:
    try:
        xml = para_element.xml if hasattr(para_element, 'xml') else ''
        if 'w:type="page"' in xml:
            return True
    except Exception:
        pass
    return False


def _get_indent_x(para) -> int:
    try:
        indent = para.paragraph_format.left_indent
        if indent:
            return round(indent.pt)
    except Exception:
        pass
    raw  = para.text
    tabs = len(raw) - len(raw.lstrip('\t'))
    return tabs * 20 if tabs else 0


# ── Haupt-Extraktionsklasse ───────────────────────────────────────────────────

class MasterWordExtractor:

    def extract(self, doc_path: str) -> WordExtractResult:
        path_obj = Path(doc_path)

        if not path_obj.exists():
            return WordExtractResult(
                ok=False,
                error=f'Datei nicht gefunden: {doc_path}',
                filename=path_obj.name,
            )

        suffix = path_obj.suffix.lower()

        if suffix == '.docx':
            return self._extract_docx(path_obj)
        elif suffix == '.doc':
            return self._extract_doc_via_libreoffice(path_obj)
        else:
            return WordExtractResult(
                ok=False,
                error=f'Format nicht unterstützt: {suffix} — nur .docx und .doc',
                filename=path_obj.name,
            )

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def _extract_docx(self, path_obj: Path) -> WordExtractResult:
        start = time.time()
        try:
            from docx import Document
        except ImportError:
            return WordExtractResult(
                ok=False,
                error='python-docx nicht installiert',
                filename=path_obj.name,
            )
        try:
            doc = Document(str(path_obj))
        except Exception as e:
            return WordExtractResult(
                ok=False,
                error=f'DOCX konnte nicht geöffnet werden: {e}',
                filename=path_obj.name,
            )

        spans:    List[dict] = []
        tables:   List[dict] = []
        headings: List[str]  = []
        page = 1
        y    = 0

        for block in doc.element.body:
            tag = block.tag.split('}')[-1] if '}' in block.tag else block.tag

            if tag == 'p':
                if _has_page_break_from_element(block):
                    page += 1
                    y = page * Y_PAGE_OFFSET
                y += Y_STEP
                span_dict, is_heading = self._para_to_span(block, doc, page, y)
                if span_dict:
                    spans.append(span_dict)
                    if is_heading:
                        headings.append(span_dict['text'])

            elif tag == 'tbl':
                table_data = self._table_to_dict(block, doc, page, y)
                if table_data:
                    tables.append(table_data)
                    for row_spans in table_data.get('row_spans', []):
                        for s in row_spans:
                            s['y']    = y
                            s['page'] = page
                            spans.append(s)
                            y += Y_STEP

        return self._build_result(
            spans, tables, headings, page,
            path_obj.name, 'docx',
            round(time.time() - start, 3),
        )

    # ── DOC via LibreOffice ───────────────────────────────────────────────────
    # Befehl identisch zu tasks.py (getestet, funktioniert):
    #   libreoffice --headless --nofirststartwizard --norestore
    #               --convert-to pdf <doc_dest> --outdir <pdf_dir>

    def _extract_doc_via_libreoffice(self, path_obj: Path) -> WordExtractResult:
        start = time.time()

        with tempfile.TemporaryDirectory(prefix='abpe_doc_') as tmp_dir:
            tmp_path = Path(tmp_dir)

            # Exakt gleicher Befehl wie in tasks.py
            lo_result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    '--nofirststartwizard',
                    '--norestore',
                    '--convert-to', 'pdf',
                    str(path_obj),          # Datei VOR --outdir (wie in tasks.py)
                    '--outdir', str(tmp_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,             # gleicher Timeout wie tasks.py
            )

            if lo_result.returncode != 0:
                return WordExtractResult(
                    ok=False,
                    error=f'LibreOffice Exit {lo_result.returncode}: {lo_result.stderr[:300]}',
                    filename=path_obj.name,
                    source_format='doc',
                )

            # PDF suchen — gleiche Logik wie tasks.py
            pdf_name = path_obj.stem + '.pdf'
            pdf_path = tmp_path / pdf_name
            if not pdf_path.exists():
                pdfs = list(tmp_path.glob('*.pdf'))
                if not pdfs:
                    return WordExtractResult(
                        ok=False,
                        error='LibreOffice: kein PDF erzeugt',
                        filename=path_obj.name,
                        source_format='doc',
                    )
                pdf_path = pdfs[0]

            lo_duration = round(time.time() - start, 2)
            logger.info(
                f"[MasterWordExtractor] DOC→PDF: {path_obj.name} "
                f"→ {pdf_path.name} ({pdf_path.stat().st_size // 1024} KB) "
                f"in {lo_duration}s"
            )

            # PDF → Spans via pdf_extractor (im RAM)
            try:
                from apps.cv_extractor.services.pdf_extractor import PDFExtractor
                pdf_result = PDFExtractor().extract(str(pdf_path))
            except Exception as e:
                return WordExtractResult(
                    ok=False,
                    error=f'pdf_extractor nach DOC-Konvertierung: {e}',
                    filename=path_obj.name,
                    source_format='doc',
                )

            # Spans normalisieren
            spans:    List[dict] = []
            headings: List[str]  = []
            for s in (pdf_result.spans or []):
                text = (s.text or '').strip()
                if not text:
                    continue
                bold   = bool(getattr(s, 'bold',   False))
                size   = float(getattr(s, 'size',  DEFAULT_SIZE))
                is_hdg = bold and size >= HEADING_SIZE_MIN
                span = {
                    'page':   int(getattr(s, 'page', 1)),
                    'y':      int(getattr(s, 'y',    0)),
                    'x':      int(getattr(s, 'x',    0)),
                    'size':   size,
                    'bold':   bold,
                    'italic': bool(getattr(s, 'italic', False)),
                    'font':   str(getattr(s, 'font',   'Arial')),
                    'text':   text,
                    'width':  float(getattr(s, 'x1', 0) or 0)
                               - float(getattr(s, 'x0', 0) or 0),
                }
                spans.append(span)
                if is_hdg:
                    headings.append(text)

            total_duration = round(time.time() - start, 3)
            result = self._build_result(
                spans, [], headings,
                getattr(pdf_result, 'page_count', 1),
                path_obj.name, 'doc', total_duration,
            )
            result.meta['lo_duration_s'] = lo_duration
            result.meta['lo_stderr']     = lo_result.stderr[:200] if lo_result.stderr else ''
            # tmp_dir → automatisch gelöscht durch context manager
            return result

    # ── Paragraph → Span ─────────────────────────────────────────────────────

    def _para_to_span(self, para_element, doc, page: int, y: int):
        try:
            from docx.text.paragraph import Paragraph as DocxParagraph
            para = DocxParagraph(para_element, doc)
        except Exception:
            return None, False

        text = para.text.strip()
        if not text:
            return None, False

        x      = _get_indent_x(para)
        style  = para.style
        size   = _resolve_size(style)
        bold   = _resolve_bold(style) or _is_heading_style(style)
        italic = False
        font   = _resolve_font(style)

        for run in para.runs:
            if not run.text.strip():
                continue
            if run.bold is not None:
                bold = bool(run.bold)
            if run.italic:
                italic = True
            try:
                if run.font.size:
                    size = round(run.font.size.pt, 1)
            except Exception:
                pass
            try:
                if run.font.name:
                    font = run.font.name
            except Exception:
                pass
            break

        if _is_heading_style(style) and size < HEADING_SIZE_MIN:
            size = HEADING_SIZE_MIN
            bold = True

        is_heading = bold and size >= HEADING_SIZE_MIN

        return {
            'page':   page,
            'y':      y,
            'x':      x,
            'size':   size,
            'bold':   bold,
            'italic': italic,
            'font':   font,
            'text':   text,
            'width':  0.0,
        }, is_heading

    # ── Tabelle → Dict ────────────────────────────────────────────────────────

    def _table_to_dict(self, tbl_element, doc, page: int, y: int) -> Optional[dict]:
        try:
            from docx.table import Table as DocxTable
            table = DocxTable(tbl_element, doc)
        except Exception:
            return None

        all_rows:  List[List[str]]  = []
        row_spans: List[List[dict]] = []
        raw_lines: List[str]        = []

        for row in table.rows:
            row_texts:     List[str]  = []
            row_span_list: List[dict] = []
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                row_texts.append(cell_text)
                raw_lines.append(cell_text)
                row_span_list.append({
                    'page':   page,
                    'y':      y,
                    'x':      col_idx * 100,
                    'size':   DEFAULT_SIZE,
                    'bold':   False,
                    'italic': False,
                    'font':   'Arial',
                    'text':   cell_text,
                    'width':  0.0,
                })
            if row_texts:
                all_rows.append(row_texts)
                row_spans.append(row_span_list)

        if not all_rows:
            return None

        return {
            'headers':   all_rows[0],
            'rows':      all_rows[1:] if len(all_rows) > 1 else [],
            'row_spans': row_spans,
            'raw_text':  '\n'.join(raw_lines),
        }

    # ── Gemeinsame Ergebnis-Konstruktion ──────────────────────────────────────

    def _build_result(
        self,
        spans:         List[dict],
        tables:        List[dict],
        headings:      List[str],
        page_count:    int,
        filename:      str,
        source_format: str,
        duration:      float,
    ) -> WordExtractResult:

        plain_text = '\n'.join(
            s['text'] for s in spans if s.get('text', '').strip()
        )
        if len(plain_text) > MAX_PLAIN_TEXT_CHARS:
            plain_text = plain_text[:MAX_PLAIN_TEXT_CHARS]

        char_count = len(plain_text.replace(' ', '').replace('\n', ''))

        logger.info(
            f"[MasterWordExtractor] {filename} ({source_format}): "
            f"{len(spans)} Spans | {page_count} Seite(n) | "
            f"{char_count} Zeichen | {len(tables)} Tabellen | "
            f"{duration}s"
        )

        return WordExtractResult(
            ok            = True,
            filename      = filename,
            source_format = source_format,
            page_count    = page_count,
            span_count    = len(spans),
            char_count    = char_count,
            plain_text    = plain_text,
            first_500     = plain_text[:500],
            spans         = spans,
            tables        = tables,
            headings      = headings,
            meta          = {
                'duration_s': duration,
                'suffix':     '.' + source_format,
            },
        )


# ── Singleton ─────────────────────────────────────────────────────────────────

master_word_extractor = MasterWordExtractor()


# ── Schnelltest ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys

    path = sys.argv[1] if len(sys.argv) > 1 else None
    if not path:
        print("Usage: python master_word_extractor.py <datei.docx|datei.doc>")
        sys.exit(1)

    result = master_word_extractor.extract(path)

    if not result.ok:
        print(f"FEHLER: {result.error}")
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"Datei:         {result.filename}")
    print(f"Format:        {result.source_format}")
    print(f"Seiten:        {result.page_count}")
    print(f"Spans:         {result.span_count}")
    print(f"Zeichen:       {result.char_count}")
    print(f"Tabellen:      {len(result.tables)}")
    print(f"Überschriften: {len(result.headings)}")
    if result.meta.get('lo_duration_s'):
        print(f"LO-Konv.:      {result.meta['lo_duration_s']}s")
    print(f"\n--- ERSTE 500 ZEICHEN ---")
    print(result.first_500)
    print(f"\n--- ÜBERSCHRIFTEN ---")
    for h in result.headings[:10]:
        print(f"  • {h}")
    if result.tables:
        print(f"\n--- ERSTE TABELLE ---")
        t = result.tables[0]
        print(f"  Header: {t['headers']}")
        for row in t['rows'][:3]:
            print(f"  Row:    {row}")
    print(f"\n--- SPANS (erste 5) ---")
    for s in result.spans[:5]:
        b = 'B' if s['bold'] else '.'
        i = 'I' if s['italic'] else '.'
        print(f"  p{s['page']:02d}|y={s['y']:5}|x={s['x']:4}|"
              f"sz={s['size']:5.1f}|{b}{i}|{s['text'][:60]}")
    print(f"{'='*60}")
