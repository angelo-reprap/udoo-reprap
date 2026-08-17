"""
word_builder.py - Basis-Klasse für alle Word-Templates.
Enthält alle XML/DOCX Aufbau-Methoden.
Jedes Template-Layout erbt davon und überschreibt nur was anders ist.
"""
import os
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


# ── XML Helpers ───────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def _set_cell_borders(cell, color="E5E7EB", size=4):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement("w:"+side)
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),str(size))
        el.set(qn("w:space"),"0");    el.set(qn("w:color"),color)
        borders.append(el)
    tcPr.append(borders)

def _set_no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement("w:"+side); el.set(qn("w:val"),"none")
        borders.append(el)
    tcPr.append(borders)

def _cell_margin(cell, top=60, bottom=60, left=120, right=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for name, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement("w:"+name)
        el.set(qn("w:w"),str(val)); el.set(qn("w:type"),"dxa")
        mar.append(el)
    tcPr.append(mar)

def _set_table_width(table, width_dxa):
    tbl = table._tbl; tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),str(width_dxa)); tblW.set(qn("w:type"),"dxa")
    tblPr.append(tblW); tbl.insert(0, tblPr)

def _set_col_width(cell, width_dxa):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),str(width_dxa)); tcW.set(qn("w:type"),"dxa")
    tcPr.append(tcW)

def _add_para_border_bottom(para, color="163258", size=6):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),str(size))
    bot.set(qn("w:space"),"4");    bot.set(qn("w:color"),color)
    pBdr.append(bot); pPr.append(pBdr)

def _left_border_only(cell, color="163258", size=12):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","right","bottom"):
        el = OxmlElement("w:"+side); el.set(qn("w:val"),"none")
        borders.append(el)
    left = OxmlElement("w:left")
    left.set(qn("w:val"),"single"); left.set(qn("w:sz"),str(size))
    left.set(qn("w:space"),"4");    left.set(qn("w:color"),color)
    borders.append(left); tcPr.append(borders)

def _left_bar_para(para, color="163258"):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"12")
    left.set(qn("w:space"),"8");    left.set(qn("w:color"),color)
    pBdr.append(left); pPr.append(pBdr)

def _badge_bg(run, color="163258"):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),color);  rPr.append(shd)

def _page_field(para, field_name, font_size=8, color=None):
    from docx.shared import RGBColor as RC
    r = para.add_run()
    for ftype, name in [("begin", field_name), ("end", None)]:
        fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), ftype)
        r._r.append(fld)
        if name:
            instr = OxmlElement("w:instrText"); instr.text = name
            r._r.append(instr)
    r.font.size = Pt(font_size)
    if color: r.font.color.rgb = color
    return r


# ── Base Builder ─────────────────────────────────────────────

class WordBuilder:
    """
    Basis-Klasse. Erbt von hier:
    - AidProfileLayout  (aid-profile)
    - AidShortLayout    (aid-short)
    - CustomLayout      (custom templates)
    """

    def __init__(self, style: dict):
        self.style   = style
        b = style["brand"]; f = style["fonts"]
        self.BLUE     = _hex_to_rgb(b["blue"])
        self.GRAY     = _hex_to_rgb(b["gray"])
        self.BLACK    = _hex_to_rgb(b["black"])
        self.BOX_BG   = b["box_bg"]
        self.BOX_BDR  = b["box_border"]
        self.TECH_BG  = b["tech_bg"]
        self.FONT     = f["body"]
        self.SZ       = f["size_body"]
        self.SZ_SM    = f["size_small"]
        self.SZ_H     = f["size_heading"]
        self.SZ_SH    = f["size_subheading"]
        self.SZ_AID   = f["size_aid"]
        self.SZ_HL    = f["size_headline"]
        self.SZ_HDR   = f["size_header"]
        self.BLUE_HEX = b["blue"]

    # ── Typo helpers ─────────────────────────────────────────

    def run(self, para, text, bold=False, italic=False, size=None, color=None):
        r = para.add_run(text)
        r.font.name = self.FONT
        r.font.size = Pt(size or self.SZ)
        r.bold = bold; r.italic = italic
        r.font.color.rgb = color or self.BLACK
        return r

    def heading(self, doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        self.run(p, text, bold=True, size=self.SZ_H, color=self.BLUE)
        _add_para_border_bottom(p, self.BLUE_HEX)
        return p

    def subheading(self, doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.3)
        self.run(p, text, bold=True, size=self.SZ_SH, color=self.BLUE)
        _left_bar_para(p, self.BLUE_HEX)
        return p

    # ── Header ───────────────────────────────────────────────

    def setup_header(self, section, aid):
        hdr  = section.header
        para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        para.clear()
        para.paragraph_format.space_after = Pt(4)
        hcfg = self.style["header"]
        if hcfg.get("show_aid"):
            self.run(para, "Qualifikationsprofil: " + aid + "   ",
                     size=self.SZ_HDR, color=self.GRAY)
        if hcfg.get("show_url"):
            self.run(para, hcfg.get("url","www.abcona.de") + "   ",
                     size=self.SZ_HDR, color=self.GRAY)
        if hcfg.get("show_page_number"):
            self.run(para, "Seite ", size=self.SZ_HDR, color=self.GRAY)
            _page_field(para, "PAGE",     self.SZ_HDR, self.GRAY)
            self.run(para, " von ",   size=self.SZ_HDR, color=self.GRAY)
            _page_field(para, "NUMPAGES", self.SZ_HDR, self.GRAY)
        _add_para_border_bottom(para, self.BLUE_HEX)

    # ── Page 1 ───────────────────────────────────────────────

    def build_page1(self, doc, data, logo_path):
        pers = data["personal"]; company = data["company"]
        if logo_path and os.path.exists(logo_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(logo_path, width=Cm(8))
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.run(p, "[Logo abcona]", italic=True, color=self.GRAY, size=9)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "active business consulting agency",
                 italic=True, color=self.GRAY, size=9)
        p.paragraph_format.space_after = Pt(24)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, data["aid"], bold=True, size=self.SZ_AID, color=self.BLUE)
        p.paragraph_format.space_after = Pt(8)

        if pers.get("headline"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.run(p, "Schwerpunkt: " + pers["headline"],
                     bold=True, size=self.SZ_HL, color=self.BLUE)
            p.paragraph_format.space_after = Pt(28)

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(tbl, 9638)
        cell = tbl.cell(0,0)
        _left_border_only(cell, self.BLUE_HEX)
        _set_cell_bg(cell, "F8FAFC")
        _cell_margin(cell, top=120, bottom=120, left=300, right=160)

        lines = [
            (company.get("name","abcona e. K."), True,  self.BLUE,  12),
            ("active business consulting agency", False, self.GRAY,  9),
            ("", False, self.BLACK, 2),
            ("Bornhohl 26",                       False, self.BLACK, self.SZ),
            ("61449 Steinbach",                   False, self.BLACK, self.SZ),
            ("", False, self.BLACK, 2),
            ("Telefon +49 (0) 61 71 - 8867 - 00", False, self.BLACK, self.SZ),
            ("Fax +49 (0) 61 71 - 8867 - 09",    False, self.BLACK, self.SZ),
            ("", False, self.BLACK, 2),
            ("E-Mail office@abcona.de",            False, self.BLACK, self.SZ),
            ("Internet http://www.abcona.de",      False, self.BLACK, self.SZ),
        ]
        first = True
        for text, bold, color, size in lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(2)
            if text: self.run(p, text, bold=bold, color=color, size=size)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.run(p, "Stand: " + data["date"],
                 italic=True, color=self.GRAY, size=9)
        p.paragraph_format.space_before = Pt(20)

    # ── Page 2 ───────────────────────────────────────────────

    def build_page2(self, doc, data):
        sec = self.style["sections"]
        self.heading(doc, "Persönliche Daten")

        rows_data = [
            ("Name:", (data["personal"].get("first_name","") + " " +
                       data["personal"].get("last_name","")).strip() or "-"),
            ("Geburtsjahr:",         str(data["personal"].get("birth_year") or "-")),
            ("Staatsangehörigkeit:", data["personal"].get("nationality") or "Deutsch"),
            ("Sprachen:",            ", ".join(data["languages"]) if data["languages"] else "-"),
            ("EDV Erfahrung seit:",  str(data["personal"].get("edv_experience_since") or "-")),
            ("verfügbar:",           data["personal"].get("availability") or "nach Absprache"),
            ("Einsatzort:",          data["personal"].get("location") or "nach Absprache"),
        ]
        tbl = doc.add_table(rows=len(rows_data), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(tbl, 9638)
        for i, (label, value) in enumerate(rows_data):
            lc, vc = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
            _set_no_border(lc); _set_no_border(vc)
            _set_col_width(lc, 2600); _set_col_width(vc, 7038)
            _cell_margin(lc, top=40, bottom=40, left=0,   right=100)
            _cell_margin(vc, top=40, bottom=40, left=100, right=0)
            lp = lc.paragraphs[0]
            self.run(lp, label, bold=True, size=self.SZ)
            _add_para_border_bottom(lp, self.BOX_BDR, size=2)
            vp = vc.paragraphs[0]
            self.run(vp, value, size=self.SZ)
            _add_para_border_bottom(vp, self.BOX_BDR, size=2)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        self.heading(doc, "Ausbildung")
        for e in (data["education"] or [{"period":"-","description":""}]):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            self.run(p, (e.get("period") or "") + "   ",
                     bold=True, color=self.BLUE, size=self.SZ)
            self.run(p, e.get("description") or "", size=self.SZ)

        if sec.get("show_certifications", True):
            self.heading(doc, "Fachbereiche & Zertifizierungen")
            left_items  = data["focus_areas"]    or ["-"]
            right_items = data["certifications"] or ["-"]
            max_rows = max(len(left_items), len(right_items))
            tbl2 = doc.add_table(rows=max_rows+1, cols=2)
            tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
            _set_table_width(tbl2, 9638); half = 4779
            lc, rc = tbl2.rows[0].cells[0], tbl2.rows[0].cells[1]
            for c in (lc, rc): _set_no_border(c); _set_col_width(c, half)
            _cell_margin(lc, top=0, bottom=60, left=0,   right=200)
            _cell_margin(rc, top=0, bottom=60, left=200, right=0)
            self.run(lc.paragraphs[0], "Fachbereiche",    bold=True, size=11, color=self.BLUE)
            self.run(rc.paragraphs[0], "Zertifizierungen", bold=True, size=11, color=self.BLUE)
            for i in range(max_rows):
                lc2, rc2 = tbl2.rows[i+1].cells[0], tbl2.rows[i+1].cells[1]
                for c in (lc2, rc2): _set_no_border(c); _set_col_width(c, half)
                _cell_margin(lc2, top=30, bottom=30, left=0,   right=200)
                _cell_margin(rc2, top=30, bottom=30, left=200, right=0)
                if i < len(left_items):
                    self.run(lc2.paragraphs[0], "· " + left_items[i], size=self.SZ)
                if i < len(right_items):
                    self.run(rc2.paragraphs[0], "· " + right_items[i], size=self.SZ)

        if sec.get("show_courses", True) and data.get("courses"):
            self.heading(doc, "Schulungen / Kurse")
            courses = data["courses"]
            tbl3 = doc.add_table(rows=(len(courses)+1)//2, cols=2)
            tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
            _set_table_width(tbl3, 9638)
            for i, course in enumerate(courses):
                cell = tbl3.rows[i//2].cells[i%2]
                _set_no_border(cell)
                _cell_margin(cell, top=40, bottom=40, left=0, right=100)
                p = cell.paragraphs[0]
                self.run(p, course, size=self.SZ)
                _add_para_border_bottom(p, self.BOX_BDR, size=2)

        if sec.get("show_industries", True) and data.get("industries"):
            self.heading(doc, "Branchen")
            p = doc.add_paragraph()
            self.run(p, ", ".join(data["industries"]), size=self.SZ)
            p.paragraph_format.space_after = Pt(4)

    # ── Experiences ──────────────────────────────────────────

    def build_experiences(self, doc, data):
        sec   = self.style["sections"]
        limit = sec.get("experiences_per_page", 99)
        exps  = data.get("experiences", [])[:limit]
        self.heading(doc, "Berufliche Erfahrungen")
        for exp in exps:
            self._exp_block(doc, exp, show_tech=True)

    def _exp_block(self, doc, exp, show_tech=True):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(tbl, 9638)
        cell = tbl.cell(0,0)
        _set_cell_bg(cell, self.BOX_BG)
        _set_cell_borders(cell, self.BOX_BDR, size=4)
        _cell_margin(cell, top=80, bottom=80, left=180, right=180)

        p = cell.paragraphs[0]
        r_date = p.add_run((exp.get("period") or "") + "  ")
        r_date.font.name = self.FONT; r_date.font.size = Pt(self.SZ_SM)
        r_date.bold = True; r_date.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _badge_bg(r_date, self.BLUE_HEX)
        r_co = p.add_run("  " + (exp.get("company") or ""))
        r_co.font.name = self.FONT; r_co.font.size = Pt(11)
        r_co.bold = True; r_co.font.color.rgb = self.BLACK
        p.paragraph_format.space_after = Pt(4)

        if exp.get("role"):
            p2 = cell.add_paragraph()
            self.run(p2, exp["role"], bold=True, size=self.SZ, color=self.BLUE)
            p2.paragraph_format.space_after = Pt(4)

        for act in (exp.get("activities") or []):
            p3 = cell.add_paragraph()
            self.run(p3, "· " + act, size=self.SZ)
            p3.paragraph_format.left_indent = Cm(0.3)
            p3.paragraph_format.space_after = Pt(2)

        if show_tech and exp.get("technologies"):
            p4 = cell.add_paragraph()
            _set_cell_bg(cell, self.TECH_BG)
            self.run(p4, "Technologien: ", bold=True, size=self.SZ_SM, color=self.BLUE)
            self.run(p4, ", ".join(exp["technologies"]), size=self.SZ_SM)
            p4.paragraph_format.space_after = Pt(2)

        for p in list(cell.paragraphs):
            if not p.text.strip() and p != cell.paragraphs[-1]:
                try: p._element.getparent().remove(p._element)
                except: pass

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Products ─────────────────────────────────────────────

    def build_products(self, doc, data):
        if not self.style["sections"].get("show_products", True): return
        products = data.get("products", [])
        if not products: return
        self.heading(doc, "Produkte | Standards | Erfahrungen")
        rows = (len(products)+1)//2
        tbl  = doc.add_table(rows=rows, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(tbl, 9638)
        for i, product in enumerate(products):
            cell = tbl.rows[i//2].cells[i%2]
            _set_no_border(cell); _set_cell_bg(cell, self.BOX_BG)
            _cell_margin(cell, top=40, bottom=40, left=100, right=100)
            p = cell.paragraphs[0]
            _left_bar_para(p, self.BLUE_HEX)
            p.paragraph_format.left_indent = Cm(0.2)
            self.run(p, product, size=self.SZ_SM)

    # ── Skills ───────────────────────────────────────────────

    def build_skills(self, doc, data):
        if not self.style["sections"].get("show_skills", True): return
        sections = data.get("skills_sections", [])
        if not sections: return
        self.heading(doc, "Technische Kenntnisse")
        for sec in sections:
            if not sec.get("skills"): continue
            self.subheading(doc, sec["name"])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_after = Pt(4)
            self.run(p, " · ".join(sec["skills"]), size=self.SZ)

    # ── Document orchestration ───────────────────────────────

    def build_document(self, doc, data, logo_path):
        sec   = self.style["sections"]
        order = sec.get("order", ["page1","page2","experiences","products","skills"])

        for section in doc.sections:
            section.page_width    = Cm(21.0)
            section.page_height   = Cm(29.7)
            section.left_margin   = Cm(self.style["layout"]["margin_cm"])
            section.right_margin  = Cm(self.style["layout"]["margin_cm"])
            section.top_margin    = Cm(self.style["layout"]["top_margin_cm"])
            section.bottom_margin = Cm(self.style["layout"]["bottom_margin_cm"])
            self.setup_header(section, data["aid"])

        doc.styles["Normal"].font.name = self.FONT
        doc.styles["Normal"].font.size = Pt(self.SZ)

        for i, step in enumerate(order):
            more = i < len(order) - 1
            if   step == "page1":
                self.build_page1(doc, data, logo_path)
                if sec.get("page_break_after_page1") and more:
                    doc.add_page_break()
            elif step == "page2":
                self.build_page2(doc, data)
                if sec.get("page_break_after_page2") and more:
                    doc.add_page_break()
            elif step == "experiences":
                self.build_experiences(doc, data)
                if sec.get("page_break_after_experiences") and more:
                    doc.add_page_break()
            elif step == "products":
                self.build_products(doc, data)
                if sec.get("page_break_after_products") and more:
                    doc.add_page_break()
            elif step == "skills":
                self.build_skills(doc, data)
        return doc
